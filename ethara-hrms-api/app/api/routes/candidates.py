import concurrent.futures
import logging
import os
import re
import shutil
import subprocess
import tempfile
from datetime import UTC, date, datetime, time
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from typing import Annotated

from fastapi import (
    APIRouter,
    BackgroundTasks,
    Depends,
    File,
    Form,
    HTTPException,
    Query,
    Request,
    UploadFile,
    status,
)
from fastapi.encoders import jsonable_encoder
from fastapi.responses import FileResponse, JSONResponse, RedirectResponse
from pydantic import EmailStr, TypeAdapter, ValidationError
from sqlalchemy import select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from app.api.deps import get_current_user, require_permissions, user_has_any_role
from app.core.config import get_settings
from app.core.database import SessionLocal, get_db
from app.core.limiter import limiter
from app.core.permissions import Permission
from app.core.timezone import app_date_stamp, format_app_datetime
from app.db.models import (
    AdminSetting,
    ApAssignment,
    ApAttempt,
    Candidate,
    CandidateStage,
    ContractStatus,
    Document,
    Role,
    User,
)
from app.schemas.candidate import (
    AdvanceStageRequest,
    CandidateDetail,
    CandidateListResponse,
    CandidatePortalSelfOverview,
    CandidateStatsResponse,
    CandidateSummary,
    CreateCandidateRequest,
    PortalApplyRequest,
    PortalProfileUpdateRequest,
    UpdateCandidateRequest,
    scrub_candidate_preview_fields,
    scrub_internal_candidate_fields,
)
from app.schemas.common import MessageResponse
from app.schemas.resources import validate_password_strength
from app.schemas.workflow import CandidateIdCardFormRead, CandidateIdCardFormSubmitRequest
from app.services import account_security
from app.services import candidates as candidate_service
from app.services import vertex_ai
from app.services.audit import log_audit
from app.services import workflows

router = APIRouter(prefix="/candidates", tags=["candidates"])
logger = logging.getLogger(__name__)


def _run_registration_screening_background(candidate_id: str) -> None:
    with SessionLocal() as screening_db:
        try:
            workflows.run_resume_screening(screening_db, candidate_id=candidate_id)
            screening_db.commit()
        except Exception as exc:  # noqa: BLE001 - registration response already succeeded
            logger.warning(
                "Post-registration resume screening failed for candidate %s: %s",
                candidate_id,
                exc,
            )
            try:
                screening_db.rollback()
                candidate = screening_db.get(Candidate, candidate_id)
                if candidate is None:
                    return
                payload = dict(candidate.screening_payload or {})
                payload.update(
                    {
                        "status": "failed",
                        "lastScreeningErrorAt": datetime.now(UTC).isoformat(),
                    }
                )
                candidate.screening_payload = payload
                candidate.llm_status = "failed"
                screening_db.add(candidate)
                screening_db.commit()
            except Exception as mark_exc:  # noqa: BLE001 - keep background failure isolated
                try:
                    screening_db.rollback()
                except Exception:
                    pass
                logger.warning(
                    "Unable to mark post-registration screening failed for candidate %s: %s",
                    candidate_id,
                    mark_exc,
                )

# DoS guards for the public registration / OCR path.
def _bounded_int_env(name: str, default: int, *, minimum: int, maximum: int) -> int:
    try:
        value = int(os.getenv(name, str(default)))
    except (TypeError, ValueError):
        return default
    return max(minimum, min(maximum, value))


_MAX_OCR_PAGES = 15  # cap per-document OCR pages; real resumes/ID docs are only a few pages
_TESSERACT_THREAD_LIMIT = _bounded_int_env("TESSERACT_THREAD_LIMIT", 1, minimum=1, maximum=4)
_TESSERACT_DIRECT_TIMEOUT_SECONDS = _bounded_int_env(
    "TESSERACT_DIRECT_TIMEOUT_SECONDS", 10, minimum=2, maximum=30
)
os.environ.setdefault("OMP_THREAD_LIMIT", str(_TESSERACT_THREAD_LIMIT))
os.environ.setdefault("OMP_NUM_THREADS", str(_TESSERACT_THREAD_LIMIT))
try:
    from PIL import Image as _PILImage

    # Reject decompression-bomb images early. PIL's default (~89MP) raises only above 2x;
    # ~50MP comfortably covers legitimate photos/scans while bounding memory.
    _PILImage.MAX_IMAGE_PIXELS = 50_000_000
except Exception:  # pragma: no cover - Pillow always present, but never hard-fail import
    pass

PORTAL_ID_CARD_ALLOWED_STAGES = {
    CandidateStage.CONTRACT_SIGNED,
    CandidateStage.INDUCTION_COMPLETED,
    CandidateStage.IT_EMAIL_CREATED,
    CandidateStage.WELCOME_MAIL_SENT,
    CandidateStage.STATUTORY_FORMS_SENT,
    CandidateStage.STATUTORY_FORMS_SUBMITTED,
    CandidateStage.COMPLIANCE_VERIFIED,
    CandidateStage.ONBOARDING_COMPLETED,
}


def _date_range_bounds(
    created_from: date | None,
    created_to: date | None,
) -> tuple[datetime | None, datetime | None]:
    if created_from and created_to and created_to < created_from:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="createdTo must be on or after createdFrom.",
        )
    start = datetime.combine(created_from, time.min, tzinfo=UTC) if created_from else None
    end = datetime.combine(created_to, time.max, tzinfo=UTC) if created_to else None
    return start, end


def _require_candidate_user(current_user: User) -> User:
    if not user_has_any_role(current_user, {Role.CANDIDATE}):
        raise HTTPException(status_code=403, detail="Candidate account required")
    return current_user


# Roles that may read their own scoped candidates but must NOT see internal recruiting
# data (resume scores/summaries, screening payloads, evaluation notes/scores, audit
# logs, escalations, CTC/salary).
_RESTRICTED_CANDIDATE_VIEWER_ROLES = (Role.VENDOR, Role.EMPLOYEE_REFERRER)
_FULL_CANDIDATE_DETAIL_ROLES = (Role.SUPER_ADMIN, Role.ADMIN, Role.HR, Role.TA)


def _is_restricted_candidate_viewer(current_user: User) -> bool:
    return (
        user_has_any_role(current_user, _RESTRICTED_CANDIDATE_VIEWER_ROLES)
        and not _has_full_candidate_detail_access(current_user)
    )


def _has_full_candidate_detail_access(current_user: User) -> bool:
    return user_has_any_role(current_user, _FULL_CANDIDATE_DETAIL_ROLES)


def _is_preview_candidate_viewer(current_user: User) -> bool:
    return not _has_full_candidate_detail_access(current_user) and not _is_restricted_candidate_viewer(current_user)


def _require_full_candidate_detail_access(current_user: User) -> None:
    if _has_full_candidate_detail_access(current_user) or _is_restricted_candidate_viewer(current_user):
        return
    raise HTTPException(
        status_code=status.HTTP_403_FORBIDDEN,
        detail="Only Admin, HR, and TA users can open full candidate details.",
    )


def _require_candidate_management_access(current_user: User) -> None:
    if _has_full_candidate_detail_access(current_user):
        return
    raise HTTPException(
        status_code=status.HTTP_403_FORBIDDEN,
        detail="Only Admin, HR, and TA users can manage candidate details.",
    )


def _require_id_card_portal_access(candidate: object) -> None:
    current_stage = getattr(candidate, "current_stage", None)
    contract = getattr(candidate, "contract", None)
    contract_status = getattr(contract, "status", None)
    if current_stage not in PORTAL_ID_CARD_ALLOWED_STAGES or contract_status != ContractStatus.SIGNED:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="ID card details are available after your contract and NDA are signed.",
        )


# Matches Aadhaar numbers: 12-digit run, or groups of 4 separated by space/hyphen/dot
# Also handles common OCR errors where digits may be separated by non-standard whitespace
AADHAAR_PATTERN = re.compile(
    r"(?<!\d)"
    r"(?:\d{4}[ \t\-\.]{0,2}\d{4}[ \t\-\.]{0,2}\d{4}"  # 4-4-4 with flexible inline separator
    r"|\d{12})"  # or solid 12-digit
    r"(?!\d)"
)
AADHAAR_CONTEXT_PATTERNS: tuple[tuple[re.Pattern[str], int], ...] = (
    (re.compile(r"\baadhaar\b", re.IGNORECASE), 60),
    (re.compile(r"\buidai\b", re.IGNORECASE), 45),
    (re.compile(r"\bunique\s+identification\b", re.IGNORECASE), 35),
    (re.compile(r"\bgovernment\s+of\s+india\b", re.IGNORECASE), 30),
    (re.compile(r"\b(?:dob|d0b|date\s+of\s+birth|year\s+of\s+birth)\b", re.IGNORECASE), 30),
    (re.compile(r"\b(?:male|female|transgender)\b", re.IGNORECASE), 15),
)
AADHAAR_NON_NUMBER_CONTEXT = re.compile(
    r"\b(?:vid|virtual\s+id|enrolment|enrollment|eid|mobile|phone|ifsc|account|pan)\b",
    re.IGNORECASE,
)
PAN_PATTERN = re.compile(
    r"(?<![A-Z0-9])"
    r"([A-Z]{5})[\s\-\.]*([0-9]{4})[\s\-\.]*([A-Z])"
    r"(?![A-Z0-9])",
    re.IGNORECASE,
)
PAN_COMPACT_PATTERN = re.compile(r"[A-Z]{5}[0-9]{4}[A-Z]")
PAN_ALNUM_WINDOW_PATTERN = re.compile(r"(?=([A-Z0-9]{10}))")
PAN_LETTER_OCR_MAP = str.maketrans({"0": "O", "1": "I", "2": "Z", "5": "S", "8": "B"})
PAN_DIGIT_OCR_MAP = str.maketrans({"O": "0", "I": "1", "L": "1", "S": "5", "B": "8", "Z": "2"})
PAN_CONTEXT_TERMS = (
    "PERMANENTACCOUNT",
    "ACCOUNTNUMBERCARD",
    "INCOMETAX",
    "TAXDEPARTMENT",
    "GOVTOFINDIA",
    "GOVERNMENTOFINDIA",
    "PANCARD",
)
PAN_FUZZY_FALSE_PREFIXES = {
    "ACCOU",
    "BIRTH",
    "BRANC",
    "DATED",
    "IFSCO",
    "SALBR",
    "VALID",
}
ADDRESS_POSTAL_RE = re.compile(r"\b[1-9][0-9]{5}\b")
ADDRESS_INLINE_RE = re.compile(r"\baddress\s*[:\-]?\s*(.+)$", re.IGNORECASE)
ADDRESS_STOP_WORDS = {
    "aadhaar",
    "dob",
    "date of birth",
    "government of india",
    "income tax",
    "permanent account",
    "ifsc",
    "account number",
    "cancelled",
    "signature",
    "vid",
    "uidai",
}

# Matches DOB in multiple formats found on Aadhaar cards:
#   "DOB: 01/01/1990", "Date of Birth: 1-Jan-1990", "Year of Birth: 1990"
#   "जन्म तिथि: 01/01/1990" (Hindi), plain DD/MM/YYYY near "birth"
DOB_PATTERN = re.compile(
    r"(?:"
    r"(?:date\s*of\s*birth|d\s*\.?\s*[o0]\s*\.?\s*b\.?|birth\s*date|जन्म\s*(?:तिथि|दिनांक))"
    r"\s*[:\-]?\s*"
    r"([0-9OoIl|SsBbZz]{1,2}[/\-\.\s]+[0-9OoIl|SsBbZz]{1,2}[/\-\.\s]+[0-9OoIl|SsBbZz]{4}"  # DD/MM/YYYY
    r"|[0-9OoIl|SsBbZz]{4}[/\-\.\s]+[0-9OoIl|SsBbZz]{1,2}[/\-\.\s]+[0-9OoIl|SsBbZz]{1,2}"  # YYYY/MM/DD
    r"|\d{1,2}[\s\-](?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s\-]\d{4}"  # 01-Jan-1990
    r")"
    r"|(?:year\s*of\s*birth|yob)\s*[:\-]?\s*([0-9OoIl|SsBbZz]{4})"  # Year of Birth: 1990
    r")",
    re.IGNORECASE,
)

# Additional standalone date pattern — catches dates on Aadhaar where label may not be OCR'd
_STANDALONE_DATE_PATTERN = re.compile(
    r"\b([0-9OoIl|SsBbZz]{1,2})[/\-\.\s]+([0-9OoIl|SsBbZz]{1,2})[/\-\.\s]+([0-9OoIl|SsBbZz]{4})\b"
)
TEXT_SUFFIXES = {".txt", ".csv", ".md", ".json", ".xml"}
TEXT_CONTENT_TYPES = {"text/plain", "text/csv", "application/json", "application/xml"}
IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}
IMAGE_CONTENT_TYPES = {
    "image/jpeg",
    "image/png",
    "image/webp",
    "image/bmp",
    "image/tiff",
}
OCR_LANGUAGE = "eng"
OCR_DPI = 300
PYMUPDF_FILE_TYPES = {
    ".pdf": "pdf",
    ".jpg": "jpeg",
    ".jpeg": "jpeg",
    ".png": "png",
    ".webp": "webp",
    ".bmp": "bmp",
    ".tif": "tiff",
    ".tiff": "tiff",
}


# Public-registration upload allowlists. Resumes accept documents; Aadhaar/cheque
# accept images or PDFs. Mirrors the employee-registration allowlist so the public
# self-registration path is held to the same content-type/extension discipline.
ALLOWED_RESUME_CONTENT_TYPES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
ALLOWED_RESUME_EXTENSIONS = {".pdf", ".doc", ".docx"}
ALLOWED_DOCUMENT_IMAGE_CONTENT_TYPES = {
    "application/pdf",
    "image/jpeg",
    "image/png",
    "image/webp",
}
ALLOWED_DOCUMENT_IMAGE_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".webp"}


def _max_upload_bytes() -> int:
    return get_settings().max_upload_size_mb * 1024 * 1024


def validate_candidate_upload(
    file: UploadFile | None,
    *,
    label: str,
    allowed_content_types: set[str],
    allowed_extensions: set[str],
    required: bool = True,
) -> None:
    """Reject uploads whose declared content-type/extension is not on the allowlist,
    or whose body exceeds the configured size cap. Used by the public self-registration
    path so a caller cannot smuggle executables/scripts or oversize bodies past the
    file handling. Validation only — does not consume the stream for later readers."""
    if file is None or not file.filename:
        if required:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail=f"{label} upload is required.",
            )
        return

    content_type = (file.content_type or "").split(";")[0].strip().lower()
    suffix = Path(file.filename or "").suffix.lower()
    if content_type not in allowed_content_types or suffix not in allowed_extensions:
        allowed = ", ".join(sorted(ext.lstrip(".").upper() for ext in allowed_extensions))
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail=f"Invalid {label} — only the following file types are allowed: {allowed}.",
        )

    size = _enforce_upload_size(file, label=label)
    if size == 0:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail=f"Invalid {label} — the file is empty.",
        )


def _enforce_upload_size(file: UploadFile, *, label: str = "file") -> int:
    """Reject an upload whose body exceeds the configured cap, then rewind so later
    readers (OCR/parse/storage) still see the full content. Returns the byte length."""
    max_bytes = _max_upload_bytes()
    file.file.seek(0, os.SEEK_END)
    size = file.file.tell()
    file.file.seek(0)
    if size > max_bytes:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail=f"{label} exceeds the maximum upload size of {get_settings().max_upload_size_mb} MB.",
        )
    return size


def save_candidate_upload(file: UploadFile | None, folder: str) -> str | None:
    if file is None or not file.filename:
        return None
    from app.services.integrations import StorageService

    file_url, _storage_path = StorageService().save_upload(file, folder=folder)
    file.file.seek(0)
    return file_url


def record_candidate_upload_document(
    db: Session,
    *,
    candidate: Candidate,
    file: UploadFile | None,
    type_: str,
    file_url: str | None,
    ocr_result: dict | None = None,
) -> None:
    if file is None or not file.filename or not file_url:
        return
    file_size: int | None = None
    try:
        file.file.seek(0, os.SEEK_END)
        file_size = file.file.tell()
    except Exception:
        file_size = getattr(file, "size", None)
    finally:
        file.file.seek(0)
    document = Document(
        candidate_id=candidate.id,
        type=type_,
        file_name=file.filename,
        file_url=file_url,
        file_size=file_size,
        mime_type=file.content_type,
        status="uploaded",
    )
    # When Vertex AI produced a verdict, persist it so HR sees a "needs_review"
    # flag on documents whose type didn't match what was uploaded. Only populated
    # on the AI path (ocr_result carries "verification"); local-OCR uploads keep
    # the default ocr_status="pending", so behaviour is unchanged when AI is off.
    verification = (ocr_result or {}).get("verification")
    if verification:
        document.ocr_provider = "vertex_gemini"
        # Persist only the type-verification verdict (detected type, match,
        # confidence, issues, notes) — NOT extracted_fields, which can hold the
        # raw Aadhaar number. The number is handled by validate_aadhaar_identity
        # (hash/last4 only) and is never stored in full here.
        document.llm_extracted_data = {
            k: v for k, v in verification.items() if k != "extracted_fields"
        }
        document.ocr_status = (
            "needs_review"
            if ocr_result.get("matchesExpectedCategory") is False
            else (ocr_result.get("ocrStatus") or "extracted")
        )
    db.add(document)


def read_upload_bytes(file: UploadFile) -> bytes:
    # Enforce the configured size cap before reading the body into memory so an
    # oversize upload cannot be slurped by the OCR/parse/register read path.
    _enforce_upload_size(file)
    content = file.file.read()
    file.file.seek(0)
    return content


def _join_ocr_text_passes(text_passes: list[str]) -> str:
    return "\n".join(text.strip() for text in text_passes if text and text.strip())


_EMAIL_ADAPTER = TypeAdapter(EmailStr)


def _validate_registration_email(value: str) -> str:
    """Validate + normalize an email from a public registration form.

    The self-service registration endpoints accept ``personalEmail`` as a raw
    ``Form(str)`` — unlike the admin JSON paths, which are typed ``EmailStr`` — so a
    non-address like ``"shubham"`` would be stored verbatim. That junk later fails
    response-model email validation and 500s the candidate/user list endpoints (a
    whole page goes blank because of one bad row). Reject malformed input with 422
    BEFORE any candidate/portal-user record is created, and normalise to a trimmed,
    lower-cased address so duplicate checks stay consistent.
    """
    normalized = (value or "").strip().lower()
    try:
        return _EMAIL_ADAPTER.validate_python(normalized)
    except ValidationError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="Enter a valid email address.",
        ) from exc


def _strip_nul(text: str) -> str:
    """Remove NUL (0x00) bytes from extracted text.

    PostgreSQL text/varchar columns cannot store NUL bytes — an INSERT with one
    raises ``psycopg.DataError`` and aborts the whole transaction (candidate
    registration 500s). PDF/OCR extraction occasionally emits stray NULs, so scrub
    them before the value reaches the DB. Other text is left untouched.
    """
    return text.replace("\x00", "") if text else text


def _extract_resume_text(content: bytes, suffix: str | None, content_type: str | None) -> str:
    """Extract resume text and scrub NUL bytes that Postgres text columns reject."""
    return _strip_nul(_extract_resume_text_impl(content, suffix, content_type))


def _extract_resume_text_impl(content: bytes, suffix: str | None, content_type: str | None) -> str:
    """Extract resume text from text PDFs, scanned PDFs/images, or DOC/DOCX files."""
    suffix = (suffix or "").lower()
    ct = (content_type or "").lower()
    if suffix == ".pdf" or ct == "application/pdf":
        text = extract_pdf_text(content)
        if not text.strip():
            try:
                text = extract_text_with_pymupdf_ocr(content, filetype="pdf", full_ocr=True)
            except Exception:
                text = ""
        if not text.strip():
            try:
                text = _join_ocr_text_passes(
                    _extract_text_with_rapidocr(content, image_upload=False, max_pages=5)
                )
            except Exception:
                text = ""
        return text
    if suffix in IMAGE_SUFFIXES or ct in IMAGE_CONTENT_TYPES:
        filetype = PYMUPDF_FILE_TYPES.get(suffix)
        if not filetype and ct in IMAGE_CONTENT_TYPES:
            filetype = "jpeg" if ct == "image/jpeg" else ct.split("/", maxsplit=1)[1]
        text = ""
        if filetype:
            try:
                text = extract_text_with_pymupdf_ocr(content, filetype=filetype, full_ocr=True)
            except Exception:
                text = ""
        if not text.strip():
            try:
                text = _join_ocr_text_passes(_extract_text_with_rapidocr(content, image_upload=True))
            except Exception:
                text = ""
        return text
    if suffix in {".doc", ".docx"} or "word" in ct or "officedocument" in ct:
        try:
            import docx as _docx
            doc = _docx.Document(BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            try:
                return content.decode("utf-8", errors="ignore")
            except Exception:
                return ""
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def is_image_upload(file: UploadFile) -> bool:
    suffix = Path(file.filename or "").suffix.lower()
    content_type = (file.content_type or "").split(";")[0].lower()
    return suffix in IMAGE_SUFFIXES or content_type in IMAGE_CONTENT_TYPES


def resolve_pymupdf_filetype(file: UploadFile) -> str | None:
    suffix = Path(file.filename or "").suffix.lower()
    content_type = (file.content_type or "").split(";")[0].lower()

    if suffix in PYMUPDF_FILE_TYPES:
        return PYMUPDF_FILE_TYPES[suffix]
    if content_type == "application/pdf":
        return "pdf"
    if content_type in IMAGE_CONTENT_TYPES:
        if content_type == "image/jpeg":
            return "jpeg"
        return content_type.split("/", maxsplit=1)[1]
    return None


def resolve_tesseract_command() -> str | None:
    settings = get_settings()
    if settings.tesseract_command:
        return settings.tesseract_command
    return shutil.which("tesseract")


@lru_cache
def get_available_tesseract_languages() -> set[str]:
    command = resolve_tesseract_command()
    if not command:
        return set()
    try:
        result = subprocess.run(
            [command, "--list-langs"],
            capture_output=True,
            text=True,
            check=False,
            timeout=10,
        )
    except Exception:
        return set()
    if result.returncode != 0:
        return set()
    return {
        line.strip()
        for line in result.stdout.splitlines()
        if line.strip() and not line.startswith("List of available languages")
    }


def get_ocr_languages() -> str:
    settings = get_settings()
    configured = [language.strip() for language in (settings.ocr_languages or "eng").split("+")]
    available = get_available_tesseract_languages()
    selected = [language for language in configured if language and language in available]
    return "+".join(selected) if selected else "eng"


def get_ocr_dpi() -> int:
    settings = get_settings()
    return settings.ocr_dpi or OCR_DPI


def is_local_ocr_runtime_available() -> bool:
    if not resolve_tesseract_command():
        return False
    try:
        import pymupdf  # noqa: F401
    except ImportError:
        return False
    return True


def build_aadhaar_ocr_message(*, extracted: bool, image_upload: bool) -> str:
    if extracted:
        return "Extracted readable Aadhaar details."
    if image_upload and is_local_ocr_runtime_available():
        return (
            "We could not clearly read enough Aadhaar details from this image. "
            "Try a brighter, tightly cropped front-side photo, or continue by entering your Aadhaar number manually."
        )
    if image_upload:
        return (
            "We could not auto-read this Aadhaar image on the current server setup. "
            "You can continue by entering your Aadhaar number manually, and the document will be reviewed."
        )
    return (
        "Could not extract Aadhaar details from this file. "
        "The profile can still be created for manual verification."
    )


def _preprocess_image_for_ocr(content: bytes) -> bytes:
    """
    Enhance image contrast/sharpness so Tesseract reads Aadhaar digits more reliably.
    Requires Pillow. Falls back to original bytes if unavailable.
    """
    try:
        from io import BytesIO
        from PIL import Image, ImageEnhance, ImageFilter, ImageOps
        img = ImageOps.exif_transpose(Image.open(BytesIO(content))).convert("L")  # greyscale
        # Sharpen edges — helps with embossed/blurred card scans
        img = img.filter(ImageFilter.SHARPEN)
        # Boost contrast to make digits stand out
        img = ImageEnhance.Contrast(img).enhance(2.0)
        # Scale up if the image is small (better for Tesseract)
        w, h = img.size
        if max(w, h) < 1500:
            scale = 1500 / max(w, h)
            img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        buf = BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return content


def _document_image_ocr_payloads(content: bytes) -> list[bytes]:
    """Build a small set of document image variants for OCR.

    Phone uploads often carry EXIF orientation, compression, shadows, or a
    sideways card. RapidOCR is fast enough that a few variants make extraction
    more reliable without requiring the tesseract binary.
    """
    payloads: list[bytes] = []
    seen: set[bytes] = set()

    def add_payload(data: bytes) -> None:
        if data and data not in seen:
            seen.add(data)
            payloads.append(data)

    add_payload(content)
    try:
        from PIL import Image, ImageEnhance, ImageFilter, ImageOps

        base = ImageOps.exif_transpose(Image.open(BytesIO(content))).convert("RGB")

        def encode(image: Image.Image) -> bytes:
            buf = BytesIO()
            image.save(buf, format="PNG")
            return buf.getvalue()

        def scale_for_ocr(image: Image.Image, target: int = 2200) -> Image.Image:
            width, height = image.size
            longest = max(width, height)
            if longest >= target:
                return image
            scale = target / max(longest, 1)
            return image.resize((int(width * scale), int(height * scale)), Image.LANCZOS)

        base = scale_for_ocr(base)
        add_payload(encode(base))

        gray = ImageOps.grayscale(base)
        enhanced = ImageOps.autocontrast(gray)
        enhanced = enhanced.filter(ImageFilter.SHARPEN)
        enhanced = ImageEnhance.Contrast(enhanced).enhance(2.4)
        add_payload(encode(enhanced))

        threshold = enhanced.point(lambda pixel: 255 if pixel > 150 else 0)
        add_payload(encode(threshold))

        if base.height > base.width:
            for angle in (90, 270):
                add_payload(encode(enhanced.rotate(angle, expand=True)))
    except Exception:
        processed = _preprocess_image_for_ocr(content)
        add_payload(processed)

    return payloads[:8]


def _aadhaar_image_ocr_payloads(content: bytes) -> list[bytes]:
    return _document_image_ocr_payloads(content)


def _pan_image_ocr_payloads(content: bytes) -> list[bytes]:
    return _document_image_ocr_payloads(content)


def extract_text_with_pymupdf_ocr(
    content: bytes,
    *,
    filetype: str,
    full_ocr: bool,
) -> str:
    if not content or not filetype:
        return ""

    try:
        import pymupdf
    except ImportError:
        return ""

    # For image uploads (non-PDF), pre-process to improve OCR accuracy
    processed_content = content
    if filetype in ("jpeg", "png", "webp", "bmp", "tiff"):
        processed_content = _preprocess_image_for_ocr(content)

    try:
        document = pymupdf.open(stream=processed_content, filetype=filetype)
    except Exception:
        # Fallback to original if preprocessing changed the format
        try:
            document = pymupdf.open(stream=content, filetype=filetype)
        except Exception:
            return ""

    with document:
        chunks: list[str] = []
        try:
            page_count = document.page_count
        except Exception:
            return ""

        # Bound per-page OCR work so a many-page (but small) PDF can't pin worker CPUs on
        # the public registration endpoint (DoS). Real resumes/ID docs are a few pages.
        page_count = min(page_count, _MAX_OCR_PAGES)
        for page_number in range(page_count):
            try:
                page = document.load_page(page_number)
                textpage = page.get_textpage_ocr(
                    language=get_ocr_languages(),
                    dpi=get_ocr_dpi(),
                    full=full_ocr,
                )
                text = page.get_text("text", textpage=textpage, sort=True).strip()
            except Exception:
                text = ""
            if text:
                chunks.append(text)
        return "\n".join(chunks)


def extract_pdf_text(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(BytesIO(content))
    except Exception:
        return ""
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_readable_text(file: UploadFile) -> str:
    content = read_upload_bytes(file)
    suffix = Path(file.filename or "").suffix.lower()
    content_type = (file.content_type or "").split(";")[0].lower()

    if suffix in TEXT_SUFFIXES or content_type in TEXT_CONTENT_TYPES:
        return content.decode("utf-8", errors="ignore")

    if suffix == ".pdf" or content_type == "application/pdf":
        text = extract_pdf_text(content)
        if extract_aadhaar_numbers_from_text(text) or parse_date_of_birth(text):
            return text
        return extract_text_with_pymupdf_ocr(content, filetype="pdf", full_ocr=not bool(text.strip()))

    if suffix in IMAGE_SUFFIXES or content_type in IMAGE_CONTENT_TYPES:
        filetype = resolve_pymupdf_filetype(file)
        return extract_text_with_pymupdf_ocr(content, filetype=filetype or "", full_ocr=True)

    return ""


_OCR_DIGIT_MAP = str.maketrans("OoIlSsBbZz", "0011558820")


def _normalize_ocr_digits(text: str) -> str:
    """Fix common OCR character substitutions in digit sequences."""
    # Replace common letter-for-digit errors in runs that look like numbers
    return re.sub(
        r"[0-9OoIlSsBbZz]{4,}",
        lambda m: m.group(0).translate(_OCR_DIGIT_MAP),
        text,
    )


def _aadhaar_number_candidates_from_text(text: str) -> dict[str, dict[str, int]]:
    """Return Aadhaar-like numbers with context scores.

    OCR can surface multiple 12-digit values from QR, receipt, or reference
    text. Keep every plausible candidate, but rank numbers near Aadhaar/DOB/UIDAI
    context above unrelated numeric IDs.
    """
    combined = text + "\n" + _normalize_ocr_digits(text)
    candidates: dict[str, dict[str, int]] = {}
    for match in AADHAAR_PATTERN.finditer(combined):
        digits = re.sub(r"\D", "", match.group(0))
        if len(digits) != 12:
            continue

        start, end = match.span()
        line_start = combined.rfind("\n", 0, start) + 1
        line_end = combined.find("\n", end)
        if line_end == -1:
            line_end = len(combined)
        line = combined[line_start:line_end].lower()
        prefix = combined[max(0, start - 48) : start].lower()

        if re.search(r"\b(?:vid|virtual\s+id)\b", line) or re.search(
            r"\b(?:vid|virtual\s+id)\b", prefix
        ):
            continue

        context = combined[max(0, start - 180) : min(len(combined), end + 180)]
        score = 10
        for pattern, weight in AADHAAR_CONTEXT_PATTERNS:
            if pattern.search(context):
                score += weight
            if pattern.search(line):
                score += weight

        if AADHAAR_NON_NUMBER_CONTEXT.search(line):
            score -= 50

        try:
            from app.services.ocr import is_valid_verhoeff

            if is_valid_verhoeff(digits):
                score += 20
        except Exception:
            pass

        entry = candidates.setdefault(
            digits,
            {"score": 0, "count": 0, "first_index": start},
        )
        entry["score"] += score
        entry["count"] += 1
        entry["first_index"] = min(entry["first_index"], start)
    return candidates


def _rank_aadhaar_numbers_from_text(text: str) -> list[str]:
    candidates = _aadhaar_number_candidates_from_text(text)
    return [
        number
        for number, _meta in sorted(
            candidates.items(),
            key=lambda item: (
                -item[1]["score"],
                -item[1]["count"],
                item[1]["first_index"],
                item[0],
            ),
        )
    ]


def extract_aadhaar_numbers_from_text(text: str) -> set[str]:
    return set(_aadhaar_number_candidates_from_text(text))


def _parse_date_parts(day: int, month: int, year: int) -> str | None:
    try:
        return datetime(year, month, day, tzinfo=UTC).date().isoformat()
    except ValueError:
        return None


_OCR_DATE_DIGIT_TRANS = str.maketrans({
    "O": "0",
    "o": "0",
    "I": "1",
    "l": "1",
    "|": "1",
    "S": "5",
    "s": "5",
    "B": "8",
    "b": "8",
    "Z": "2",
    "z": "2",
})


def _normalize_ocr_date_digits(value: str) -> str:
    return value.translate(_OCR_DATE_DIGIT_TRANS)


def parse_date_of_birth(text: str) -> str | None:
    match = DOB_PATTERN.search(text)
    if not match:
        return None

    # "Year of Birth" match — group(2)
    if match.lastindex and match.lastindex >= 2 and match.group(2):
        year = int(_normalize_ocr_date_digits(match.group(2)))
        if 1900 < year < 2025:
            return f"{year}-01-01"  # approximate — day/month unknown
        return None

    raw = match.group(1)
    if not raw:
        return None

    # Named-month format e.g. "01-Jan-1990"
    month_map = {
        "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
        "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
    }
    named = re.match(r"(\d{1,2})[\s\-]([a-z]+)[\s\-](\d{4})", raw, re.IGNORECASE)
    if named:
        mon_str = named.group(2).lower()[:3]
        if mon_str in month_map:
            return _parse_date_parts(int(named.group(1)), month_map[mon_str], int(named.group(3)))

    # Numeric separator formats
    raw_clean = _normalize_ocr_date_digits(raw)
    raw_clean = re.sub(r"[.\-\s]+", "/", raw_clean)
    raw_clean = re.sub(r"/+", "/", raw_clean).strip("/")
    parts = raw_clean.split("/")
    if len(parts) != 3:
        return None
    try:
        p0, p1, p2 = int(parts[0]), int(parts[1]), int(parts[2])
    except ValueError:
        return None

    # YYYY/MM/DD
    if p0 > 31:
        return _parse_date_parts(p2, p1, p0)
    # DD/MM/YYYY (standard Aadhaar format)
    return _parse_date_parts(p0, p1, p2)


def normalize_aadhaar_date_of_birth(value: str | None) -> str | None:
    if not value:
        return None
    return parse_date_of_birth(f"DOB: {value}") or _fallback_dob_from_standalone(value) or value


def parse_iso_date(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value).replace(tzinfo=UTC)
    except ValueError:
        return None


def parse_experience_years(value: str | None, *, required: bool) -> int | None:
    if value is None or not value.strip():
        if required:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Experience years is required for experienced candidates.",
            )
        return None

    raw = value.strip()
    if not raw.isdigit():
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="Experience years must be a whole number.",
        )

    years = int(raw)
    if years < 1:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="Experience years must be at least 1.",
        )
    return years


def _fallback_dob_from_standalone(text: str) -> str | None:
    """
    Last-resort DOB extraction: scan all DD/MM/YYYY patterns and return the one
    that falls in a plausible birth-year range (1940–2010).
    """
    for match in _STANDALONE_DATE_PATTERN.finditer(text):
        day = int(_normalize_ocr_date_digits(match.group(1)))
        month = int(_normalize_ocr_date_digits(match.group(2)))
        year = int(_normalize_ocr_date_digits(match.group(3)))
        if 1940 <= year <= 2010 and 1 <= month <= 12 and 1 <= day <= 31:
            result = _parse_date_parts(day, month, year)
            if result:
                return result
    return None


def _run_tesseract_direct(image_bytes: bytes, mime_type: str = "image/png") -> str:
    """
    Run Tesseract directly via subprocess as a fallback when PyMuPDF OCR fails.
    Writes image to a temp file, runs tesseract, reads stdout text.
    """
    import tempfile
    tess_cmd = resolve_tesseract_command()
    if not tess_cmd:
        return ""
    tmp_in_path = ""
    tmp_out_path = ""
    try:
        suffix = ".png" if "png" in mime_type else ".jpg"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp_in:
            tmp_in.write(image_bytes)
            tmp_in_path = tmp_in.name
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp_out:
            tmp_out_path = tmp_out.name.replace(".txt", "")

        langs = get_ocr_languages()
        whitelist = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz/:-. "
        subprocess.run(
            [tess_cmd, tmp_in_path, tmp_out_path,
             "-l", langs, "--psm", "6",
             "--oem", "3",
             "-c", f"tessedit_char_whitelist={whitelist}"],
            capture_output=True,
            check=False,
            timeout=_TESSERACT_DIRECT_TIMEOUT_SECONDS,
            env={
                **os.environ,
                "OMP_THREAD_LIMIT": str(_TESSERACT_THREAD_LIMIT),
                "OMP_NUM_THREADS": str(_TESSERACT_THREAD_LIMIT),
            },
        )
        out_file = Path(tmp_out_path + ".txt")
        text = out_file.read_text(errors="ignore") if out_file.exists() else ""
        return text.strip()
    except Exception:
        return ""
    finally:
        for p in [tmp_in_path, tmp_out_path + ".txt"]:
            try:
                if p:
                    os.unlink(p)
            except Exception:
                pass


def _normalize_aadhaar_number(value: str | None) -> str | None:
    if not value:
        return None
    digits = re.sub(r"\D", "", value)
    if len(digits) == 12:
        return digits
    return None


_NAME_LINE_RE = re.compile(r"[A-Z][A-Za-z \-\.\']{2,60}")
_NOISE_WORDS = frozenset(
    "government india republic aadhaar unique identification authority enrolment"
    " address dob male female date birth year age".split()
)


def parse_name_from_text(text: str) -> str | None:
    candidates: list[str] = []
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not _NAME_LINE_RE.fullmatch(line):
            continue
        words = line.split()
        if len(words) < 2 or len(words) > 6:
            continue
        lowered = {w.lower() for w in words}
        if lowered & _NOISE_WORDS:
            continue
        if any(ch.isdigit() for ch in line):
            continue
        candidates.append(line)
    return candidates[0] if candidates else None


def _normalize_name_for_match(name: str) -> str:
    return re.sub(r"\s+", " ", name.strip().casefold())


def validate_aadhaar_identity(
    *,
    entered_name: str,
    entered_aadhaar: str,
    entered_dob: str | None,
    ocr_result: dict,
) -> dict:
    ocr_status = ocr_result.get("ocrStatus", "needs_review")
    ocr_numbers: list[str] = []
    for value in [
        ocr_result.get("aadhaarNumber"),
        *(ocr_result.get("aadhaarNumberCandidates") or []),
    ]:
        digits = re.sub(r"\D", "", str(value or ""))
        if len(digits) == 12 and digits not in ocr_numbers:
            ocr_numbers.append(digits)
    ocr_number = ocr_numbers[0] if ocr_numbers else ""
    ocr_dob = ocr_result.get("dateOfBirth") or ""
    ocr_name = ocr_result.get("cardHolderName") or ""
    entered_digits = re.sub(r"\D", "", entered_aadhaar)

    if ocr_status == "needs_review":
        return {
            "validationStatus": "needs_review",
            "ocrName": ocr_name or None,
            "mismatchReason": "OCR could not read the Aadhaar document — flagged for manual review.",
        }

    mismatches: list[str] = []

    if ocr_numbers and entered_digits and entered_digits not in ocr_numbers:
        mismatches.append(
            f"Aadhaar number on document (ending {ocr_number[-4:]}) does not match the entered number."
        )

    if ocr_dob and entered_dob:
        ocr_dob_norm = re.sub(r"\D", "", ocr_dob)
        entered_dob_norm = re.sub(r"\D", "", entered_dob)
        if ocr_dob_norm and entered_dob_norm and ocr_dob_norm != entered_dob_norm:
            mismatches.append("Date of birth on the document does not match the entered date of birth.")

    if ocr_name and entered_name:
        if _normalize_name_for_match(ocr_name) != _normalize_name_for_match(entered_name):
            mismatches.append("Name on the Aadhaar document does not match the entered name.")

    if mismatches:
        return {
            "validationStatus": "failed",
            "ocrName": ocr_name or None,
            "mismatchReason": " | ".join(mismatches),
        }

    return {
        "validationStatus": "passed",
        "ocrName": ocr_name or None,
        "mismatchReason": None,
    }


def _downgrade_candidate_aadhaar_failure_for_review(aadhaar_validation: dict) -> dict:
    """Candidate registration must not be blocked by a brittle OCR mismatch."""
    if aadhaar_validation.get("validationStatus") != "failed":
        return aadhaar_validation
    reason = aadhaar_validation.get("mismatchReason") or "OCR result did not match the entered details."
    return {
        **aadhaar_validation,
        "validationStatus": "needs_review",
        "mismatchReason": f"{reason} Flagged for manual review; account creation allowed.",
    }


def _render_pdf_pages_for_vision(content: bytes, *, max_pages: int = 2) -> list[bytes]:
    try:
        import pymupdf
    except ImportError:
        return []

    try:
        document = pymupdf.open(stream=content, filetype="pdf")
    except Exception:
        return []

    rendered_pages: list[bytes] = []
    with document:
        page_count = min(document.page_count, max_pages)
        dpi = max(get_ocr_dpi(), 220)
        for page_number in range(page_count):
            try:
                page = document.load_page(page_number)
                pixmap = page.get_pixmap(dpi=dpi, alpha=False)
                rendered_pages.append(pixmap.tobytes("png"))
            except Exception:
                continue
    return rendered_pages


# ── Concurrency control for CPU-bound OCR ─────────────────────────────────────
# Aadhaar OCR is CPU-bound (ONNX inference). Without a bound, every concurrent
# upload during a registration burst spawns its own worker thread, oversubscribes
# the cores, and every job then misses the request timeout — which surfaced as
# "OCR works locally but not in production". A single shared, bounded pool caps
# how many OCR jobs run at once; excess requests queue briefly instead of
# thrashing the box. Sized to the available cores (each job is pinned to one
# ONNX thread below, so N jobs map to ~N cores).
# Per-worker bound. With multiple Gunicorn workers, total OCR parallelism is
# (workers * this), so keep it to about half the cores per worker — each OCR job
# is pinned to one ONNX thread, so this maps to ~half the cores per worker and
# leaves headroom for the box's other request work. Override via OCR_MAX_CONCURRENCY.
_OCR_MAX_CONCURRENCY = int(
    os.getenv("OCR_MAX_CONCURRENCY") or max(2, (os.cpu_count() or 4) // 2)
)
_ocr_executor = concurrent.futures.ThreadPoolExecutor(
    max_workers=_OCR_MAX_CONCURRENCY,
    thread_name_prefix="aadhaar-ocr",
)


@lru_cache(maxsize=1)
def _get_rapidocr_engine():
    """Primary local Aadhaar OCR engine: RapidOCR (ONNX, pure-library, no system
    binary and no external API).

    Each call is restricted to a single intra-op thread so that N concurrent OCR
    jobs (bounded by ``_ocr_executor``) map cleanly to ~N CPU cores instead of
    every job spawning ``cpu_count`` threads and oversubscribing the host.
    """
    try:
        from rapidocr_onnxruntime import RapidOCR
    except Exception:
        return None

    # Newer RapidOCR releases accept ONNX thread limits as kwargs; older ones
    # don't, so fall back to the default constructor if they're rejected.
    for kwargs in ({"intra_op_num_threads": 1, "inter_op_num_threads": 1}, {}):
        try:
            return RapidOCR(**kwargs)
        except Exception:
            continue
    return None


def warm_ocr_engine() -> bool:
    """Eagerly initialise the RapidOCR engine at process startup so the first
    real upload does not pay the one-off ONNX model-load cost. Under a cold
    start that cost can exceed the per-request OCR timeout and look to the user
    like OCR is broken. Called once per worker from the app startup hook."""
    return _get_rapidocr_engine() is not None


def _extract_text_with_rapidocr(
    content: bytes,
    *,
    image_upload: bool,
    max_pages: int = 2,
) -> list[str]:
    engine = _get_rapidocr_engine()
    if engine is None:
        return []

    payloads = [content] if image_upload else _render_pdf_pages_for_vision(content, max_pages=max_pages)
    extracted_texts: list[str] = []

    for payload in payloads:
        try:
            result, _ = engine(payload)
        except Exception:
            continue

        lines = [
            line[1].strip()
            for line in (result or [])
            if len(line) > 1 and isinstance(line[1], str) and line[1].strip()
        ]
        if lines:
            extracted_texts.append("\n".join(lines))

    return extracted_texts


@lru_cache(maxsize=1)
def _get_apple_vision_ocr_binary() -> str | None:
    if get_settings().app_env.lower() == "test":
        return None
    if shutil.which("swiftc") is None:
        return None

    try:
        import platform

        if platform.system() != "Darwin":
            return None
    except Exception:
        return None

    source = """
import Foundation
import Vision
import AppKit

let args = CommandLine.arguments
guard args.count > 1 else {
    fputs("missing path\\n", stderr)
    exit(1)
}

let path = args[1]
guard let image = NSImage(contentsOfFile: path) else {
    fputs("could not load image\\n", stderr)
    exit(2)
}

guard let tiff = image.tiffRepresentation,
      let bitmap = NSBitmapImageRep(data: tiff),
      let cgImage = bitmap.cgImage else {
    fputs("could not build cgimage\\n", stderr)
    exit(3)
}

let request = VNRecognizeTextRequest()
request.recognitionLevel = .accurate
request.usesLanguageCorrection = false
let handler = VNImageRequestHandler(cgImage: cgImage, options: [:])
try handler.perform([request])

let texts = (request.results ?? []).compactMap { observation in
    observation.topCandidates(1).first?.string
}
print(texts.joined(separator: "\\n"))
"""

    cache_root = Path(tempfile.gettempdir()) / "ethara_ocr"
    cache_root.mkdir(parents=True, exist_ok=True)
    source_path = cache_root / "apple_vision_ocr.swift"
    binary_path = cache_root / "apple_vision_ocr"

    try:
        if not source_path.exists() or source_path.read_text() != source:
            source_path.write_text(source)
        if not binary_path.exists() or binary_path.stat().st_mtime < source_path.stat().st_mtime:
            subprocess.run(
                ["swiftc", str(source_path), "-o", str(binary_path)],
                check=True,
                capture_output=True,
                text=True,
                timeout=120,
            )
        return str(binary_path)
    except Exception:
        return None


def _extract_text_with_apple_vision(content: bytes, *, image_upload: bool) -> list[str]:
    binary = _get_apple_vision_ocr_binary()
    if binary is None:
        return []

    payloads = [content] if image_upload else _render_pdf_pages_for_vision(content)
    extracted_texts: list[str] = []

    for payload in payloads:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as handle:
            temp_path = Path(handle.name)

        try:
            if image_upload:
                from PIL import Image

                image = Image.open(BytesIO(payload))
                image.save(temp_path, format="PNG")
            else:
                temp_path.write_bytes(payload)

            result = subprocess.run(
                [binary, str(temp_path)],
                check=True,
                capture_output=True,
                text=True,
                timeout=120,
            )
            text = result.stdout.strip()
            if text:
                extracted_texts.append(text)
        except Exception:
            continue
        finally:
            try:
                temp_path.unlink(missing_ok=True)
            except Exception:
                pass

    return extracted_texts


def _parse_aadhaar_text_passes(text_passes: list[str]) -> tuple[list[str], str | None, str | None]:
    combined_text = "\n\n".join(pass_text for pass_text in text_passes if pass_text.strip())
    normalized = _normalize_ocr_digits(combined_text)
    full_text = f"{combined_text}\n{normalized}" if normalized else combined_text

    numbers = _rank_aadhaar_numbers_from_text(full_text)
    date_of_birth = parse_date_of_birth(full_text)
    if not date_of_birth:
        date_of_birth = _fallback_dob_from_standalone(full_text)

    return numbers, date_of_birth, parse_name_from_text(combined_text)


def _extract_aadhaar_via_gemini_vision(
    file: UploadFile,
    *,
    content: bytes,
    content_type: str,
) -> dict | None:
    settings = get_settings()
    if not settings.gemini_ocr_fallback or not settings.gemini_api_key:
        return None

    vision_payloads: list[tuple[bytes, str]] = []
    if is_image_upload(file):
        mime_type = content_type if content_type in IMAGE_CONTENT_TYPES else "image/png"
        vision_payloads.append((content, mime_type))
    elif Path(file.filename or "").suffix.lower() == ".pdf" or content_type == "application/pdf":
        vision_payloads.extend((page_bytes, "image/png") for page_bytes in _render_pdf_pages_for_vision(content))

    if not vision_payloads:
        return None

    from app.services.integrations import LLMService

    service = LLMService()
    for image_bytes, mime_type in vision_payloads:
        try:
            result = service.extract_aadhaar_via_vision(image_bytes, mime_type)
        except Exception:
            continue

        aadhaar_number = _normalize_aadhaar_number(result.get("aadhaarNumber"))
        raw_dob = result.get("dateOfBirth")
        date_of_birth = None
        if raw_dob:
            date_of_birth = parse_date_of_birth(f"DOB: {raw_dob}") or _fallback_dob_from_standalone(raw_dob)

        if aadhaar_number or date_of_birth:
            return {
                "aadhaarNumber": aadhaar_number,
                "dateOfBirth": date_of_birth,
                "cardHolderName": result.get("cardHolderName"),
                "ocrStatus": "extracted",
                "message": "Aadhaar details extracted successfully via Gemini Vision.",
            }

    return None


def _finalize_aadhaar_result(
    numbers: list[str],
    date_of_birth: str | None,
    card_holder_name: str | None,
    extracted_message: str | None = None,
) -> dict:
    """Build the OCR response payload (status + user-facing message) from the
    fields parsed by any engine. Shared by the early RapidOCR path and the full
    fallback chain so the status logic lives in exactly one place."""
    has_number = bool(numbers)
    has_dob = bool(date_of_birth)
    has_name = bool(card_holder_name)

    if has_number:
        ocr_status = "extracted"
        message = extracted_message or "Aadhaar number extracted successfully."
    elif has_dob or has_name:
        ocr_status = "partial"
        parts = []
        if has_dob:
            parts.append("date of birth")
        if has_name:
            parts.append("name")
        message = (
            f"Partial extraction: {' and '.join(parts)} detected. "
            "Please enter your Aadhaar number manually — it will be verified on submission."
        )
    else:
        ocr_status = "needs_review"
        message = (
            "Could not extract details from this Aadhaar document. "
            "Please ensure the image is clear and well-lit, then try again. "
            "You can also enter your details manually."
        )

    return {
        "aadhaarNumber": numbers[0] if numbers else None,
        "aadhaarNumberCandidates": numbers[:5],
        "dateOfBirth": date_of_birth,
        "cardHolderName": card_holder_name,
        "ocrStatus": ocr_status,
        "message": message,
    }


def _vertex_ocr_extract(content: bytes, content_type: str, expected_category: str) -> dict | None:
    """AI-primary registration OCR. When Vertex AI is enabled, verify the upload is
    the expected document type and extract its fields, returning the raw verdict
    (ok=True). Returns None when Vertex is disabled or the call failed, so the
    caller falls back to the local OCR library."""
    if not vertex_ai.is_enabled():
        return None
    verdict = vertex_ai.verify_and_extract(content, content_type or None, expected_category)
    return verdict if verdict.get("ok") else None


def _vertex_ocr_message(verdict: dict, label: str, extracted: bool) -> str:
    if verdict.get("matches_expected_category") is False:
        detected = verdict.get("detected_document_type") or "an unrecognised document"
        return (
            f"This file does not look like the expected {label} (detected: {detected}). "
            "Please double-check the document you uploaded."
        )
    if extracted:
        return f"{label} details extracted successfully."
    return verdict.get("validation_notes") or (
        f"Could not read {label} details. Please upload a clearer image or enter them manually."
    )


def extract_aadhaar_fields(file: UploadFile) -> dict:
    content = read_upload_bytes(file)
    suffix = Path(file.filename or "").suffix.lower()
    content_type = (file.content_type or "").split(";")[0].lower()
    image_upload = is_image_upload(file)

    verdict = _vertex_ocr_extract(content, content_type, "aadhaar")
    if verdict is not None:
        fields = verdict.get("extracted_fields") or {}
        aadhaar_raw = fields.get("aadhaar_number")
        if aadhaar_raw:
            digits_only = re.sub(r"\D", "", str(aadhaar_raw))
            aadhaar_raw = digits_only if len(digits_only) == 12 else str(aadhaar_raw)
        matched = verdict.get("matches_expected_category")
        date_of_birth = normalize_aadhaar_date_of_birth(fields.get("date_of_birth"))
        has_fields = bool(aadhaar_raw or date_of_birth)
        return {
            "aadhaarNumber": aadhaar_raw,
            "aadhaarNumberCandidates": [aadhaar_raw]
            if _normalize_aadhaar_number(aadhaar_raw)
            else [],
            "dateOfBirth": date_of_birth,
            "cardHolderName": fields.get("name"),
            "ocrStatus": "extracted" if (has_fields and matched is not False) else "needs_review",
            "message": _vertex_ocr_message(verdict, "Aadhaar card", has_fields),
            "detectedDocumentType": verdict.get("detected_document_type"),
            "matchesExpectedCategory": matched,
            "verification": verdict,
        }

    all_text_passes: list[str] = []

    # Try the bounded pure-library OCR path before Tesseract. Tesseract remains
    # available as a fallback, but it is intentionally no longer the first thing
    # an Aadhaar image upload does on small hosts.
    if image_upload:
        for payload in _aadhaar_image_ocr_payloads(content):
            try:
                all_text_passes.extend(_extract_text_with_rapidocr(payload, image_upload=True))
            except Exception:
                continue
            numbers, date_of_birth, card_holder_name = _parse_aadhaar_text_passes(all_text_passes)
            if numbers:
                return _finalize_aadhaar_result(numbers, date_of_birth, card_holder_name)

    from app.services import ocr as ocr_service

    if ocr_service.is_available() and resolve_tesseract_command():
        try:
            if image_upload:
                result = ocr_service.extract_aadhaar_from_image_bytes(content)
            else:
                pdf_text = extract_pdf_text(content)
                aadhaar_in_text = extract_aadhaar_numbers_from_text(pdf_text)
                dob_in_text = parse_date_of_birth(pdf_text) if pdf_text else None

                if aadhaar_in_text or dob_in_text:
                    numbers = _rank_aadhaar_numbers_from_text(pdf_text)
                    return {
                        "aadhaarNumber": numbers[0] if numbers else None,
                        "aadhaarNumberCandidates": numbers[:5],
                        "dateOfBirth": dob_in_text,
                        "cardHolderName": parse_name_from_text(pdf_text),
                        "ocrStatus": "extracted",
                        "message": "Aadhaar details extracted successfully.",
                    }

                ocr_result = ocr_service.ocr_pdf_bytes(content)
                result = ocr_service._build_aadhaar_result(ocr_result)

            aadhaar_raw = result.get("aadhaarNumber")
            date_of_birth = normalize_aadhaar_date_of_birth(result.get("dateOfBirth"))
            ocr_status = result.get("ocrStatus", "needs_review")

            if aadhaar_raw:
                digits_only = re.sub(r"\D", "", aadhaar_raw)
                if len(digits_only) == 12:
                    aadhaar_raw = digits_only

            if aadhaar_raw or date_of_birth:
                return {
                    "aadhaarNumber": aadhaar_raw,
                    "aadhaarNumberCandidates": [aadhaar_raw] if aadhaar_raw else [],
                    "dateOfBirth": date_of_birth,
                    "cardHolderName": result.get("cardHolderName") or result.get("name"),
                    "ocrStatus": ocr_status,
                    "message": result.get("message", "Aadhaar details extracted successfully."),
                }
        except Exception:
            pass

    if suffix == ".pdf" or content_type == "application/pdf":
        pdf_text = extract_pdf_text(content)
        if pdf_text.strip():
            all_text_passes.append(pdf_text)

    filetype = resolve_pymupdf_filetype(file)
    if filetype:
        proc_content = _preprocess_image_for_ocr(content) if image_upload else content
        pymupdf_text = extract_text_with_pymupdf_ocr(proc_content, filetype=filetype, full_ocr=True)
        if pymupdf_text.strip():
            all_text_passes.append(pymupdf_text)
        if image_upload and proc_content != content:
            orig_text = extract_text_with_pymupdf_ocr(content, filetype=filetype, full_ocr=True)
            if orig_text.strip() and orig_text != pymupdf_text:
                all_text_passes.append(orig_text)

    if image_upload and resolve_tesseract_command():
        for enhance_factor in (2.5, 1.8, 1.0):
            try:
                from PIL import Image, ImageEnhance, ImageFilter
                img = Image.open(BytesIO(content)).convert("L")
                img = img.filter(ImageFilter.SHARPEN)
                img = ImageEnhance.Contrast(img).enhance(enhance_factor)
                w, h = img.size
                if max(w, h) < 2000:
                    scale = 2000 / max(w, h)
                    img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
                buf = BytesIO()
                img.save(buf, format="PNG")
                tess_text = _run_tesseract_direct(buf.getvalue(), "image/png")
                if tess_text.strip():
                    all_text_passes.append(tess_text)
                    break
            except Exception:
                pass

    numbers, date_of_birth, card_holder_name = _parse_aadhaar_text_passes(all_text_passes)

    if not numbers and not date_of_birth:
        # Image uploads already ran RapidOCR as the primary engine above; only
        # PDFs reach it here. Apple Vision is a macOS-only dev fallback and is a
        # no-op on the Linux hosts.
        if not image_upload:
            all_text_passes.extend(_extract_text_with_rapidocr(content, image_upload=False))
        all_text_passes.extend(_extract_text_with_apple_vision(content, image_upload=image_upload))
        numbers, date_of_birth, card_holder_name = _parse_aadhaar_text_passes(all_text_passes)

    extracted_message: str | None = None

    # Optional cloud fallback. Self-gated by GEMINI_OCR_FALLBACK; disabled in
    # production so OCR stays a pure-library path with no external API calls.
    if not (numbers or date_of_birth):
        gemini_result = _extract_aadhaar_via_gemini_vision(
            file,
            content=content,
            content_type=content_type,
        )
        if gemini_result:
            aadhaar_number = gemini_result.get("aadhaarNumber")
            if aadhaar_number:
                numbers = [aadhaar_number]
            date_of_birth = gemini_result.get("dateOfBirth") or date_of_birth
            card_holder_name = (
                gemini_result.get("cardHolderName")
                or gemini_result.get("name")
                or card_holder_name
            )
            extracted_message = gemini_result.get("message")

    return _finalize_aadhaar_result(
        numbers, date_of_birth, card_holder_name, extracted_message
    )


def extract_aadhaar_numbers(file: UploadFile) -> set[str]:
    """Extract all valid Aadhaar numbers from a file (used for validation)."""
    result = extract_aadhaar_fields(file)
    num = result.get("aadhaarNumber")
    if num:
        return {re.sub(r"\D", "", num)}
    return set()


def _normalize_pan_number(value: str | None) -> str | None:
    normalized = re.sub(r"[^A-Z0-9]", "", (value or "").upper())
    if PAN_COMPACT_PATTERN.fullmatch(normalized):
        return normalized
    return None


def _normalize_pan_ocr_window(value: str) -> str | None:
    compact = re.sub(r"[^A-Z0-9]", "", (value or "").upper())
    if len(compact) != 10:
        return None

    candidate = (
        compact[:5].translate(PAN_LETTER_OCR_MAP)
        + compact[5:9].translate(PAN_DIGIT_OCR_MAP)
        + compact[9].translate(PAN_LETTER_OCR_MAP)
    )
    return _normalize_pan_number(candidate)


def _has_pan_context(text: str) -> bool:
    normalized = re.sub(r"[^A-Z]", "", (text or "").upper())
    return any(term in normalized for term in PAN_CONTEXT_TERMS)


def _pan_number_candidates_from_text(text: str) -> dict[str, int]:
    upper = (text or "").upper()
    has_context = _has_pan_context(upper)
    candidates: dict[str, int] = {}

    def add_candidate(value: str | None, score: int, *, fuzzy: bool = False) -> None:
        pan_number = _normalize_pan_number(value)
        if not pan_number:
            return
        if fuzzy and pan_number[:5] in PAN_FUZZY_FALSE_PREFIXES:
            return
        candidates[pan_number] = max(candidates.get(pan_number, 0), score)

    for match in PAN_PATTERN.finditer(upper):
        add_candidate("".join(match.groups()), 100)

    compact = re.sub(r"[^A-Z0-9]", "", upper)
    if has_context:
        for match in PAN_COMPACT_PATTERN.finditer(compact):
            add_candidate(match.group(0), 80)

        for match in PAN_ALNUM_WINDOW_PATTERN.finditer(compact):
            add_candidate(_normalize_pan_ocr_window(match.group(1)), 40, fuzzy=True)

    return candidates


def extract_pan_numbers_from_text(text: str) -> set[str]:
    return set(_pan_number_candidates_from_text(text))


def _parse_pan_text_passes(text_passes: list[str]) -> list[str]:
    combined_text = "\n\n".join(pass_text for pass_text in text_passes if pass_text.strip())
    candidates = _pan_number_candidates_from_text(combined_text)
    return [
        pan_number
        for pan_number, _score in sorted(
            candidates.items(),
            key=lambda item: (-item[1], item[0]),
        )
    ]


def _finalize_pan_result(numbers: list[str]) -> dict:
    if numbers:
        return {
            "panNumber": numbers[0],
            "ocrStatus": "extracted",
            "message": "PAN number extracted successfully.",
        }

    return {
        "panNumber": None,
        "ocrStatus": "needs_review",
        "message": (
            "Could not extract the PAN number from this document. "
            "Please ensure the image is clear and well-lit, then try again, "
            "or enter the PAN number manually."
        ),
    }


def extract_pan_fields(file: UploadFile) -> dict:
    content = read_upload_bytes(file)
    suffix = Path(file.filename or "").suffix.lower()
    content_type = (file.content_type or "").split(";")[0].lower()
    image_upload = is_image_upload(file)

    verdict = _vertex_ocr_extract(content, content_type, "pan")
    if verdict is not None:
        fields = verdict.get("extracted_fields") or {}
        pan_raw = fields.get("pan_number")
        pan_number = re.sub(r"\s+", "", str(pan_raw)).upper() if pan_raw else None
        matched = verdict.get("matches_expected_category")
        has_fields = bool(pan_number)
        return {
            "panNumber": pan_number,
            "ocrStatus": "extracted" if (has_fields and matched is not False) else "needs_review",
            "message": _vertex_ocr_message(verdict, "PAN card", has_fields),
            "detectedDocumentType": verdict.get("detected_document_type"),
            "matchesExpectedCategory": matched,
            "verification": verdict,
        }

    all_text_passes: list[str] = []

    # Phone camera / WhatsApp images are common for PAN uploads. Run the pure
    # OCR image engine first, including a contrast-enhanced variant, before the
    # heavier PDF/text fallbacks.
    if image_upload:
        for payload in _pan_image_ocr_payloads(content):
            try:
                all_text_passes.extend(_extract_text_with_rapidocr(payload, image_upload=True))
            except Exception:
                continue
            numbers = _parse_pan_text_passes(all_text_passes)
            if numbers:
                return _finalize_pan_result(numbers)

    if suffix in TEXT_SUFFIXES or content_type in TEXT_CONTENT_TYPES:
        try:
            text = content.decode("utf-8", errors="ignore")
        except Exception:
            text = ""
        if text.strip():
            all_text_passes.append(text)

    if suffix == ".pdf" or content_type == "application/pdf":
        pdf_text = extract_pdf_text(content)
        if pdf_text.strip():
            all_text_passes.append(pdf_text)

    numbers = _parse_pan_text_passes(all_text_passes)
    if numbers:
        return _finalize_pan_result(numbers)

    filetype = resolve_pymupdf_filetype(file)
    if filetype:
        try:
            proc_content = _preprocess_image_for_ocr(content) if image_upload else content
            pymupdf_text = extract_text_with_pymupdf_ocr(
                proc_content,
                filetype=filetype,
                full_ocr=True,
            )
            if pymupdf_text.strip():
                all_text_passes.append(pymupdf_text)
            if image_upload and proc_content != content:
                original_text = extract_text_with_pymupdf_ocr(
                    content,
                    filetype=filetype,
                    full_ocr=True,
                )
                if original_text.strip() and original_text != pymupdf_text:
                    all_text_passes.append(original_text)
        except Exception:
            pass

    numbers = _parse_pan_text_passes(all_text_passes)
    if numbers:
        return _finalize_pan_result(numbers)

    if image_upload and resolve_tesseract_command():
        for enhance_factor in (2.5, 1.8, 1.0):
            try:
                from PIL import Image, ImageEnhance, ImageFilter
                img = Image.open(BytesIO(content)).convert("L")
                img = img.filter(ImageFilter.SHARPEN)
                img = ImageEnhance.Contrast(img).enhance(enhance_factor)
                w, h = img.size
                if max(w, h) < 2000:
                    scale = 2000 / max(w, h)
                    img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
                buf = BytesIO()
                img.save(buf, format="PNG")
                tess_text = _run_tesseract_direct(buf.getvalue(), "image/png")
                if tess_text.strip():
                    all_text_passes.append(tess_text)
                    break
            except Exception:
                pass

    numbers = _parse_pan_text_passes(all_text_passes)
    if numbers:
        return _finalize_pan_result(numbers)

    try:
        all_text_passes.extend(_extract_text_with_rapidocr(content, image_upload=image_upload))
    except Exception:
        pass
    try:
        all_text_passes.extend(_extract_text_with_apple_vision(content, image_upload=image_upload))
    except Exception:
        pass

    return _finalize_pan_result(_parse_pan_text_passes(all_text_passes))


def _is_address_stop_line(line: str) -> bool:
    lower = line.lower()
    if any(word in lower for word in ADDRESS_STOP_WORDS):
        return True
    if re.search(r"\b\d{4}\s?\d{4}\s?\d{4}\b", line):
        return True
    return False


def _clean_address_line(line: str) -> str:
    cleaned = re.sub(r"\s+", " ", line or "").strip(" .,:;-")
    cleaned = re.sub(r"^(?:address|addr)\s*[:\-]\s*", "", cleaned, flags=re.IGNORECASE)
    return cleaned.strip(" .,:;-")


def _parse_address_text(text: str) -> dict:
    raw_lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    lines = [_clean_address_line(line) for line in raw_lines]
    lines = [line for line in lines if line]
    collected: list[str] = []
    collecting = False

    for raw_line, line in zip(raw_lines, lines, strict=False):
        inline_match = ADDRESS_INLINE_RE.search(raw_line)
        if inline_match:
            inline_value = _clean_address_line(inline_match.group(1))
            if inline_value and not _is_address_stop_line(inline_value):
                collected.append(inline_value)
            collecting = True
            continue

        if "address" in raw_line.lower():
            collecting = True
            continue

        if not collecting:
            continue
        if _is_address_stop_line(line):
            if collected:
                break
            continue
        collected.append(line)
        if ADDRESS_POSTAL_RE.search(line) and len(collected) >= 2:
            break
        if len(collected) >= 6:
            break

    if not collected:
        for index, line in enumerate(lines):
            if ADDRESS_POSTAL_RE.search(line):
                start = max(0, index - 4)
                window = lines[start : index + 1]
                collected = [item for item in window if not _is_address_stop_line(item)]
                break

    deduped: list[str] = []
    seen: set[str] = set()
    for line in collected:
        key = re.sub(r"[^a-z0-9]", "", line.lower())
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append(line)

    address = ", ".join(deduped) if deduped else None
    postal_match = ADDRESS_POSTAL_RE.search("\n".join(deduped) or text or "")
    return {
        "address": address,
        "addressLines": deduped,
        "postalCode": postal_match.group(0) if postal_match else None,
    }


def _finalize_address_result(parsed: dict) -> dict:
    address = parsed.get("address")
    postal_code = parsed.get("postalCode")
    if address:
        return {
            "address": address,
            "addressLines": parsed.get("addressLines") or [],
            "postalCode": postal_code,
            "ocrStatus": "extracted",
            "message": "Address extracted from document.",
        }
    return {
        "address": None,
        "addressLines": [],
        "postalCode": postal_code,
        "ocrStatus": "needs_review",
        "message": (
            "Could not extract address from this document. "
            "Please upload a clearer address proof or enter the address manually."
        ),
    }


def extract_address_fields(file: UploadFile) -> dict:
    content = read_upload_bytes(file)
    suffix = Path(file.filename or "").suffix.lower()
    content_type = (file.content_type or "").split(";")[0].lower()
    image_upload = is_image_upload(file)

    verdict = _vertex_ocr_extract(content, content_type, "address_proof")
    if verdict is not None:
        fields = verdict.get("extracted_fields") or {}
        address = fields.get("address")
        matched = verdict.get("matches_expected_category")
        has_fields = bool(address)
        return {
            "address": address,
            "addressLines": [line for line in str(address).splitlines() if line.strip()] if address else [],
            "postalCode": fields.get("postal_code"),
            "ocrStatus": "extracted" if (has_fields and matched is not False) else "needs_review",
            "message": _vertex_ocr_message(verdict, "address proof", has_fields),
            "detectedDocumentType": verdict.get("detected_document_type"),
            "matchesExpectedCategory": matched,
            "verification": verdict,
        }

    all_text_passes: list[str] = []

    if image_upload:
        for payload in _document_image_ocr_payloads(content):
            try:
                all_text_passes.extend(_extract_text_with_rapidocr(payload, image_upload=True))
            except Exception:
                continue
            parsed = _parse_address_text(_join_ocr_text_passes(all_text_passes))
            if parsed.get("address"):
                return _finalize_address_result(parsed)

    if suffix in TEXT_SUFFIXES or content_type in TEXT_CONTENT_TYPES:
        try:
            text = content.decode("utf-8", errors="ignore")
        except Exception:
            text = ""
        if text.strip():
            all_text_passes.append(text)

    if suffix == ".pdf" or content_type == "application/pdf":
        pdf_text = extract_pdf_text(content)
        if pdf_text.strip():
            all_text_passes.append(pdf_text)

    parsed = _parse_address_text(_join_ocr_text_passes(all_text_passes))
    if parsed.get("address"):
        return _finalize_address_result(parsed)

    filetype = resolve_pymupdf_filetype(file)
    if filetype:
        try:
            proc_content = _preprocess_image_for_ocr(content) if image_upload else content
            pymupdf_text = extract_text_with_pymupdf_ocr(
                proc_content,
                filetype=filetype,
                full_ocr=True,
            )
            if pymupdf_text.strip():
                all_text_passes.append(pymupdf_text)
        except Exception:
            pass

    if not image_upload:
        try:
            all_text_passes.extend(_extract_text_with_rapidocr(content, image_upload=False))
        except Exception:
            pass

    try:
        all_text_passes.extend(_extract_text_with_apple_vision(content, image_upload=image_upload))
    except Exception:
        pass

    return _finalize_address_result(_parse_address_text(_join_ocr_text_passes(all_text_passes)))


def validate_aadhaar_upload(file: UploadFile, entered_aadhaar_number: str) -> None:
    """
    Soft validation: if OCR found numbers and none matches the entered one, warn.
    We don't hard-block since OCR can fail on scanned/low-quality cards.
    """
    uploaded_numbers = extract_aadhaar_numbers(file)
    normalized_entered = re.sub(r"\D", "", entered_aadhaar_number)
    if uploaded_numbers and normalized_entered not in uploaded_numbers:
        raise HTTPException(
            status_code=400,
            detail="Aadhaar card OCR did not match the entered Aadhaar number. "
                   "Please ensure you uploaded the correct Aadhaar card.",
        )


@router.post("/aadhaar/ocr")
@limiter.limit("60/minute")
def extract_aadhaar_card(
    request: Request,
    aadhaar_card: Annotated[UploadFile, File(alias="aadhaarCard")],
) -> dict:
    # Run OCR on the shared, bounded pool with a 50-second timeout. The bound
    # caps concurrent CPU-bound OCR jobs (see _ocr_executor) so a registration
    # burst queues briefly instead of oversubscribing the cores and timing out.
    try:
        future = _ocr_executor.submit(extract_aadhaar_fields, aadhaar_card)
        return future.result(timeout=50)
    except concurrent.futures.TimeoutError:
        return {
            "aadhaarNumber": None,
            "dateOfBirth": None,
            "cardHolderName": None,
            "ocrStatus": "needs_review",
            "message": "OCR timed out. Please enter your Aadhaar details manually below.",
        }
    except Exception:
        return {
            "aadhaarNumber": None,
            "dateOfBirth": None,
            "cardHolderName": None,
            "ocrStatus": "needs_review",
            "message": "Could not process document. Please enter your Aadhaar details manually.",
        }


@router.post("/pan/ocr")
@limiter.limit("60/minute")
def extract_candidate_pan_card(
    request: Request,
    pan_card: Annotated[UploadFile, File(alias="panCard")],
) -> dict:
    try:
        future = _ocr_executor.submit(extract_pan_fields, pan_card)
        return future.result(timeout=50)
    except concurrent.futures.TimeoutError:
        return {
            "panNumber": None,
            "ocrStatus": "needs_review",
            "message": "PAN OCR timed out. Please enter the PAN number manually.",
        }
    except Exception:
        return {
            "panNumber": None,
            "ocrStatus": "needs_review",
            "message": "PAN details could not be extracted. Please upload a clearer PAN card image or enter the PAN manually.",
        }


@router.post("/address/ocr")
@limiter.limit("60/minute")
def extract_candidate_address_proof(
    request: Request,
    address_proof: Annotated[UploadFile, File(alias="addressProof")],
) -> dict:
    try:
        future = _ocr_executor.submit(extract_address_fields, address_proof)
        return future.result(timeout=50)
    except concurrent.futures.TimeoutError:
        return {
            "address": None,
            "addressLines": [],
            "postalCode": None,
            "ocrStatus": "needs_review",
            "message": "Address OCR timed out. Please enter the address manually.",
        }
    except Exception:
        return {
            "address": None,
            "addressLines": [],
            "postalCode": None,
            "ocrStatus": "needs_review",
            "message": "Could not extract address from this document. Please upload a clearer address proof or enter the address manually.",
        }


# ── Cancelled cheque OCR ──────────────────────────────────────────────────────

_CHEQUE_IFSC_RE = re.compile(r'\b[A-Z]{4}0[A-Z0-9]{6}\b')
_CHEQUE_ACC_LABELED = re.compile(
    r'(?:a/?c|acc(?:ount)?)[\s\.#:]*(?:no\.?|number|#)?[\s:]*(\d{9,18})',
    re.IGNORECASE,
)
_CHEQUE_ACC_MICR = re.compile(r'[|⑆](\d{6,18})[|⑆]')
_CHEQUE_ACC_FALLBACK = re.compile(r'\b(\d{9,18})\b')
_CHEQUE_HOLDER_RE = re.compile(
    r'\b(?:account\s*holder\s*name|account\s*holder|account\s*name|a/?c\s*name|name)'
    r'[\s:\.\-]+([A-Za-z][A-Za-z \.\-\']{2,60})',
    re.IGNORECASE,
)
_CHEQUE_BANK_RE = re.compile(r'\b([A-Z][A-Za-z\s]{2,40}(?:Bank|Banking))\b', re.IGNORECASE)

_KNOWN_BANKS = [
    "HDFC Bank", "ICICI Bank", "State Bank of India", "Axis Bank",
    "Kotak Mahindra Bank", "Bank of Baroda", "Punjab National Bank",
    "Canara Bank", "Union Bank of India", "IndusInd Bank", "Yes Bank",
    "Bank of India", "Central Bank of India", "UCO Bank", "Indian Bank",
    "Indian Overseas Bank", "IDBI Bank", "Federal Bank", "South Indian Bank",
    "RBL Bank", "Bandhan Bank", "AU Small Finance Bank",
    "Equitas Small Finance Bank", "Jana Small Finance Bank",
]


def _parse_cheque_text(text: str) -> dict:
    upper = text.upper()

    ifsc_matches = _CHEQUE_IFSC_RE.findall(upper)
    ifsc_code = ifsc_matches[0] if ifsc_matches else None

    account_number: str | None = None
    m = _CHEQUE_ACC_LABELED.search(text)
    if m:
        account_number = m.group(1)
    if not account_number:
        m = _CHEQUE_ACC_MICR.search(text)
        if m:
            account_number = m.group(1)
    if not account_number:
        candidates = [c for c in _CHEQUE_ACC_FALLBACK.findall(text) if 9 <= len(c) <= 18]
        if candidates:
            account_number = max(candidates, key=len)

    bank_name: str | None = None
    for name in _KNOWN_BANKS:
        if name.upper() in upper:
            bank_name = name
            break
    if not bank_name:
        bm = _CHEQUE_BANK_RE.search(text)
        if bm:
            bank_name = bm.group(1).strip()

    holder_name: str | None = None
    hm = _CHEQUE_HOLDER_RE.search(text)
    if hm:
        holder_name = re.sub(r"\s+", " ", hm.group(1)).strip(" .:-")

    return {
        "accountNumber": account_number,
        "ifscCode": ifsc_code,
        "accountHolderName": holder_name,
        "bankName": bank_name,
    }


def _extract_cheque_via_gemini_vision(
    file: UploadFile,
    *,
    content: bytes,
    content_type: str,
) -> dict | None:
    settings = get_settings()
    if not settings.gemini_ocr_fallback or not settings.gemini_api_key:
        return None

    vision_payloads: list[tuple[bytes, str]] = []
    if is_image_upload(file):
        mime = content_type if content_type in IMAGE_CONTENT_TYPES else "image/png"
        vision_payloads.append((content, mime))
    elif Path(file.filename or "").suffix.lower() == ".pdf" or content_type == "application/pdf":
        vision_payloads.extend(
            (page_bytes, "image/png") for page_bytes in _render_pdf_pages_for_vision(content)
        )

    if not vision_payloads:
        return None

    from app.services.integrations import LLMService

    service = LLMService()
    for image_bytes, mime_type in vision_payloads:
        try:
            result = service.extract_cheque_via_vision(image_bytes, mime_type)
        except Exception:
            continue

        account_number = re.sub(r"\D", "", str(result.get("accountNumber") or "")) or None
        raw_ifsc = str(result.get("ifscCode") or "").upper().replace(" ", "")
        ifsc_code = raw_ifsc if raw_ifsc and _CHEQUE_IFSC_RE.fullmatch(raw_ifsc) else None

        if account_number or ifsc_code:
            return {
                "accountNumber": account_number,
                "ifscCode": ifsc_code,
                "accountHolderName": result.get("accountHolderName"),
                "bankName": result.get("bankName"),
                "ocrStatus": "extracted",
                "message": "Bank details extracted successfully via Gemini Vision.",
            }
    return None


def extract_cheque_fields(file: UploadFile) -> dict:
    content = read_upload_bytes(file)
    suffix = Path(file.filename or "").suffix.lower()
    content_type = (file.content_type or "").split(";")[0].lower()
    image_upload = is_image_upload(file)

    verdict = _vertex_ocr_extract(content, content_type, "bank_proof")
    if verdict is not None:
        fields = verdict.get("extracted_fields") or {}
        account_number = fields.get("account_number")
        ifsc = fields.get("ifsc")
        matched = verdict.get("matches_expected_category")
        has_fields = bool(account_number or ifsc)
        return {
            "accountNumber": re.sub(r"\s+", "", str(account_number)) if account_number else None,
            "ifscCode": re.sub(r"\s+", "", str(ifsc)).upper() if ifsc else None,
            "accountHolderName": fields.get("account_holder_name"),
            "bankName": fields.get("bank_name"),
            "ocrStatus": "extracted" if (has_fields and matched is not False) else "needs_review",
            "message": _vertex_ocr_message(verdict, "bank proof", has_fields),
            "detectedDocumentType": verdict.get("detected_document_type"),
            "matchesExpectedCategory": matched,
            "verification": verdict,
        }

    all_text_passes: list[str] = []
    parsed: dict = {}

    if image_upload:
        for payload in _document_image_ocr_payloads(content):
            try:
                all_text_passes.extend(_extract_text_with_rapidocr(payload, image_upload=True))
            except Exception:
                continue
            parsed = _parse_cheque_text(_join_ocr_text_passes(all_text_passes))
            if parsed.get("accountNumber") and parsed.get("ifscCode"):
                break

    filetype = resolve_pymupdf_filetype(file)
    if filetype:
        proc = _preprocess_image_for_ocr(content) if image_upload else content
        t = extract_text_with_pymupdf_ocr(proc, filetype=filetype, full_ocr=True)
        if t.strip():
            all_text_passes.append(t)
        if image_upload and proc != content:
            t2 = extract_text_with_pymupdf_ocr(content, filetype=filetype, full_ocr=True)
            if t2.strip() and t2 != t:
                all_text_passes.append(t2)

    if suffix == ".pdf" or content_type == "application/pdf":
        pdf_text = extract_pdf_text(content)
        if pdf_text.strip():
            all_text_passes.append(pdf_text)

    if image_upload and resolve_tesseract_command():
        for enhance_factor in (2.5, 1.8, 1.0):
            try:
                from PIL import Image, ImageEnhance, ImageFilter
                img = Image.open(BytesIO(content)).convert("L")
                img = img.filter(ImageFilter.SHARPEN)
                img = ImageEnhance.Contrast(img).enhance(enhance_factor)
                w, h = img.size
                if max(w, h) < 2000:
                    scale = 2000 / max(w, h)
                    img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
                buf = BytesIO()
                img.save(buf, format="PNG")
                tess_text = _run_tesseract_direct(buf.getvalue(), "image/png")
                if tess_text.strip():
                    all_text_passes.append(tess_text)
                    break
            except Exception:
                pass

    combined = "\n\n".join(t for t in all_text_passes if t.strip())
    parsed = _parse_cheque_text(combined) if combined else parsed
    account_number = parsed.get("accountNumber")
    ifsc_code = parsed.get("ifscCode")
    holder_name = parsed.get("accountHolderName")
    bank_name = parsed.get("bankName")

    if not account_number and not ifsc_code:
        all_text_passes.extend(_extract_text_with_rapidocr(content, image_upload=image_upload))
        all_text_passes.extend(_extract_text_with_apple_vision(content, image_upload=image_upload))
        combined = "\n\n".join(t for t in all_text_passes if t.strip())
        if combined:
            p2 = _parse_cheque_text(combined)
            account_number = account_number or p2.get("accountNumber")
            ifsc_code = ifsc_code or p2.get("ifscCode")
            holder_name = holder_name or p2.get("accountHolderName")
            bank_name = bank_name or p2.get("bankName")

    if not account_number and not ifsc_code:
        gem = _extract_cheque_via_gemini_vision(file, content=content, content_type=content_type)
        if gem:
            account_number = account_number or gem.get("accountNumber")
            ifsc_code = ifsc_code or gem.get("ifscCode")
            holder_name = holder_name or gem.get("accountHolderName")
            bank_name = bank_name or gem.get("bankName")

    has_account = bool(account_number)
    has_ifsc = bool(ifsc_code)

    if has_account and has_ifsc:
        ocr_status = "extracted"
        message = "Bank details extracted from cancelled cheque."
    elif has_account or has_ifsc:
        found = []
        if has_account:
            found.append("account number")
        if has_ifsc:
            found.append("IFSC code")
        ocr_status = "partial"
        message = (
            f"Partial extraction: {' and '.join(found)} detected. "
            "Please verify and complete any remaining fields."
        )
    else:
        ocr_status = "needs_review"
        message = (
            "Could not extract bank details from this cheque. "
            "Please ensure the image is clear and well-lit, then try again, "
            "or enter your details manually."
        )

    return {
        "accountNumber": account_number,
        "ifscCode": ifsc_code,
        "accountHolderName": holder_name,
        "bankName": bank_name,
        "ocrStatus": ocr_status,
        "message": message,
    }


@router.post("/cheque/ocr")
@limiter.limit("60/minute")
def extract_cancelled_cheque(
    request: Request,
    cancelled_cheque: Annotated[UploadFile, File(alias="cancelledCheque")],
) -> dict:
    return extract_cheque_fields(cancelled_cheque)


@router.post("/aadhaar/check")
@limiter.limit("20/minute")
def check_aadhaar_duplicate(
    request: Request,
    payload: dict,
    db: Annotated[Session, Depends(get_db)],
) -> dict:
    """Check if an Aadhaar number is already registered in the system."""
    from app.core.security import fingerprint_identifier
    from app.db.models import Candidate
    from sqlalchemy import select as sa_select

    aadhaar_raw = re.sub(r"\D", "", payload.get("aadhaarNumber", ""))
    if len(aadhaar_raw) != 12:
        return {"exists": False, "message": "Invalid Aadhaar number format."}

    aadhaar_hash = fingerprint_identifier(aadhaar_raw)
    existing = db.scalar(
        sa_select(Candidate).where(
            Candidate.aadhaar_hash == aadhaar_hash,
            Candidate.current_status != "Removed",  # removed candidates can re-register
        )
    )
    if existing:
        return {
            "exists": True,
            "message": "This Aadhaar number is already registered in the system.",
        }
    return {"exists": False, "message": "Aadhaar number is available."}


@router.post("/resume/parse")
@limiter.limit("10/minute")
def parse_resume_file(
    request: Request,
    resume: Annotated[UploadFile, File()],
    position_title: str | None = Form(default=None, alias="positionTitle"),
) -> dict:
    """Extract and parse resume text using Gemini LLM."""
    from app.services.integrations import LLMService

    content = read_upload_bytes(resume)
    suffix = Path(resume.filename or "").suffix.lower()
    content_type = (resume.content_type or "").split(";")[0].lower()

    # Extract raw text from the resume
    resume_text = ""
    if suffix == ".pdf" or content_type == "application/pdf":
        resume_text = extract_pdf_text(content)
        if not resume_text.strip():
            filetype = resolve_pymupdf_filetype(resume) or "pdf"
            resume_text = extract_text_with_pymupdf_ocr(content, filetype=filetype, full_ocr=True)
    elif suffix in {".doc", ".docx"} or "word" in content_type:
        try:
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(content))
            resume_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            resume_text = content.decode("utf-8", errors="ignore")
    else:
        resume_text = content.decode("utf-8", errors="ignore")

    if not resume_text.strip():
        return {
            "resumeText": "",
            "summary": "Could not extract text from this resume.",
            "keyPoints": [],
            "skills": [],
            "totalExperienceYears": 0,
        }

    # Parse with Gemini
    svc = LLMService()
    parsed = svc.parse_resume(resume_text=resume_text, job_title=position_title)
    parsed["resumeText"] = _strip_nul(resume_text)[:20000]  # cap for storage; scrub NUL for Postgres
    return parsed


@router.post("/register")
@limiter.limit("5/minute")
def register_candidate(
    request: Request,
    background_tasks: BackgroundTasks,
    db: Annotated[Session, Depends(get_db)],
    aadhaar_card: Annotated[UploadFile, File(alias="aadhaarCard")],
    full_name: str = Form(alias="fullName"),
    gender: str = Form(),
    experience_type: str = Form(alias="experienceType"),
    experience_years: str | None = Form(default=None, alias="experienceYears"),
    personal_email: str = Form(alias="personalEmail"),
    phone: str = Form(),
    password: str = Form(min_length=8, max_length=128),
    aadhaar_number: str = Form(alias="aadhaarNumber"),
    date_of_birth: str | None = Form(default=None, alias="dateOfBirth"),
    position_id: str | None = Form(default=None, alias="positionId"),
    college_id: str | None = Form(default=None, alias="collegeId"),
    resume: Annotated[UploadFile | None, File()] = None,
) -> dict:
    # Enforce the same password-strength rules used by reset/change flows (#45).
    try:
        validate_password_strength(password)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    if not re.match(r'^[6-9]\d{9}$', phone.replace(' ', '')):
        raise HTTPException(status_code=422, detail="Phone must be a valid 10-digit Indian mobile number.")
    personal_email = _validate_registration_email(personal_email)

    # Enforce a content-type/extension allowlist and size cap on the public
    # registration uploads BEFORE any OCR/parse/storage read touches the bytes.
    validate_candidate_upload(
        aadhaar_card,
        label="Aadhaar card",
        allowed_content_types=ALLOWED_DOCUMENT_IMAGE_CONTENT_TYPES,
        allowed_extensions=ALLOWED_DOCUMENT_IMAGE_EXTENSIONS,
        required=True,
    )
    validate_candidate_upload(
        resume,
        label="Resume",
        allowed_content_types=ALLOWED_RESUME_CONTENT_TYPES,
        allowed_extensions=ALLOWED_RESUME_EXTENSIONS,
        required=True,
    )

    normalized_experience_type = experience_type.strip().lower()
    if normalized_experience_type not in {"fresher", "experienced"}:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="Experience type must be either fresher or experienced.",
        )
    resolved_experience_years = parse_experience_years(
        experience_years,
        required=normalized_experience_type == "experienced",
    )

    if resume is None or not resume.filename:
        raise HTTPException(status_code=400, detail="Resume upload is required.")

    aadhaar_ocr = extract_aadhaar_fields(aadhaar_card)

    aadhaar_validation = validate_aadhaar_identity(
        entered_name=full_name,
        entered_aadhaar=aadhaar_number,
        entered_dob=date_of_birth,
        ocr_result=aadhaar_ocr,
    )
    aadhaar_validation = _downgrade_candidate_aadhaar_failure_for_review(aadhaar_validation)

    if aadhaar_validation["validationStatus"] not in ("passed", "partial", "needs_review"):
        raise HTTPException(status_code=400, detail="Invalid Documents")

    entered_aadhaar_digits = re.sub(r"\D", "", aadhaar_number)
    ocr_candidate_digits = {
        re.sub(r"\D", "", str(value or ""))
        for value in [
            aadhaar_ocr.get("aadhaarNumber"),
            *(aadhaar_ocr.get("aadhaarNumberCandidates") or []),
        ]
    }
    if len(entered_aadhaar_digits) == 12 and entered_aadhaar_digits in ocr_candidate_digits:
        aadhaar_ocr["aadhaarNumber"] = entered_aadhaar_digits

    resume_url = save_candidate_upload(resume, "resumes")
    aadhaar_url = save_candidate_upload(aadhaar_card, "aadhaar")
    resolved_date_of_birth = parse_iso_date(date_of_birth)
    if resolved_date_of_birth is None:
        resolved_date_of_birth = parse_iso_date(aadhaar_ocr.get("dateOfBirth"))

    resume_content = read_upload_bytes(resume)
    resume_text_extracted = ""
    resume_suffix = Path(resume.filename or "").suffix.lower()
    resume_ct = (resume.content_type or "").split(";")[0].lower()

    resume_text_extracted = _extract_resume_text(resume_content, resume_suffix, resume_ct)

    resume_key_points: list[str] = []
    resume_summary: str | None = None
    if resume_text_extracted.strip():
        try:
            from app.services.integrations import LLMService
            from sqlalchemy import select as _sa_select
            from app.db.models import Position as _Position
            position_title_for_parse: str | None = None
            if position_id:
                _pos = db.scalar(_sa_select(_Position).where(_Position.id == position_id))
                if _pos:
                    position_title_for_parse = _pos.title
            parsed = LLMService().parse_resume(
                resume_text=resume_text_extracted,
                job_title=position_title_for_parse,
            )
            resume_key_points = parsed.get("keyPoints") or []
            resume_summary = parsed.get("summary") or None
        except Exception:
            pass

    candidate = candidate_service.create_candidate(
        db,
        payload={
            "full_name": full_name,
            "gender": gender,
            "experience_type": normalized_experience_type,
            "experience_years": resolved_experience_years,
            "personal_email": personal_email,
            "phone": phone,
            "portal_password": password,
            "aadhaar_number": aadhaar_number,
            "date_of_birth": resolved_date_of_birth,
            "position_id": position_id,
            "college_id": college_id,
            "resume_url": resume_url,
            "resume_text": resume_text_extracted[:20000] if resume_text_extracted else None,
            "resume_key_points": resume_key_points or None,
            "resume_summary": resume_summary,
            "aadhaar_extracted": aadhaar_ocr,
            "aadhaar_ocr_name": aadhaar_validation.get("ocrName"),
            "aadhaar_validation_status": aadhaar_validation.get("validationStatus"),
            "aadhaar_mismatch_reason": aadhaar_validation.get("mismatchReason"),
            "source_type": "direct_application",
            "current_stage": CandidateStage.RESUME_SCREENING_PENDING,
            "current_status": "Screening Pending",
        },
        actor=None,
        request=request,
    )
    record_candidate_upload_document(
        db,
        candidate=candidate,
        file=resume,
        type_="resume",
        file_url=resume_url,
    )
    record_candidate_upload_document(
        db,
        candidate=candidate,
        file=aadhaar_card,
        type_="aadhaar_card",
        file_url=aadhaar_url,
        ocr_result=aadhaar_ocr,
    )
    try:
        db.commit()
    except IntegrityError as exc:
        db.rollback()
        raise HTTPException(
            status_code=409,
            detail="A candidate with this email or Aadhaar number already exists.",
        ) from exc
    db.refresh(candidate)
    candidate.llm_status = "processing"
    screening_payload = dict(candidate.screening_payload or {})
    screening_payload["status"] = "queued"
    screening_payload["queuedAt"] = datetime.now(UTC).isoformat()
    candidate.screening_payload = screening_payload
    db.add(candidate)
    db.commit()
    db.refresh(candidate)

    # Send email verification OTP after successful registration
    otp_result = None
    if candidate.portal_user_id:
        from sqlalchemy import select as sa_select

        from app.db.models import User as UserModel
        portal_user = db.scalar(
            sa_select(UserModel).where(UserModel.id == candidate.portal_user_id)
        )
        if portal_user:
            try:
                otp_result = account_security.request_email_verification(db, user=portal_user)
                db.commit()
            except HTTPException as exc:
                db.rollback()
                logger.warning(
                    "Verification OTP dispatch failed for registered candidate %s (%s): %s",
                    candidate.id,
                    portal_user.email,
                    exc.detail,
                )
            except Exception as exc:  # noqa: BLE001 - candidate account is already created
                db.rollback()
                logger.warning(
                    "Verification OTP dispatch failed for registered candidate %s (%s): %s",
                    candidate.id,
                    portal_user.email,
                    exc,
                )

    background_tasks.add_task(_run_registration_screening_background, candidate.id)

    return {
        "candidateId": candidate.id,
        "email": candidate.personal_email,
        "message": (
            "Registration successful. Please verify your email to continue."
            if otp_result
            else (
                "Registration successful. Use resend OTP on the verification page "
                "if the code does not arrive."
            )
        ),
        "developmentCode": otp_result.development_code if otp_result else None,
        "expiresAt": (
            otp_result.expires_at.isoformat()
            if otp_result and otp_result.expires_at
            else None
        ),
    }


# ─────────────────────────────── campus drive ────────────────────────────────


def _campus_enabled(db: Session) -> bool:
    setting = db.scalar(select(AdminSetting).where(AdminSetting.key == "campus_drive"))
    return bool(setting and isinstance(setting.value, dict) and setting.value.get("enabled"))


@router.get("/campus/config")
def campus_config(db: Annotated[Session, Depends(get_db)]) -> dict:
    """Public: tells the registration page whether the campus-drive path is open."""
    return {"enabled": _campus_enabled(db)}


@router.put("/campus/config")
def set_campus_config(
    payload: dict,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.SETTINGS_WRITE))],
) -> dict:
    """Admin: turn the campus-drive registration path on/off."""
    value = {"enabled": bool(payload.get("enabled"))}
    setting = db.scalar(select(AdminSetting).where(AdminSetting.key == "campus_drive"))
    if setting is None:
        setting = AdminSetting(
            namespace="recruitment", key="campus_drive", value=value,
            description="Campus drive registration", updated_by=current_user.id,
        )
        db.add(setting)
    else:
        setting.value = value
        setting.updated_by = current_user.id
        db.add(setting)
    db.commit()
    return value


@router.post("/campus/register")
@limiter.limit("5/minute")
def register_campus(
    request: Request,
    db: Annotated[Session, Depends(get_db)],
    full_name: str = Form(alias="fullName"),
    personal_email: str = Form(alias="personalEmail"),
    phone: str = Form(),
    password: str = Form(min_length=8, max_length=128),
    college_id: str | None = Form(default=None, alias="collegeId"),
    position_id: str | None = Form(default=None, alias="positionId"),
) -> dict:
    """Light campus registration — name/email/phone/password/college only."""
    if not _campus_enabled(db):
        raise HTTPException(status_code=403, detail="Campus drive registration is not currently open.")
    try:
        validate_password_strength(password)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    if not re.match(r"^[6-9]\d{9}$", phone.replace(" ", "")):
        raise HTTPException(status_code=422, detail="Phone must be a valid 10-digit Indian mobile number.")
    personal_email = _validate_registration_email(personal_email)
    try:
        candidate = candidate_service.create_campus_candidate(
            db, full_name=full_name, personal_email=personal_email, phone=phone,
            password=password, college_id=college_id, position_id=position_id, request=request,
        )
        db.commit()
    except IntegrityError as exc:
        db.rollback()
        raise HTTPException(status_code=409, detail="A candidate with this email already exists.") from exc
    return {
        "candidateId": candidate.id,
        "email": candidate.personal_email,
        "message": "Registration successful. Sign in to take your assessment.",
    }


@router.post("/campus/complete")
@limiter.limit("5/minute")
def complete_campus_registration(
    request: Request,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
    aadhaar_card: Annotated[UploadFile, File(alias="aadhaarCard")],
    gender: str = Form(),
    experience_type: str = Form(alias="experienceType"),
    experience_years: str | None = Form(default=None, alias="experienceYears"),
    aadhaar_number: str = Form(alias="aadhaarNumber"),
    date_of_birth: str | None = Form(default=None, alias="dateOfBirth"),
    resume: Annotated[UploadFile | None, File()] = None,
) -> dict:
    """Full registration for a campus candidate who passed — same Aadhaar OCR / resume
    parsing as normal, upgrading their existing record + skipping the assessment stage."""
    candidate = candidate_service.get_campus_candidate_for_user(db, current_user)
    if candidate is None:
        raise HTTPException(status_code=404, detail="No campus registration found for your account.")
    if candidate.resume_url:
        raise HTTPException(status_code=409, detail="Your registration is already complete.")
    released_pass = db.scalar(
        select(ApAttempt)
        .join(ApAssignment, ApAttempt.assignment_id == ApAssignment.id)
        .where(
            ApAssignment.user_id == current_user.id,
            ApAttempt.result_status == "pass",
            ApAttempt.result_released_at.isnot(None),
        )
    )
    if released_pass is None:
        raise HTTPException(status_code=403, detail="Your assessment result is not available yet.")

    # ── identical document handling to the normal registration ──
    validate_candidate_upload(
        aadhaar_card, label="Aadhaar card",
        allowed_content_types=ALLOWED_DOCUMENT_IMAGE_CONTENT_TYPES,
        allowed_extensions=ALLOWED_DOCUMENT_IMAGE_EXTENSIONS, required=True,
    )
    validate_candidate_upload(
        resume, label="Resume",
        allowed_content_types=ALLOWED_RESUME_CONTENT_TYPES,
        allowed_extensions=ALLOWED_RESUME_EXTENSIONS, required=True,
    )
    normalized_experience_type = experience_type.strip().lower()
    if normalized_experience_type not in {"fresher", "experienced"}:
        raise HTTPException(status_code=422, detail="Experience type must be either fresher or experienced.")
    resolved_experience_years = parse_experience_years(
        experience_years, required=normalized_experience_type == "experienced"
    )
    if resume is None or not resume.filename:
        raise HTTPException(status_code=400, detail="Resume upload is required.")

    aadhaar_ocr = extract_aadhaar_fields(aadhaar_card)
    aadhaar_validation = validate_aadhaar_identity(
        entered_name=candidate.full_name, entered_aadhaar=aadhaar_number,
        entered_dob=date_of_birth, ocr_result=aadhaar_ocr,
    )
    aadhaar_validation = _downgrade_candidate_aadhaar_failure_for_review(aadhaar_validation)
    if aadhaar_validation["validationStatus"] not in ("passed", "partial", "needs_review"):
        raise HTTPException(status_code=400, detail="Invalid Documents")

    entered_aadhaar_digits = re.sub(r"\D", "", aadhaar_number)
    ocr_candidate_digits = {
        re.sub(r"\D", "", str(value or ""))
        for value in [
            aadhaar_ocr.get("aadhaarNumber"),
            *(aadhaar_ocr.get("aadhaarNumberCandidates") or []),
        ]
    }
    if len(entered_aadhaar_digits) == 12 and entered_aadhaar_digits in ocr_candidate_digits:
        aadhaar_ocr["aadhaarNumber"] = entered_aadhaar_digits

    resume_url = save_candidate_upload(resume, "resumes")
    aadhaar_url = save_candidate_upload(aadhaar_card, "aadhaar")
    resolved_dob = parse_iso_date(date_of_birth) or parse_iso_date(aadhaar_ocr.get("dateOfBirth"))
    resume_text_extracted = _extract_resume_text(
        read_upload_bytes(resume),
        Path(resume.filename or "").suffix.lower(),
        (resume.content_type or "").split(";")[0].lower(),
    )
    resume_key_points: list[str] = []
    resume_summary: str | None = None
    if resume_text_extracted.strip():
        try:
            from app.services.integrations import LLMService
            parsed = LLMService().parse_resume(resume_text=resume_text_extracted, job_title=None)
            resume_key_points = parsed.get("keyPoints") or []
            resume_summary = parsed.get("summary") or None
        except Exception:
            pass

    from app.core.security import fingerprint_identifier
    fields = {
        "gender": gender,
        "experience_type": normalized_experience_type,
        "experience_years": resolved_experience_years,
        "date_of_birth": resolved_dob,
        "aadhaar_hash": fingerprint_identifier(aadhaar_number) if aadhaar_number else None,
        "aadhaar_last4": (aadhaar_number or "")[-4:] or None,
        "aadhaar_extracted": aadhaar_ocr,
        "aadhaar_ocr_name": aadhaar_validation.get("ocrName"),
        "aadhaar_validation_status": aadhaar_validation.get("validationStatus"),
        "aadhaar_mismatch_reason": aadhaar_validation.get("mismatchReason"),
        "resume_url": resume_url,
        "resume_text": resume_text_extracted[:20000] if resume_text_extracted else None,
        "resume_key_points": resume_key_points or None,
        "resume_summary": resume_summary,
    }
    candidate_service.complete_campus_candidate(db, candidate=candidate, fields=fields, request=request)
    record_candidate_upload_document(
        db,
        candidate=candidate,
        file=resume,
        type_="resume",
        file_url=resume_url,
    )
    record_candidate_upload_document(
        db,
        candidate=candidate,
        file=aadhaar_card,
        type_="aadhaar_card",
        file_url=aadhaar_url,
        ocr_result=aadhaar_ocr,
    )
    try:
        db.commit()
    except IntegrityError as exc:
        db.rollback()
        raise HTTPException(status_code=409, detail="These documents are already linked to another candidate.") from exc
    db.refresh(candidate)
    return {"candidateId": candidate.id, "message": "Registration complete."}


@router.post("/{candidate_id}/screen")
@limiter.limit("10/minute")
def trigger_screening(
    request: Request,
    candidate_id: str,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.CANDIDATES_WRITE))],
) -> dict:
    """Trigger LLM resume screening for a candidate."""
    _require_candidate_management_access(current_user)
    # Screening is a recruiting-staff action. Vendors and employee referrers hold
    # CANDIDATES_WRITE so they can submit referrals, but they must not screen.
    if _is_restricted_candidate_viewer(current_user):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You are not permitted to run resume screening.",
        )
    # Enforce per-record access (404/403 for out-of-scope candidates) before screening.
    candidate_service.get_candidate_or_404(db, candidate_id, current_user=current_user)
    from app.services import workflows as wf
    candidate = wf.run_resume_screening(
        db,
        candidate_id=candidate_id,
        actor=current_user,
        request=request,
    )
    db.commit()
    db.refresh(candidate)
    return {
        "candidateId": candidate.id,
        "resumeScore": candidate.resume_score,
        "resumeSummary": candidate.resume_summary,
        "llmStatus": candidate.llm_status,
        "currentStage": candidate.current_stage.value,
        "screeningPayload": candidate.screening_payload,
    }


@router.post("", response_model=CandidateSummary)
def create_candidate(
    payload: CreateCandidateRequest,
    request: Request,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.CANDIDATES_WRITE))],
):
    candidate = candidate_service.create_candidate(
        db,
        payload=payload.model_dump(),
        actor=current_user,
        request=request,
    )
    db.commit()
    db.refresh(candidate)
    return candidate


@router.get("", response_model=CandidateListResponse)
def list_candidates(
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.CANDIDATES_READ))],
    search: str | None = None,
    source_type: str | None = Query(default=None, alias="sourceType"),
    stage: str | None = None,
    position_id: str | None = Query(default=None, alias="positionId"),
    created_from: date | None = Query(default=None, alias="createdFrom"),
    created_to: date | None = Query(default=None, alias="createdTo"),
    blacklisted: bool | None = Query(default=None),
    page: int = 1,
    limit: int = 20,
    sort_by: str = Query(default="createdAt", alias="sortBy"),
    sort_dir: str = Query(default="desc", alias="sortDir"),
):
    created_from_dt, created_to_dt = _date_range_bounds(created_from, created_to)
    items, total = candidate_service.list_candidates(
        db,
        current_user=current_user,
        search=search,
        source_type=source_type,
        stage=stage,
        position_id=position_id,
        created_from=created_from_dt,
        created_to=created_to_dt,
        blacklisted=blacklisted,
        page=page,
        limit=limit,
        sort_by=sort_by,
        sort_dir=sort_dir,
    )
    total_pages = (total + limit - 1) // limit if limit else 1
    # Restricted viewers receive hand-serialized projections so hidden fields are
    # truly omitted rather than nulled by a response_model round-trip.
    if _is_restricted_candidate_viewer(current_user) or _is_preview_candidate_viewer(current_user):
        scrubbed = []
        for item in items:
            payload = CandidateSummary.model_validate(item).model_dump(by_alias=True)
            if _is_restricted_candidate_viewer(current_user):
                payload = scrub_internal_candidate_fields(payload)
                payload["accessLevel"] = "scoped"
                payload["canOpenDetail"] = True
            else:
                payload = scrub_candidate_preview_fields(payload)
            scrubbed.append(payload)
        return JSONResponse(
            content=jsonable_encoder({
                "data": scrubbed,
                "total": total,
                "page": page,
                "limit": limit,
                "totalPages": total_pages,
            })
        )
    return {
        "data": items,
        "total": total,
        "page": page,
        "limit": limit,
        "totalPages": total_pages,
    }


@router.get("/stats", response_model=CandidateStatsResponse)
def candidate_stats(
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.CANDIDATES_READ))],
):
    stats = candidate_service.get_pipeline_stats(db, current_user=current_user)
    return {"stages": stats["stages"], "total": stats["total"], "thisMonth": stats["this_month"]}


@router.post("/employee-codes/backfill-signed", response_model=MessageResponse)
def backfill_signed_candidate_employee_codes(
    request: Request,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.CANDIDATES_WRITE))],
):
    _require_candidate_management_access(current_user)
    updated = candidate_service.backfill_signed_candidate_employee_codes(
        db,
        actor=current_user,
        request=request,
    )
    db.commit()
    return MessageResponse(
        message=f"Assigned employee codes to {len(updated)} signed candidate(s)."
    )


@router.get("/export")
def export_candidates_csv(
    request: Request,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.CANDIDATES_READ))],
    search: str | None = None,
    source_type: str | None = Query(default=None, alias="sourceType"),
    stage: str | None = None,
    position_id: str | None = Query(default=None, alias="positionId"),
    created_from: date | None = Query(default=None, alias="createdFrom"),
    created_to: date | None = Query(default=None, alias="createdTo"),
    blacklisted: bool | None = Query(default=None),
):
    """Server-side CSV export of candidates INCLUDING signed, openable links to
    each uploaded document (resume, Aadhaar card, and any other documents). The
    links are HMAC-signed and time-limited so they open directly from the
    downloaded spreadsheet without a separate login. Honours the same vendor /
    referrer access scope as the candidate list."""
    import csv as _csv
    from io import StringIO
    import json as _json

    from fastapi.responses import Response
    from sqlalchemy import select
    from sqlalchemy.orm import joinedload, selectinload

    from app.core.signed_urls import make_signed_upload_url
    from app.db.models import Document, Evaluation, PiInterviewRound, SelectionForm

    created_from_dt, created_to_dt = _date_range_bounds(created_from, created_to)
    items, _total = candidate_service.list_candidates(
        db,
        current_user=current_user,
        search=search,
        source_type=source_type,
        stage=stage,
        position_id=position_id,
        created_from=created_from_dt,
        created_to=created_to_dt,
        blacklisted=blacklisted,
        page=1,
        limit=5000,
        sort_by="created_at",
        sort_dir="desc",
    )

    # Non-full-detail users must not receive internal recruiting data in exports
    # either (resume scores/summaries, screening status, CTC, full Aadhaar/OCR
    # details, evaluations). Full recruiting users get the full export.
    restricted = _is_restricted_candidate_viewer(current_user) or _is_preview_candidate_viewer(current_user)

    # Gather all documents for the returned candidates in one query (avoid N+1).
    candidate_ids = [c.id for c in items]
    docs_by_candidate: dict[str, list[Document]] = {}
    if candidate_ids:
        for doc in db.scalars(
            select(Document).where(Document.candidate_id.in_(candidate_ids))
        ):
            docs_by_candidate.setdefault(doc.candidate_id, []).append(doc)

    selection_forms_by_candidate: dict[str, SelectionForm] = {}
    if candidate_ids:
        for form in db.scalars(
            select(SelectionForm).where(SelectionForm.candidate_id.in_(candidate_ids))
        ):
            selection_forms_by_candidate[form.candidate_id] = form

    evaluations_by_candidate: dict[str, list[Evaluation]] = {}
    if candidate_ids and not restricted:
        evaluation_query = (
            select(Evaluation)
            .options(
                joinedload(Evaluation.evaluator),
                selectinload(Evaluation.pi_rounds).joinedload(PiInterviewRound.evaluator),
            )
            .where(Evaluation.candidate_id.in_(candidate_ids))
        )
        for evaluation in db.scalars(evaluation_query):
            evaluations_by_candidate.setdefault(evaluation.candidate_id, []).append(evaluation)
        for evaluations in evaluations_by_candidate.values():
            evaluations.sort(
                key=lambda item: item.completed_at or item.created_at or datetime.min.replace(tzinfo=UTC),
                reverse=True,
            )

    def _signed(url: str | None) -> str:
        if not url:
            return ""
        try:
            return make_signed_upload_url(url)
        except Exception:
            return ""

    def _fmt_dt(value):
        return format_app_datetime(value) if value else ""

    def _json_cell(value) -> str:
        if value in (None, "", [], {}):
            return ""
        try:
            return _json.dumps(jsonable_encoder(value), ensure_ascii=False, sort_keys=True)
        except Exception:
            return str(value)

    def _clean_text(value: str | None) -> str:
        return re.sub(r"\s+", " ", value or "").strip()

    def _cell_text(value) -> str:
        if value is None:
            return ""
        return str(value).strip()

    def _selection_data(candidate_id: str) -> dict:
        form = selection_forms_by_candidate.get(candidate_id)
        form_data = getattr(form, "form_data", None)
        return form_data if isinstance(form_data, dict) else {}

    def _selection_text(candidate_id: str, *path: str) -> str:
        current = _selection_data(candidate_id)
        for part in path:
            if not isinstance(current, dict):
                return ""
            current = current.get(part)
        return _cell_text(current)

    def _profile_text(c: Candidate, attr: str, *paths: tuple[str, ...]) -> str:
        value = _cell_text(getattr(c, attr, None))
        if value:
            return value
        for path in paths:
            value = _selection_text(c.id, *path)
            if value:
                return value
        return ""

    def _profile_date(c: Candidate, attr: str, *paths: tuple[str, ...]) -> datetime | None:
        value = getattr(c, attr, None)
        if isinstance(value, datetime):
            return value
        ocr = c.aadhaar_extracted if isinstance(c.aadhaar_extracted, dict) else {}
        raw_values = [_selection_text(c.id, *path) for path in paths]
        raw_values.extend(
            _cell_text(ocr.get(key))
            for key in ("dateOfBirth", "date_of_birth", "dob")
        )
        for raw in raw_values:
            if not raw:
                continue
            parsed = parse_iso_date(raw)
            if parsed is None:
                parsed_text = parse_date_of_birth(f"DOB: {raw}")
                parsed = parse_iso_date(parsed_text)
            if parsed is not None:
                return parsed
        return None

    def _aadhaar_last4(c: Candidate, ocr: dict) -> str:
        value = _cell_text(c.aadhaar_last4)
        if value:
            return value
        for raw in (
            _selection_text(c.id, "identityDetails", "aadhaarNumber"),
            _selection_text(c.id, "personalDetails", "aadhaarNumber"),
            _selection_text(c.id, "aadhaarNumber"),
            _cell_text(ocr.get("aadhaarNumber")),
        ):
            digits = re.sub(r"\D", "", raw)
            if len(digits) == 12:
                return digits[-4:]
        return ""

    def _doc_type_label(raw: str | None) -> str:
        key = (raw or "document").strip().lower()
        return {"aadhaar_card": "aadhaar"}.get(key, key)

    def _pi_rounds_for(evaluations: list[Evaluation]) -> list[PiInterviewRound]:
        rounds: list[PiInterviewRound] = []
        for evaluation in evaluations:
            rounds.extend(evaluation.pi_rounds or [])
        return sorted(
            rounds,
            key=lambda item: (
                item.completed_at or item.scheduled_at or item.created_at or datetime.min.replace(tzinfo=UTC),
                item.round_number,
            ),
            reverse=True,
        )

    def _evaluation_details(evaluations: list[Evaluation]) -> str:
        details: list[str] = []
        for index, evaluation in enumerate(evaluations, start=1):
            pi_rounds = sorted(evaluation.pi_rounds or [], key=lambda item: item.round_number)
            pi_details = "; ".join(
                " ".join(
                    part
                    for part in [
                        f"R{round_record.round_number}",
                        f"score={round_record.score}" if round_record.score is not None else "",
                        f"status={round_record.status}" if round_record.status else "",
                        f"decision={round_record.round_decision}" if round_record.round_decision else "",
                        f"verdict={round_record.final_verdict}" if round_record.final_verdict else "",
                        f"remarks={_clean_text(round_record.remarks)}" if round_record.remarks else "",
                    ]
                    if part
                )
                for round_record in pi_rounds
            )
            details.append(
                " | ".join(
                    part
                    for part in [
                        f"Evaluation {index}",
                        f"evaluator={evaluation.evaluator.name}" if evaluation.evaluator else "",
                        f"totalScore={evaluation.total_score}" if evaluation.total_score is not None else "",
                        f"recommendation={evaluation.recommendation}" if evaluation.recommendation else "",
                        f"technical={evaluation.technical_skills}" if evaluation.technical_skills is not None else "",
                        f"communication={evaluation.communication}" if evaluation.communication is not None else "",
                        f"problemSolving={evaluation.problem_solving}" if evaluation.problem_solving is not None else "",
                        f"culturalFit={evaluation.cultural_fit}" if evaluation.cultural_fit is not None else "",
                        f"attitude={evaluation.attitude}" if evaluation.attitude is not None else "",
                        f"piScore={evaluation.pi_score}" if evaluation.pi_score is not None else "",
                        f"pmsScore={evaluation.pms_score}" if evaluation.pms_score is not None else "",
                        f"completedAt={_fmt_dt(evaluation.completed_at)}" if evaluation.completed_at else "",
                        f"notes={_clean_text(evaluation.notes)}" if evaluation.notes else "",
                        f"piRounds=[{pi_details}]" if pi_details else "",
                    ]
                    if part
                )
            )
        return " || ".join(details)

    # Build a {doc-type -> file_url} map per candidate, and collect the union of
    # all document types so each document gets its OWN column (one link per cell).
    per_candidate_docs: dict[str, dict[str, str]] = {}
    doc_types: set[str] = {"resume", "aadhaar"}
    for c in items:
        type_map: dict[str, str] = {}
        if c.resume_url:
            type_map["resume"] = c.resume_url
        for d in docs_by_candidate.get(c.id, []):
            if not d.file_url:
                continue
            label = _doc_type_label(d.type)
            type_map.setdefault(label, d.file_url)  # first wins; keep resume_url if already set
            doc_types.add(label)
        per_candidate_docs[c.id] = type_map

    # Stable, readable column order: resume + aadhaar first, then the rest sorted.
    ordered_types = ["resume", "aadhaar"] + sorted(t for t in doc_types if t not in {"resume", "aadhaar"})
    doc_columns = [f"{t.replace('_', ' ').title()} Link" for t in ordered_types]

    # CSV formula-injection guard: a leading =, +, -, @, tab or CR makes spreadsheets
    # treat the cell as a formula. Prefix such string cells with a single quote so
    # they're rendered as literal text. Applied at the single write choke point below.
    def _csv_safe(value):
        if not isinstance(value, str) or not value:
            return value
        if value[0] in ("=", "+", "-", "@", "\t", "\r"):
            return "'" + value
        return value

    def _write_row(row: list) -> None:
        writer.writerow([_csv_safe(cell) for cell in row])

    buffer = StringIO()
    writer = _csv.writer(buffer)
    if restricted:
        _write_row([
            "Name", "Personal Email", "Ethara Email", "Phone", "Candidate Code", "Employee Code",
            "Gender", "Date of Birth", "Marital Status",
            "Aadhaar Last4",
            "Experience Type", "Experience (years)", "Current Company",
            "Notice Period (days)",
            "Source", "Source Id", "Position", "College", "Vendor",
            "Stage", "Status", "Duplicate",
            "Created", "Last Applied",
            *doc_columns,
        ])
    else:
        _write_row([
            "Candidate Id", "Candidate Code", "Employee Code", "Name", "Personal Email", "Ethara Email", "Phone",
            "Gender", "Date of Birth", "Marital Status",
            "Aadhaar Number (OCR)", "Aadhaar Last4", "Aadhaar Hash", "Aadhaar Name (OCR)",
            "Aadhaar DOB (OCR)", "Aadhaar OCR Status", "Aadhaar Validation",
            "Aadhaar Mismatch Reason",
            "Experience Type", "Experience (years)", "Current Company",
            "Current CTC", "Expected CTC", "Notice Period (days)",
            "Source", "Source Id", "Position Id", "Position", "College Id", "College",
            "Vendor Id", "Vendor", "Portal User Id",
            "Stage", "Status", "Priority Score", "Duplicate", "Duplicate Reason",
            "Reapplication Blocked", "Is Removed",
            "Resume Url", "Resume Score", "Resume Summary", "Resume Text",
            "Resume Key Points", "Screening Payload", "Screening Status",
            "Evaluation Count", "Latest Evaluation Score", "Latest Recommendation",
            "Latest Evaluator", "Latest Evaluation Completed",
            "Latest PI Score", "Latest PMS Score", "PI Round Count",
            "PI Round Scores", "PI Round Decisions", "PI Final Verdicts",
            "Evaluation Details",
            "Created", "Updated", "Last Applied",
            *doc_columns,
        ])

    for c in items:
        ocr = c.aadhaar_extracted if isinstance(c.aadhaar_extracted, dict) else {}
        type_map = per_candidate_docs.get(c.id, {})
        profile_name = _profile_text(c, "full_name", ("basicDetails", "fullName"), ("fullName",))
        profile_personal_email = _profile_text(
            c,
            "personal_email",
            ("basicDetails", "email"),
            ("personalEmail",),
            ("email",),
        )
        profile_phone = _profile_text(
            c,
            "phone",
            ("basicDetails", "contactNumber"),
            ("contactNumber",),
            ("phone",),
            ("mobileNumber",),
        )
        profile_gender = _profile_text(
            c,
            "gender",
            ("personalDetails", "gender"),
            ("basicDetails", "gender"),
            ("gender",),
        )
        profile_date_of_birth = _profile_date(
            c,
            "date_of_birth",
            ("basicDetails", "dateOfBirth"),
            ("personalDetails", "dateOfBirth"),
            ("identityDetails", "dateOfBirth"),
            ("dateOfBirth",),
        )
        profile_marital_status = _profile_text(
            c,
            "marital_status",
            ("personalDetails", "maritalStatus"),
            ("maritalStatus",),
        )
        profile_aadhaar_last4 = _aadhaar_last4(c, ocr)
        profile_experience_type = _profile_text(
            c,
            "experience_type",
            ("basicDetails", "experienceType"),
            ("experienceType",),
        )
        if restricted:
            row = [
                profile_name,
                profile_personal_email,
                c.ethara_email or "",
                profile_phone,
                c.candidate_code or "",
                c.employee_code or "",
                profile_gender,
                _fmt_dt(profile_date_of_birth),
                profile_marital_status,
                profile_aadhaar_last4,
                profile_experience_type,
                "" if c.experience_years is None else c.experience_years,
                c.current_company or "",
                "" if c.notice_period is None else c.notice_period,
                getattr(c.source_type, "value", c.source_type) or "",
                c.source_id or "",
                (c.position.title if c.position else ""),
                (c.college.name if c.college else ""),
                (c.vendor.name if c.vendor else ""),
                getattr(c.current_stage, "value", c.current_stage) or "",
                c.current_status or "",
                "Yes" if c.is_duplicate else "No",
                _fmt_dt(c.created_at),
                _fmt_dt(c.last_applied_at),
            ]
        else:
            evaluations = evaluations_by_candidate.get(c.id, [])
            latest_evaluation = evaluations[0] if evaluations else None
            pi_rounds = _pi_rounds_for(evaluations)
            row = [
                c.id or "",
                c.candidate_code or "",
                c.employee_code or "",
                profile_name,
                profile_personal_email,
                c.ethara_email or "",
                profile_phone,
                profile_gender,
                _fmt_dt(profile_date_of_birth),
                profile_marital_status,
                ocr.get("aadhaarNumber") or "",
                profile_aadhaar_last4,
                c.aadhaar_hash or "",
                c.aadhaar_ocr_name or ocr.get("cardHolderName") or "",
                ocr.get("dateOfBirth") or "",
                ocr.get("ocrStatus") or ocr.get("message") or "",
                c.aadhaar_validation_status or "",
                c.aadhaar_mismatch_reason or "",
                profile_experience_type,
                "" if c.experience_years is None else c.experience_years,
                c.current_company or "",
                "" if c.current_ctc is None else c.current_ctc,
                "" if c.expected_ctc is None else c.expected_ctc,
                "" if c.notice_period is None else c.notice_period,
                getattr(c.source_type, "value", c.source_type) or "",
                c.source_id or "",
                c.position_id or "",
                (c.position.title if c.position else ""),
                c.college_id or "",
                (c.college.name if c.college else ""),
                c.vendor_id or "",
                (c.vendor.name if c.vendor else ""),
                c.portal_user_id or "",
                getattr(c.current_stage, "value", c.current_stage) or "",
                c.current_status or "",
                c.priority_score if c.priority_score is not None else "",
                "Yes" if c.is_duplicate else "No",
                c.duplicate_reason or "",
                "Yes" if c.is_reapplication_blocked else "No",
                "Yes" if c.is_removed else "No",
                c.resume_url or "",
                "" if c.resume_score is None else c.resume_score,
                (c.resume_summary or "").replace("\n", " ").strip(),
                _clean_text(c.resume_text),
                _json_cell(c.resume_key_points),
                _json_cell(c.screening_payload),
                c.llm_status or "",
                len(evaluations),
                "" if latest_evaluation is None or latest_evaluation.total_score is None else latest_evaluation.total_score,
                latest_evaluation.recommendation if latest_evaluation else "",
                latest_evaluation.evaluator.name if latest_evaluation and latest_evaluation.evaluator else "",
                _fmt_dt(latest_evaluation.completed_at) if latest_evaluation else "",
                "" if latest_evaluation is None or latest_evaluation.pi_score is None else latest_evaluation.pi_score,
                "" if latest_evaluation is None or latest_evaluation.pms_score is None else latest_evaluation.pms_score,
                len(pi_rounds),
                "; ".join(str(round_record.score) for round_record in pi_rounds if round_record.score is not None),
                "; ".join(round_record.round_decision or "" for round_record in pi_rounds).strip("; "),
                "; ".join(round_record.final_verdict or "" for round_record in pi_rounds).strip("; "),
                _evaluation_details(evaluations),
                _fmt_dt(c.created_at),
                _fmt_dt(c.updated_at),
                _fmt_dt(c.last_applied_at),
            ]
        row.extend(_signed(type_map.get(t)) for t in ordered_types)
        _write_row(row)

    # Audit the export (who exported, with what filters, how many rows).
    log_audit(
        db,
        entity_type="candidate",
        entity_id="*",
        action="candidates_exported",
        actor=current_user,
        request=request,
        new_value={
            "filters": {
                "search": search,
                "sourceType": source_type,
                "stage": stage,
                "positionId": position_id,
            },
            "count": len(items),
        },
    )
    db.commit()

    filename = f"candidates_{app_date_stamp()}.csv"
    return Response(
        content=buffer.getvalue(),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/me", response_model=CandidatePortalSelfOverview)
def get_my_candidate_portal(
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
):
    # Candidate-SAFE response: CandidatePortalSelfOverview projects onto the trimmed
    # CandidatePortalApplication, dropping staff/internal fields (evaluations,
    # audit_logs, escalations, screening_payload, resume scores/text, full Aadhaar)
    # that the staff CandidateDetail carries and the portal never reads.
    user = _require_candidate_user(current_user)
    current_application = candidate_service.get_portal_candidate_detail(db, user=user)
    applications = candidate_service.list_portal_candidates(db, user=user)
    return {
        "currentApplication": current_application,
        "applications": applications,
        "emailVerified": bool(user.email_verified_at),
        "emailVerifiedAt": user.email_verified_at,
    }


@router.post("/me/compliance/refresh", response_model=CandidatePortalSelfOverview)
def refresh_my_candidate_compliance(
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
):
    """Candidate pulls the latest signing status of their Documenso statutory/compliance forms.
    When all are signed, this marks onboarding complete and issues the HRMS employee credentials."""
    user = _require_candidate_user(current_user)
    candidate = db.scalar(
        select(Candidate)
        .where(Candidate.portal_user_id == user.id)
        .order_by(Candidate.created_at.desc())
    )
    if candidate is not None:
        try:
            from app.services import compliance_documenso as compliance_esign

            compliance_esign.sync_candidate_compliance(db, candidate=candidate)
            db.commit()
        except Exception:
            db.rollback()
            logger.exception("candidate compliance sync failed for %s", candidate.id)
    return {
        "currentApplication": candidate_service.get_portal_candidate_detail(db, user=user),
        "applications": candidate_service.list_portal_candidates(db, user=user),
        "emailVerified": bool(user.email_verified_at),
        "emailVerifiedAt": user.email_verified_at,
    }


@router.patch("/me/profile", response_model=CandidateSummary)
def update_my_candidate_profile(
    payload: PortalProfileUpdateRequest,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
):
    user = _require_candidate_user(current_user)
    updated = candidate_service.update_portal_profile(
        db,
        user=user,
        payload=payload.model_dump(exclude_none=True),
    )
    db.commit()
    db.refresh(updated)
    return updated


@router.get("/me/id-card-form", response_model=CandidateIdCardFormRead)
def get_my_candidate_id_card_form(
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
):
    user = _require_candidate_user(current_user)
    candidate = candidate_service.get_portal_candidate_detail(db, user=user)
    if candidate is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Candidate profile not found")
    _require_id_card_portal_access(candidate)
    return workflows.get_candidate_id_card_form(db, candidate=candidate)


@router.post("/me/id-card-form", response_model=CandidateIdCardFormRead)
def submit_my_candidate_id_card_form(
    payload: CandidateIdCardFormSubmitRequest,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
):
    user = _require_candidate_user(current_user)
    candidate = candidate_service.get_portal_candidate_detail(db, user=user)
    if candidate is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Candidate profile not found")
    _require_id_card_portal_access(candidate)
    updated = workflows.submit_candidate_id_card_form(
        db,
        candidate=candidate,
        actor=user,
        payload=payload.model_dump(),
    )
    db.commit()
    return updated


@router.post("/me/apply", response_model=CandidateSummary)
def apply_for_position(
    payload: PortalApplyRequest,
    request: Request,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
):
    user = _require_candidate_user(current_user)
    updated = candidate_service.apply_to_position(
        db,
        user=user,
        position_id=payload.position_id,
        request=request,
    )
    db.commit()
    db.refresh(updated)
    return updated


@router.get("/{candidate_id}", response_model=CandidateDetail)
def get_candidate(
    candidate_id: str,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.CANDIDATES_READ))],
):
    _require_full_candidate_detail_access(current_user)
    candidate = candidate_service.get_candidate_or_404(db, candidate_id, current_user=current_user)
    # Vendors / employee referrers see a scrubbed detail (internal recruiting fields
    # omitted). Staff fall through to the normal response_model with output unchanged.
    if _is_restricted_candidate_viewer(current_user):
        scrubbed = scrub_internal_candidate_fields(
            CandidateDetail.model_validate(candidate).model_dump(by_alias=True)
        )
        return JSONResponse(content=jsonable_encoder(scrubbed))
    return candidate


@router.get("/{candidate_id}/resume/download")
def download_candidate_resume(
    candidate_id: str,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.CANDIDATES_READ))],
):
    _require_full_candidate_detail_access(current_user)
    candidate = candidate_service.get_candidate_or_404(db, candidate_id, current_user=current_user)
    resume_url = (candidate.resume_url or "").strip()
    if not resume_url:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Resume not found")
    if not resume_url.startswith("/uploads/"):
        from app.services.integrations import StorageService

        download_url = StorageService().presigned_download_url(resume_url)
        if download_url:
            return RedirectResponse(download_url)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Resume is stored externally and cannot be proxied from this server.",
        )

    settings = get_settings()
    relative = resume_url.removeprefix("/")
    local_path = settings.local_storage_path.parent / relative
    if not local_path.exists():
        local_path = settings.local_storage_path / relative.removeprefix("uploads/")
    if not local_path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Resume file not found")

    media_type = candidate.resume_url.lower()
    if media_type.endswith(".pdf"):
        mime_type = "application/pdf"
    elif media_type.endswith(".doc"):
        mime_type = "application/msword"
    elif media_type.endswith(".docx"):
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        mime_type = "application/octet-stream"

    return FileResponse(
        path=str(local_path),
        filename=local_path.name,
        media_type=mime_type,
    )


@router.patch("/{candidate_id}", response_model=CandidateSummary)
def update_candidate(
    candidate_id: str,
    payload: UpdateCandidateRequest,
    request: Request,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.CANDIDATES_WRITE))],
):
    _require_candidate_management_access(current_user)
    candidate = candidate_service.get_candidate_or_404(db, candidate_id, current_user=current_user)
    updated = candidate_service.update_candidate(
        db,
        candidate=candidate,
        payload=payload.model_dump(exclude_unset=True),
        actor=current_user,
        request=request,
    )
    db.commit()
    db.refresh(updated)
    return updated


@router.delete("/{candidate_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_candidate(
    candidate_id: str,
    request: Request,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.CANDIDATES_WRITE))],
):
    _require_candidate_management_access(current_user)
    candidate = candidate_service.get_candidate_or_404(db, candidate_id, current_user=current_user)
    candidate_service.soft_delete_candidate(db, candidate=candidate, actor=current_user, request=request)
    db.commit()


@router.post("/{candidate_id}/advance-stage", response_model=CandidateSummary)
def advance_stage(
    candidate_id: str,
    payload: AdvanceStageRequest,
    request: Request,
    db: Annotated[Session, Depends(get_db)],
    current_user: Annotated[User, Depends(require_permissions(Permission.CANDIDATES_WRITE))],
):
    _require_candidate_management_access(current_user)
    # Stage progression is a recruiting-staff action. Vendors and employee referrers
    # hold CANDIDATES_WRITE for their own referrals but must not advance stages.
    if _is_restricted_candidate_viewer(current_user):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You are not permitted to advance candidate stages.",
        )
    candidate = candidate_service.get_candidate_or_404(db, candidate_id, current_user=current_user)
    updated = candidate_service.advance_stage(
        db,
        candidate=candidate,
        to_stage=payload.to_stage,
        notes=payload.notes,
        actor=current_user,
        request=request,
    )
    db.commit()
    db.refresh(updated)
    return updated
