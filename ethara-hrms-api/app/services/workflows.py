from __future__ import annotations

import csv
import copy
import mimetypes
import re
from io import BytesIO, StringIO
from pathlib import Path
from datetime import UTC, datetime, timedelta
from typing import Any
from urllib.request import urlopen

from fastapi import HTTPException, Request, UploadFile, status
from sqlalchemy import Select, func, or_, select
from sqlalchemy.orm import Session, joinedload, selectinload

from app.core.config import get_settings
from app.core.timezone import format_app_datetime
from app.db.models import (
    Candidate,
    CandidateIdCardForm,
    CandidateStage,
    ComplianceForm,
    Contract,
    ContractStatus,
    Document,
    Escalation,
    EscalationStatus,
    Evaluation,
    ITRequest,
    Notification,
    NotificationType,
    PiInterviewRound,
    Position,
    Role,
    SelectionForm,
    StageLog,
    User,
)
from app.core.signed_urls import make_signed_upload_url
from app.services.audit import log_audit
from app.services.integrations import LLMService, OCRService, StorageService


SLA_HOURS: dict[CandidateStage, int] = {
    CandidateStage.RESUME_SCREENING_PENDING: 48,
    CandidateStage.EVALUATION_ASSIGNED: 72,
    CandidateStage.SELECTION_FORM_SENT: 96,
    CandidateStage.CONTRACT_SENT: 120,
    CandidateStage.IT_EMAIL_CREATED: 24,
    CandidateStage.STATUTORY_FORMS_SENT: 168,
}
DEFAULT_EVALUATOR_EMAIL = "evaluator@ethara.ai"
# Staff roles allowed to re-open / overwrite an already-finalized evaluation or
# to complete a PI round assigned to another evaluator.
_EVALUATION_STAFF_ROLES = {Role.SUPER_ADMIN, Role.ADMIN, Role.LEADERSHIP, Role.HR, Role.TA}
MAX_PI_ROUNDS = 5
PI_ACTIVE_STATUSES = {"scheduled", "rescheduled"}
PI_FINAL_VERDICTS = {"selected", "rejected"}
SELECTION_FORM_VERIFICATION_QUEUE_KEY = "verificationQueue"


def _role_value(role: Role | str | None) -> str:
    if role is None:
        return ""
    return role.value if isinstance(role, Role) else str(role)


def _user_role_values(user: User) -> set[str]:
    return {_role_value(user.role)} | {_role_value(role) for role in (user.roles or [])}


def _has_any_role(user: User, roles: set[Role]) -> bool:
    allowed = {_role_value(role) for role in roles}
    return bool(_user_role_values(user) & allowed)


def _normalize_selection_form_payload(form_data: dict[str, Any]) -> dict[str, Any]:
    normalized = dict(form_data)
    # Server-managed status; never trust a value posted by the browser.
    normalized.pop(SELECTION_FORM_VERIFICATION_QUEUE_KEY, None)
    references = normalized.get("references")
    if not isinstance(references, list):
        return normalized

    cleaned_references: list[Any] = []
    for index, reference in enumerate(references):
        if not isinstance(reference, dict):
            cleaned_references.append(reference)
            continue

        cleaned_reference = dict(reference)
        phone_raw = str(cleaned_reference.get("phone") or "").strip()
        if phone_raw:
            digits = re.sub(r"\D", "", phone_raw)
            if len(digits) != 10:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail=f"Reference {index + 1} phone number must be exactly 10 digits.",
                )
            cleaned_reference["phone"] = digits

        cleaned_references.append(cleaned_reference)

    normalized["references"] = cleaned_references
    return normalized


def _selection_record(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _selection_text(*values: Any) -> str:
    for value in values:
        if isinstance(value, str) and value.strip():
            return value.strip()
    return ""


def _selection_datetime(value: Any) -> datetime | None:
    text = _selection_text(value)
    if not text:
        return None
    try:
        if len(text) == 10:
            return datetime.fromisoformat(text).replace(tzinfo=UTC)
        parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
        return parsed if parsed.tzinfo else parsed.replace(tzinfo=UTC)
    except ValueError:
        return None


def _selection_form_data(record: SelectionForm) -> dict[str, Any]:
    return copy.deepcopy(record.form_data) if isinstance(record.form_data, dict) else {}


def _clean_selection_string(value: Any) -> str | None:
    if isinstance(value, str) and value.strip():
        return value.strip()
    return None


def selection_form_document_metadata(raw_entry: Any) -> dict[str, str | None] | None:
    if isinstance(raw_entry, str):
        file_name = _clean_selection_string(raw_entry)
        if file_name:
            return {"file_name": file_name, "file_url": None, "document_id": None}
    if isinstance(raw_entry, dict):
        file_name = (
            _clean_selection_string(raw_entry.get("fileName"))
            or _clean_selection_string(raw_entry.get("file_name"))
            or _clean_selection_string(raw_entry.get("name"))
        )
        file_url = _clean_selection_string(raw_entry.get("fileUrl")) or _clean_selection_string(raw_entry.get("file_url"))
        document_id = _clean_selection_string(raw_entry.get("documentId")) or _clean_selection_string(raw_entry.get("document_id"))
        if file_name or file_url or document_id:
            return {"file_name": file_name, "file_url": file_url, "document_id": document_id}
    return None


def selection_form_document_entry(record: SelectionForm, document_key: str) -> dict[str, str | None]:
    form_data = record.form_data if isinstance(record.form_data, dict) else {}
    documents = form_data.get("documentsUploaded")
    if not isinstance(documents, dict) or document_key not in documents:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Selection form document not found")

    metadata = selection_form_document_metadata(documents.get(document_key))
    if metadata:
        return metadata
    raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Selection form document not found")


def lookup_selection_form_document(
    db: Session,
    *,
    candidate_id: str,
    document_key: str,
    metadata: dict[str, str | None],
) -> Document | None:
    document_id = metadata.get("document_id")
    if document_id:
        document = db.get(Document, document_id)
        if document and document.candidate_id == candidate_id:
            return document

    documents = list(
        db.scalars(
            select(Document)
            .where(Document.candidate_id == candidate_id)
            .order_by(Document.created_at.desc())
        )
    )

    file_url = metadata.get("file_url")
    if file_url:
        for document in documents:
            if document.file_url == file_url:
                return document

    file_name = metadata.get("file_name")
    if file_name:
        named_matches = [document for document in documents if document.file_name == file_name]
        preferred_types = {
            document_key,
            f"selection_form_{document_key}",
            f"selection-form-{document_key}",
            f"selection_{document_key}",
        }
        for document in named_matches:
            if document.type in preferred_types:
                return document
        if named_matches:
            return named_matches[0]

    return None


def find_selection_form_document(
    db: Session,
    *,
    candidate_id: str,
    document_key: str,
    metadata: dict[str, str | None],
) -> Document:
    document = lookup_selection_form_document(
        db,
        candidate_id=candidate_id,
        document_key=document_key,
        metadata=metadata,
    )
    if document is not None:
        return document
    raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Uploaded document file not found")


def _is_relative_to(path: Path, parent: Path) -> bool:
    try:
        path.relative_to(parent)
        return True
    except ValueError:
        return False


def resolve_local_upload_path(file_url: str) -> Path | None:
    if not file_url.startswith("/uploads/"):
        return None

    relative_path = Path(file_url.removeprefix("/uploads/"))
    if relative_path.is_absolute() or ".." in relative_path.parts:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found on server")

    settings = get_settings()
    bases = [settings.local_storage_path, Path.cwd() / "uploads"]
    seen: set[Path] = set()
    for base in bases:
        resolved_base = base.resolve()
        if resolved_base in seen:
            continue
        seen.add(resolved_base)
        candidate = (resolved_base / relative_path).resolve()
        if _is_relative_to(candidate, resolved_base) and candidate.exists():
            return candidate
    return None


def selection_form_document_media_type(document: Document) -> str:
    value = (document.mime_type or "").split(";", maxsplit=1)[0].strip().lower()
    if value:
        return value
    guessed, _ = mimetypes.guess_type(document.file_name or document.file_url or "")
    return guessed or "application/octet-stream"


def selection_form_document_file_available(document: Document) -> bool:
    file_url = document.file_url or ""
    if file_url.startswith("/uploads/"):
        return resolve_local_upload_path(file_url) is not None
    return bool(file_url)


def _read_selection_form_document_content(document: Document) -> bytes | None:
    file_url = document.file_url or ""
    local_path = resolve_local_upload_path(file_url)
    if local_path and local_path.exists():
        return local_path.read_bytes()

    download_url = StorageService().presigned_download_url(file_url)
    if not download_url:
        return None
    try:
        with urlopen(download_url, timeout=30) as response:  # nosec B310 - generated S3 URL for stored upload.
            return response.read(15 * 1024 * 1024)
    except OSError:
        return None


def _selection_form_document_verification_update(result: dict[str, Any]) -> dict[str, Any]:
    ocr_status = str(result.get("ocrStatus") or "skipped")
    matched = result.get("matchesExpectedCategory")
    verification_status = (
        "verified"
        if matched is True or ocr_status == "extracted"
        else "needs_review"
        if matched is False or ocr_status == "needs_review"
        else ocr_status
    )
    return {
        "verificationStatus": verification_status,
        "ocrStatus": ocr_status,
        "needsReview": verification_status == "needs_review",
        "detectedDocumentType": result.get("detectedDocumentType"),
        "matchesExpectedCategory": matched,
        "verificationMessage": result.get("message") or "",
        "verifiedAt": datetime.now(UTC).isoformat(),
    }


def persist_selection_form_document_verification(
    db: Session,
    *,
    record: SelectionForm,
    document_key: str,
    actor: User,
) -> dict[str, Any]:
    from app.services import employees as employee_service

    metadata = selection_form_document_entry(record, document_key)
    document = find_selection_form_document(
        db,
        candidate_id=record.candidate_id,
        document_key=document_key,
        metadata=metadata,
    )
    content = _read_selection_form_document_content(document)
    result = employee_service.verify_document_content(
        content=content,
        mime_type=selection_form_document_media_type(document),
        document_type=f"selection_form_{document_key}",
    )
    verification = _selection_form_document_verification_update(result)
    if verification["verificationStatus"] == "verified":
        document.status = "verified"
        document.ocr_status = "extracted"
        document.verified_by = actor.id
        document.verified_at = datetime.now(UTC)
    elif verification["verificationStatus"] == "needs_review":
        document.status = "needs_review"
        document.ocr_status = "needs_review"
    document.extracted_data = {
        "detectedDocumentType": result.get("detectedDocumentType"),
        "matchesExpectedCategory": result.get("matchesExpectedCategory"),
        "message": result.get("message") or "",
        "verifiedAt": verification["verifiedAt"],
    }
    db.add(document)

    form_data = _selection_form_data(record)
    documents = form_data.get("documentsUploaded")
    if not isinstance(documents, dict):
        documents = {}
    current_entry = documents.get(document_key)
    current_payload = copy.deepcopy(current_entry) if isinstance(current_entry, dict) else {}
    current_metadata = selection_form_document_metadata(current_entry) or {}
    documents[document_key] = {
        **current_payload,
        "fileName": document.file_name or current_metadata.get("file_name") or metadata.get("file_name"),
        "documentId": document.id,
        "fileUrl": document.file_url,
        "mimeType": document.mime_type,
        **verification,
        "verifiedBy": actor.id,
    }
    form_data["documentsUploaded"] = documents
    record.form_data = form_data
    db.add(record)
    return {"result": result, "verification": verification, "document": document}


def selection_form_required_document_keys(form_data: dict[str, Any]) -> set[str]:
    basic = form_data.get("basicDetails")
    basic_details = basic if isinstance(basic, dict) else {}
    experience_type = str(basic_details.get("experienceType") or "").strip().lower()
    bank = form_data.get("bankDetails")
    bank_details = bank if isinstance(bank, dict) else {}
    required = {
        "passport_size_photo",
        "marksheet_10th",
        "marksheet_12th",
        "graduation",
        "pan_doc",
        "aadhaar_doc",
        "permanent_address_proof",
    }
    if experience_type and experience_type not in {"fresher", "freshers", "entry_level"}:
        required.update({"experience_letter_1", "relieving_letter", "payslips"})
    # NOTE: cancelled_cheque is intentionally NOT required here. Bank details (and
    # the cancelled cheque) are OPTIONAL on the candidate selection form — the UI
    # explicitly lets candidates submit without them, and HR collects/validates
    # bank info separately via the Penny Drop flow. Requiring the cheque for
    # auto-validation whenever any bank field was filled silently stranded every
    # such candidate at "needs_review" (e.g. a candidate who typed their account
    # number but, as the form permits, did not upload a cheque). If a cheque IS
    # uploaded it is still verified; it just no longer blocks auto-validation.
    _ = bank_details  # retained for clarity; bank docs are not gating.
    return required


def selection_form_document_queue_counts(record: SelectionForm) -> dict[str, int]:
    form_data = record.form_data if isinstance(record.form_data, dict) else {}
    documents = form_data.get("documentsUploaded")
    submitted_keys = set()
    if isinstance(documents, dict):
        submitted_keys = {key for key, value in documents.items() if selection_form_document_metadata(value)}
    required_keys = selection_form_required_document_keys(form_data)
    return {
        "required": len(required_keys),
        "submitted": len(submitted_keys),
        "missing": len(required_keys - submitted_keys),
        "total": len(submitted_keys),
    }


def selection_form_verification_queue_state(record: SelectionForm) -> dict[str, Any]:
    form_data = record.form_data if isinstance(record.form_data, dict) else {}
    raw_state = form_data.get(SELECTION_FORM_VERIFICATION_QUEUE_KEY)
    state = raw_state if isinstance(raw_state, dict) else {}
    status_value = str(state.get("status") or "").strip()
    message = str(state.get("message") or "").strip()
    if record.validated_at:
        status_value = "validated"
        message = "Document checks are complete and the form is validated."
    elif not record.submitted_at and not status_value:
        status_value = "not_started"
    elif record.submitted_at and not status_value:
        status_value = "submitted"
    if not message:
        if status_value in {"queued", "processing"}:
            message = "Document checks are in queue. Please wait while we verify the uploaded files."
        elif status_value == "validated":
            message = "Document checks are complete and the form is validated."
        elif status_value == "needs_review":
            message = "Document checks are complete. HR will review the items that need attention."
        elif status_value == "failed":
            message = "Document checks could not be completed automatically. HR will review the form."
        elif status_value == "submitted":
            message = "Selection form is submitted and awaiting HR review."
        else:
            message = ""
    counts = selection_form_document_queue_counts(record)
    return {
        "status": status_value,
        "message": message,
        "taskId": state.get("taskId"),
        "queuedAt": state.get("queuedAt"),
        "startedAt": state.get("startedAt"),
        "completedAt": state.get("completedAt"),
        **counts,
    }


def set_selection_form_verification_queue_state(
    record: SelectionForm,
    *,
    status_value: str,
    message: str,
    task_id: str | None = None,
) -> None:
    form_data = _selection_form_data(record)
    current = form_data.get(SELECTION_FORM_VERIFICATION_QUEUE_KEY)
    state = copy.deepcopy(current) if isinstance(current, dict) else {}
    now = datetime.now(UTC).isoformat()
    state.update({
        "status": status_value,
        "message": message,
    })
    if task_id:
        state["taskId"] = task_id
    if status_value == "queued":
        state["queuedAt"] = now
        state.pop("startedAt", None)
        state.pop("completedAt", None)
    elif status_value == "processing":
        state["startedAt"] = now
    elif status_value in {"validated", "needs_review", "failed", "skipped"}:
        state["completedAt"] = now
    state.update(selection_form_document_queue_counts(record))
    form_data[SELECTION_FORM_VERIFICATION_QUEUE_KEY] = state
    record.form_data = form_data


def clear_selection_form_verification_queue_state(record: SelectionForm) -> None:
    form_data = _selection_form_data(record)
    form_data.pop(SELECTION_FORM_VERIFICATION_QUEUE_KEY, None)
    record.form_data = form_data


def selection_form_has_documents_to_queue(record: SelectionForm) -> bool:
    form_data = record.form_data if isinstance(record.form_data, dict) else {}
    documents = form_data.get("documentsUploaded")
    return isinstance(documents, dict) and any(selection_form_document_metadata(value) for value in documents.values())


def process_selection_form_document_verification(
    db: Session,
    *,
    selection_form: SelectionForm,
    actor: User,
) -> dict[str, Any]:
    form_data = selection_form.form_data if isinstance(selection_form.form_data, dict) else {}
    documents = form_data.get("documentsUploaded")
    if not isinstance(documents, dict):
        set_selection_form_verification_queue_state(
            selection_form,
            status_value="skipped",
            message="No uploaded documents were found for automatic checks.",
        )
        db.add(selection_form)
        return {"status": "skipped", "verified": 0, "needsReview": 0, "failed": 0}

    required_keys = selection_form_required_document_keys(form_data)
    submitted_keys = {key for key, value in documents.items() if selection_form_document_metadata(value)}
    missing_keys = sorted(required_keys - submitted_keys)
    all_ready = not missing_keys
    summary = {"verified": 0, "needsReview": 0, "failed": 0, "skipped": 0}
    results: list[dict[str, Any]] = []

    for document_key in sorted(submitted_keys):
        try:
            outcome = persist_selection_form_document_verification(
                db,
                record=selection_form,
                document_key=document_key,
                actor=actor,
            )
        except Exception:
            all_ready = False
            summary["failed"] += 1
            results.append({"documentKey": document_key, "status": "failed"})
            continue

        verification = outcome["verification"]
        verification_status = verification.get("verificationStatus")
        if verification_status == "verified" and verification.get("matchesExpectedCategory") is True:
            summary["verified"] += 1
        elif verification_status == "skipped":
            # 'skipped' = no AI verifier exists for this document type (experience
            # letters, relieving letters, payslips, certifications, etc.). The file
            # IS present (presence is separately enforced via missing_keys above),
            # so the mere absence of an AI category check must NOT block
            # auto-validation — otherwise every experienced candidate (whose
            # required experience docs are all non-AI-verifiable) is stuck forever.
            summary["skipped"] += 1
        else:
            all_ready = False
            summary["needsReview"] += 1
        results.append({"documentKey": document_key, "status": verification_status})

    if missing_keys:
        summary["failed"] += len(missing_keys)

    if all_ready:
        validate_selection_form(db, selection_form=selection_form, actor=actor)
        status_value = "validated"
        message = "Document checks are complete and the form is validated."
    else:
        status_value = "needs_review"
        message = "Document checks are complete. HR will review the form before validation."

    set_selection_form_verification_queue_state(
        selection_form,
        status_value=status_value,
        message=message,
    )
    db.add(selection_form)
    return {
        "status": status_value,
        "allReady": all_ready,
        "missingKeys": missing_keys,
        "results": results,
        **summary,
    }


def _sync_candidate_from_selection_form(candidate: Candidate, form_data: dict[str, Any]) -> None:
    basic = _selection_record(form_data.get("basicDetails"))
    personal = _selection_record(form_data.get("personalDetails"))
    identity = _selection_record(form_data.get("identityDetails"))

    full_name = _selection_text(basic.get("fullName"))
    if full_name:
        candidate.full_name = full_name
    email = _selection_text(basic.get("email"))
    if email:
        candidate.personal_email = email
    phone = re.sub(r"\D", "", _selection_text(basic.get("contactNumber")))
    if phone:
        candidate.phone = phone
    date_of_birth = _selection_datetime(basic.get("dateOfBirth"))
    if date_of_birth:
        candidate.date_of_birth = date_of_birth
    gender = _selection_text(personal.get("gender"))
    if gender:
        candidate.gender = gender
    marital_status = _selection_text(personal.get("maritalStatus"))
    if marital_status:
        candidate.marital_status = marital_status
    experience_type = _selection_text(basic.get("experienceType"))
    if experience_type:
        candidate.experience_type = experience_type

    aadhaar_number = re.sub(
        r"\D",
        "",
        _selection_text(identity.get("aadhaarNumber"), personal.get("aadhaarNumber")),
    )
    aadhaar_name = full_name or _selection_text(candidate.full_name)
    aadhaar_dob = _selection_text(basic.get("dateOfBirth"))
    if aadhaar_number or aadhaar_name or aadhaar_dob:
        aadhaar_extracted = dict(candidate.aadhaar_extracted or {})
        if aadhaar_number:
            aadhaar_extracted["aadhaarNumber"] = aadhaar_number
            if len(aadhaar_number) >= 4:
                candidate.aadhaar_last4 = aadhaar_number[-4:]
        if aadhaar_name:
            aadhaar_extracted["cardHolderName"] = aadhaar_name
            candidate.aadhaar_ocr_name = aadhaar_name
        if aadhaar_dob:
            aadhaar_extracted["dateOfBirth"] = aadhaar_dob
        aadhaar_extracted["ocrStatus"] = "selection_form"
        aadhaar_extracted["message"] = "Shown from submitted selection form."
        candidate.aadhaar_extracted = aadhaar_extracted
        if not candidate.aadhaar_validation_status:
            candidate.aadhaar_validation_status = "selection_form"


def _normalize_pi_round_decision(value: str | None) -> str | None:
    if value is None:
        return None
    normalized = value.strip().lower().replace(" ", "_")
    if not normalized:
        return None
    if normalized in {"next_round", "continue", "continued", "proceed", "proceed_to_next", "proceed_next_round"}:
        return "proceed_to_next_round"
    if normalized in {"selected", "pass", "passed"}:
        return "selected"
    if normalized in {"rejected", "reject", "failed", "fail"}:
        return "rejected"
    return normalized


def _normalize_pi_final_verdict(value: str | None) -> str | None:
    if value is None:
        return None
    normalized = value.strip().lower().replace(" ", "_")
    if normalized in {"pass", "passed", "selected"}:
        return "selected"
    if normalized in {"fail", "failed", "reject", "rejected"}:
        return "rejected"
    return normalized if normalized in PI_FINAL_VERDICTS else None


def serialize_pi_round(round_record: PiInterviewRound) -> dict[str, Any]:
    return {
        "id": round_record.id,
        "evaluationId": round_record.evaluation_id,
        "candidateId": round_record.candidate_id,
        "evaluatorId": round_record.evaluator_id,
        "roundNumber": round_record.round_number,
        "panelLabel": round_record.panel_label,
        "subject": round_record.subject,
        "scheduledAt": round_record.scheduled_at,
        "completedAt": round_record.completed_at,
        "status": round_record.status,
        "mode": round_record.mode,
        "durationMinutes": round_record.duration_minutes,
        "score": round_record.score,
        "remarks": round_record.remarks,
        "notes": round_record.remarks,
        "roundDecision": round_record.round_decision,
        "noFurtherPiRequired": round_record.no_further_pi_required,
        "finalVerdict": round_record.final_verdict,
        "panelMembers": round_record.panel_members or [],
        "evaluatorName": round_record.evaluator.name if round_record.evaluator else None,
        "createdAt": round_record.created_at,
        "updatedAt": round_record.updated_at,
    }


def get_latest_pi_round(evaluation: Evaluation) -> PiInterviewRound | None:
    rounds = sorted(evaluation.pi_rounds or [], key=lambda item: (item.round_number, item.updated_at or item.created_at or datetime.min.replace(tzinfo=UTC)))
    return rounds[-1] if rounds else None


def ensure_legacy_pi_rounds(db: Session, *, evaluation: Evaluation) -> list[PiInterviewRound]:
    rounds = sorted(evaluation.pi_rounds or [], key=lambda item: item.round_number)
    if rounds:
        return rounds

    has_legacy_pi_state = any([
        bool(evaluation.interview_subject),
        evaluation.interview_scheduled_at is not None,
        bool(evaluation.interview_status),
        bool(evaluation.interview_notes),
        evaluation.pi_score is not None,
    ])
    if not has_legacy_pi_state:
        return []

    candidate = evaluation.candidate or db.get(Candidate, evaluation.candidate_id)
    recommendation = (evaluation.recommendation or "").strip().lower()
    status_value = (evaluation.interview_status or "").strip().lower()
    inferred_final_verdict = _normalize_pi_final_verdict(recommendation)
    if inferred_final_verdict is None and candidate is not None:
        if candidate.current_stage == CandidateStage.EVALUATION_FAILED:
            inferred_final_verdict = "rejected"
        elif candidate.current_stage == CandidateStage.EVALUATION_PASSED and status_value == "completed":
            inferred_final_verdict = "selected"

    round_status = status_value or ("completed" if inferred_final_verdict else "scheduled")
    if round_status not in {"scheduled", "rescheduled", "completed", "cancelled", "no_further_pi_required"}:
        round_status = "completed" if inferred_final_verdict else "scheduled"
    if inferred_final_verdict:
        round_status = "no_further_pi_required" if round_status in {"completed", "no_further_pi_required"} else round_status

    legacy_round = PiInterviewRound(
        evaluation_id=evaluation.id,
        candidate_id=evaluation.candidate_id,
        evaluator_id=evaluation.evaluator_id,
        round_number=1,
        subject=evaluation.interview_subject,
        scheduled_at=evaluation.interview_scheduled_at,
        completed_at=evaluation.completed_at if round_status in {"completed", "no_further_pi_required"} else None,
        status=round_status,
        mode=evaluation.interview_mode,
        duration_minutes=60,
        score=evaluation.pi_score,
        remarks=evaluation.interview_notes,
        round_decision=inferred_final_verdict or None,
        no_further_pi_required=bool(inferred_final_verdict),
        final_verdict=inferred_final_verdict,
    )
    db.add(legacy_round)
    db.flush()
    rounds = [legacy_round]
    evaluation.pi_rounds = rounds
    return rounds


def sync_evaluation_interview_summary(
    evaluation: Evaluation,
    *,
    latest_round: PiInterviewRound | None = None,
) -> None:
    latest_round = latest_round or get_latest_pi_round(evaluation)
    if latest_round is None:
        return
    evaluation.interview_subject = latest_round.subject
    evaluation.interview_scheduled_at = latest_round.scheduled_at
    evaluation.interview_status = latest_round.status
    evaluation.interview_notes = latest_round.remarks
    evaluation.interview_mode = latest_round.mode
    evaluation.pi_score = latest_round.score
    if latest_round.final_verdict:
        evaluation.recommendation = latest_round.final_verdict


def stage_to_status(stage: CandidateStage) -> str:
    labels = {
        CandidateStage.NEW_APPLICATION: "New Application",
        CandidateStage.SOURCE_TAGGED: "Source Tagged",
        CandidateStage.RESUME_UPLOADED: "Resume Uploaded",
        CandidateStage.RESUME_SCREENING_PENDING: "Screening Pending",
        CandidateStage.RESUME_SHORTLISTED: "Resume Shortlisted",
        CandidateStage.RESUME_REJECTED: "Resume Rejected",
        CandidateStage.EVALUATION_ASSIGNED: "Evaluation Assigned",
        CandidateStage.EVALUATION_IN_PROGRESS: "Evaluation In Progress",
        CandidateStage.EVALUATION_PASSED: "Evaluation Passed",
        CandidateStage.EVALUATION_FAILED: "Evaluation Failed",
        CandidateStage.SELECTION_FORM_SENT: "Selection Form Sent",
        CandidateStage.SELECTION_FORM_SUBMITTED: "Selection Form Submitted",
        CandidateStage.SELECTION_FORM_VALIDATED: "Selection Form Validated",
        CandidateStage.CONTRACT_SENT: "Contract Sent",
        CandidateStage.CONTRACT_SIGNED: "Contract Signed",
        CandidateStage.INDUCTION_COMPLETED: "Induction Completed",
        CandidateStage.IT_EMAIL_CREATED: "IT Email Created",
        CandidateStage.WELCOME_MAIL_SENT: "Welcome Mail Sent",
        CandidateStage.STATUTORY_FORMS_SENT: "Compliance Forms Sent",
        CandidateStage.STATUTORY_FORMS_SUBMITTED: "Compliance Forms Submitted",
        CandidateStage.COMPLIANCE_VERIFIED: "Compliance Verified",
        CandidateStage.ONBOARDING_COMPLETED: "Onboarding Completed",
    }
    return labels[stage]


# Pipeline buckets — the coarse, ordered phases of the candidate journey (Applied, Screening,
# Evaluation, Selection-form, Selection-form-validation, Contract, Onboarding). The granular
# sub-stages within a bucket are routinely skipped by automated steps, so ordering/gating works
# at the bucket level. This lives in workflows.py (the lowest-level workflow module) as the
# single source of truth; candidates.py imports stage_bucket() from here.
STAGE_BUCKETS: tuple[tuple[CandidateStage, ...], ...] = (
    (CandidateStage.NEW_APPLICATION, CandidateStage.SOURCE_TAGGED, CandidateStage.RESUME_UPLOADED),
    (CandidateStage.RESUME_SCREENING_PENDING, CandidateStage.RESUME_SHORTLISTED, CandidateStage.RESUME_REJECTED),
    (CandidateStage.EVALUATION_ASSIGNED, CandidateStage.EVALUATION_IN_PROGRESS, CandidateStage.EVALUATION_PASSED, CandidateStage.EVALUATION_FAILED),
    (CandidateStage.SELECTION_FORM_SENT, CandidateStage.SELECTION_FORM_SUBMITTED),
    (CandidateStage.SELECTION_FORM_VALIDATED,),
    (CandidateStage.CONTRACT_SENT, CandidateStage.CONTRACT_SIGNED),
    (
        CandidateStage.INDUCTION_COMPLETED, CandidateStage.IT_EMAIL_CREATED, CandidateStage.WELCOME_MAIL_SENT,
        CandidateStage.STATUTORY_FORMS_SENT, CandidateStage.STATUTORY_FORMS_SUBMITTED,
        CandidateStage.COMPLIANCE_VERIFIED, CandidateStage.ONBOARDING_COMPLETED,
    ),
)


def stage_bucket(stage: CandidateStage) -> int | None:
    for i, members in enumerate(STAGE_BUCKETS):
        if stage in members:
            return i
    return None


def create_notification(
    db: Session,
    *,
    user_id: str,
    title: str,
    message: str,
    type_: NotificationType = NotificationType.INFO,
    candidate_id: str | None = None,
    entity_type: str | None = None,
    entity_id: str | None = None,
    payload: dict[str, Any] | None = None,
) -> Notification:
    notification = Notification(
        user_id=user_id,
        candidate_id=candidate_id,
        title=title,
        message=message,
        type=type_,
        entity_type=entity_type,
        entity_id=entity_id,
        payload=payload,
    )
    db.add(notification)
    db.flush()
    return notification


def notify_roles(
    db: Session,
    *,
    roles: set[Role],
    title: str,
    message: str,
    type_: NotificationType = NotificationType.INFO,
    candidate_id: str | None = None,
) -> None:
    recipients = [
        user
        for user in db.scalars(select(User).where(User.is_active.is_(True))).all()
        if _has_any_role(user, roles)
    ]
    for recipient in recipients:
        create_notification(
            db,
            user_id=recipient.id,
            candidate_id=candidate_id,
            title=title,
            message=message,
            type_=type_,
        )


def resolve_notification_route(notification: Notification, *, user: User) -> str | None:
    text = f"{notification.title} {notification.message}".strip().lower()

    if user.role == Role.CANDIDATE:
        if "selection form" in text:
            return "/portal/selection-form"
        if "contract" in text or "offer" in text or "employment agreement" in text:
            return "/portal/contract"
        if "compliance" in text or "statutory" in text or "epf" in text or "esic" in text:
            return "/portal/compliance"
        if "document" in text or "upload" in text or "verification" in text or "aadhaar" in text or "pan" in text:
            return "/portal/documents"
        if "assessment" in text:
            return "/portal/my-assessments"
        return "/portal/dashboard"

    if user.role in {Role.EMPLOYEE, Role.EMPLOYEE_REFERRER}:
        if "leave" in text:
            return "/dashboard/employee/leave"
        if "resignation" in text or "last working day" in text or "employment terminated" in text:
            return "/dashboard/employee/separation"
        if "selection form" in text:
            return "/dashboard/employee/selection-form"
        if "contract" in text or "employment agreement" in text:
            return "/dashboard/employee/contracts"
        if "compliance" in text:
            return "/dashboard/employee/compliance"
        if "document" in text or "aadhaar" in text or "resume" in text or "upload" in text:
            return "/dashboard/employee/documents"
        if "referral" in text:
            return "/dashboard/employee/referrals"
        return "/dashboard/employee"

    if user.role == Role.MANAGER:
        if "leave" in text:
            return "/dashboard/manager/leaves"
        if "resignation" in text or "last working day" in text:
            return "/dashboard/separation"
        if notification.candidate_id:
            return f"/dashboard/candidates/{notification.candidate_id}"
        return "/dashboard/manager"

    if user.role == Role.IT_TEAM:
        if "id card" in text and notification.candidate_id:
            return f"/dashboard/candidates/{notification.candidate_id}"
        if "asset" in text or "laptop" in text:
            return "/dashboard/it/assets"
        if "email" in text or "offboarding" in text or "deactivate" in text:
            return "/dashboard/it"
        return "/dashboard/it"

    if user.role == Role.OFFICE_ADMIN:
        if "id card" in text or "offboarding" in text or "termination" in text:
            return "/dashboard/office-admin"
        return "/dashboard/employees"

    if "leave" in text:
        return "/dashboard/leave"
    if "resignation" in text or "last working day" in text or "terminated" in text or "offboarding" in text:
        return "/dashboard/separation"
    if "escalat" in text or "sla breach" in text:
        return "/dashboard/escalations"
    if "assessment" in text:
        return "/dashboard/assessment-platform"
    if "evaluat" in text or "interview" in text:
        if user.role == Role.EVALUATOR:
            return "/dashboard/evaluations"
        if notification.candidate_id:
            return f"/dashboard/candidates/{notification.candidate_id}"
        return "/dashboard/evaluations"
    if "selection form" in text:
        if "employee" in text:
            return "/dashboard/employees"
        if notification.candidate_id:
            return f"/dashboard/candidates/{notification.candidate_id}"
        return "/dashboard/selection-forms"
    if "contract" in text or "employment agreement" in text:
        return "/dashboard/contracts"
    if "compliance" in text or "statutory" in text or "epf" in text or "esic" in text:
        if "employee" in text:
            return "/dashboard/employees"
        return "/dashboard/compliance"
    if "document" in text or "verification" in text or "resume" in text or "aadhaar" in text:
        if "employee" in text:
            return "/dashboard/employees"
        if notification.candidate_id:
            return f"/dashboard/candidates/{notification.candidate_id}"
        return "/dashboard/documents"
    if "id card" in text and notification.candidate_id:
        return f"/dashboard/candidates/{notification.candidate_id}"
    if "referral" in text:
        return "/dashboard/candidates"
    if notification.candidate_id:
        return f"/dashboard/candidates/{notification.candidate_id}"
    return None


def serialize_notification(notification: Notification, *, user: User) -> dict[str, Any]:
    return {
        "id": notification.id,
        "userId": notification.user_id,
        "candidateId": notification.candidate_id,
        "title": notification.title,
        "message": notification.message,
        "type": notification.type,
        "isRead": notification.is_read,
        "entityType": notification.entity_type,
        "entityId": notification.entity_id,
        "payload": notification.payload,
        "createdAt": notification.created_at,
        "candidateName": notification.candidate.full_name if notification.candidate else None,
        "route": resolve_notification_route(notification, user=user),
    }


def suggest_ethara_email(candidate: Candidate, *, db: Session) -> str:
    base = candidate.personal_email.split("@")[0].replace(" ", ".").lower()
    base = "".join(char for char in base if char.isalnum() or char == ".").strip(".") or "candidate"
    domain = "ethara.ai"
    suggestion = f"{base}@{domain}"
    counter = 1
    while db.scalar(select(func.count()).select_from(Candidate).where(Candidate.ethara_email == suggestion)):
        counter += 1
        suggestion = f"{base}{counter}@{domain}"
    return suggestion


def ensure_selection_form(db: Session, candidate_id: str) -> SelectionForm:
    record = db.scalar(select(SelectionForm).where(SelectionForm.candidate_id == candidate_id))
    if record:
        return record
    record = SelectionForm(candidate_id=candidate_id, sent_at=datetime.now(UTC))
    db.add(record)
    db.flush()
    return record


def ensure_contract(db: Session, candidate_id: str) -> Contract:
    record = db.scalar(select(Contract).where(Contract.candidate_id == candidate_id))
    if record:
        return record
    record = Contract(candidate_id=candidate_id, status=ContractStatus.DRAFT)
    db.add(record)
    db.flush()
    return record


def ensure_it_request(db: Session, candidate: Candidate, *, requested_by: str, assigned_to_id: str | None = None) -> ITRequest:
    record = db.scalar(select(ITRequest).where(ITRequest.candidate_id == candidate.id))
    if record:
        return record
    record = ITRequest(
        candidate_id=candidate.id,
        requested_by=requested_by,
        assigned_to_id=assigned_to_id,
        suggested_email=suggest_ethara_email(candidate, db=db),
        status="pending",
    )
    db.add(record)
    db.flush()
    return record


def ensure_compliance_forms(db: Session, candidate_id: str) -> list[ComplianceForm]:
    """Statutory/compliance forms are now Documenso e-sign forms (Form 11 / Form 2 / Form F),
    sent to the candidate for signature. Idempotent (skips forms already created)."""
    candidate = db.get(Candidate, candidate_id)
    if candidate is None:
        return []
    try:
        from app.services import compliance_documenso as compliance_esign

        return compliance_esign.send_candidate_compliance_forms(db, candidate=candidate)
    except Exception:
        import logging as _logging

        _logging.getLogger(__name__).warning(
            "Compliance Documenso send failed for candidate %s", candidate_id, exc_info=True
        )
        return list(
            db.scalars(select(ComplianceForm).where(ComplianceForm.candidate_id == candidate_id))
        )


def list_compliance_forms(db: Session, candidate_id: str) -> list[ComplianceForm]:
    return list(db.scalars(select(ComplianceForm).where(ComplianceForm.candidate_id == candidate_id)))


def apply_stage_side_effects(db: Session, candidate: Candidate, *, actor: User | None = None) -> None:
    if candidate.current_stage == CandidateStage.RESUME_SHORTLISTED:
        # Auto-assign to the default evaluator so the candidate appears in the
        # Evaluation queue immediately without a manual assignment step.
        if actor is not None:
            try:
                ensure_evaluation_assignment(db, candidate_id=candidate.id, actor=actor)
            except Exception:
                pass  # Non-fatal — evaluator can be assigned manually
    elif candidate.current_stage == CandidateStage.SELECTION_FORM_SENT:
        ensure_selection_form(db, candidate.id)
    elif candidate.current_stage == CandidateStage.CONTRACT_SENT:
        contract = ensure_contract(db, candidate.id)
        contract.status = ContractStatus.SENT
        contract.sent_at = datetime.now(UTC)
        db.add(contract)
    elif candidate.current_stage == CandidateStage.CONTRACT_SIGNED:
        # Auto employee-code allocation is gated off by default: GRP codes + Ethara
        # IDs are issued only via the IT-dashboard bulk-register upload, not
        # automatically on contract signing. See AUTO_EMPLOYEE_PROVISIONING.
        if get_settings().auto_employee_provisioning:
            from app.services.employees import assign_candidate_employee_code

            assign_candidate_employee_code(db, candidate)
    elif candidate.current_stage == CandidateStage.IT_EMAIL_CREATED:
        ensure_it_request(db, candidate, requested_by=actor.id if actor else "system")
    elif candidate.current_stage == CandidateStage.STATUTORY_FORMS_SENT:
        ensure_compliance_forms(db, candidate.id)
    elif candidate.current_stage == CandidateStage.ONBOARDING_COMPLETED:
        # Final stage: optionally convert the candidate into an employee with a
        # separate login. Gated off by default — the candidate stays at
        # ONBOARDING_COMPLETED and HR creates the employee via the IT bulk-register
        # upload. Non-fatal if it fails. See AUTO_EMPLOYEE_PROVISIONING.
        if get_settings().auto_employee_provisioning:
            try:
                from app.services.employees import convert_candidate_to_employee

                convert_candidate_to_employee(db, candidate=candidate, actor=actor)
            except Exception:
                import logging as _logging

                _logging.getLogger(__name__).warning(
                    "Candidate->employee conversion failed for %s", candidate.id, exc_info=True
                )


def list_notifications(db: Session, *, user: User) -> list[Notification]:
    return list(
        db.scalars(
            select(Notification)
            .options(joinedload(Notification.candidate))
            .where(Notification.user_id == user.id)
            .order_by(Notification.created_at.desc())
        )
    )


def list_escalations(db: Session, *, status_filter: str | None = None) -> list[Escalation]:
    query = (
        select(Escalation)
        .options(joinedload(Escalation.candidate), joinedload(Escalation.responsible_user))
        .order_by(Escalation.created_at.desc())
    )
    if status_filter:
        query = query.where(Escalation.status == status_filter)
    return list(db.scalars(query))


def list_it_requests(db: Session, *, status_filter: str | None = None) -> list[ITRequest]:
    query = (
        select(ITRequest)
        .join(Candidate, ITRequest.candidate_id == Candidate.id)
        .where(Candidate.is_removed.is_(False))
        .options(joinedload(ITRequest.candidate))
        .order_by(ITRequest.created_at.desc())
    )
    if status_filter:
        query = query.where(ITRequest.status == status_filter)
    return list(db.scalars(query))


def list_documents(db: Session, *, candidate_id: str) -> list[Document]:
    return list(
        db.scalars(
            select(Document)
            .where(Document.candidate_id == candidate_id)
            .order_by(Document.created_at.desc())
        )
    )


def _extract_resume_text(content: bytes, *, file_name: str, content_type: str | None) -> str:
    suffix = Path(file_name).suffix.lower()
    mime_type = (content_type or "").split(";")[0].lower()

    try:
        from app.api.routes.candidates import _extract_resume_text as extract_candidate_resume_text

        text = extract_candidate_resume_text(content, suffix, mime_type)
        if text.strip():
            return text
    except Exception:
        pass

    if suffix == ".pdf" or mime_type == "application/pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if text.strip():
                return text
        except Exception:
            pass
        return ""

    if suffix in {".doc", ".docx"} or "word" in mime_type:
        try:
            import docx
            doc = docx.Document(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
            if text.strip():
                return text
        except Exception:
            pass
        return ""

    return content.decode("utf-8", errors="ignore")


def _extract_resume_details(
    *,
    file: UploadFile,
    candidate: Candidate,
) -> dict[str, Any]:
    content = file.file.read()
    file.file.seek(0)
    if not content:
        return {
            "resumeText": "",
            "summary": None,
            "keyPoints": [],
            "skills": [],
            "totalExperienceYears": 0,
            "currentRole": None,
            "education": "",
        }

    resume_text = _extract_resume_text(
        content,
        file_name=file.filename or "resume",
        content_type=file.content_type,
    )
    if not resume_text.strip():
        return {
            "resumeText": "",
            "summary": None,
            "keyPoints": [],
            "skills": [],
            "totalExperienceYears": 0,
            "currentRole": None,
            "education": "",
        }

    parsed = LLMService().parse_resume(
        resume_text=resume_text,
        job_title=candidate.position.title if candidate.position else None,
    )
    parsed["resumeText"] = resume_text[:20000]
    return parsed


def _normalize_screening_recommendation(value: str | None) -> str | None:
    if value is None:
        return None
    normalized = value.strip().lower()
    mapping = {
        "shortlist": "shortlisted",
        "shortlisted": "shortlisted",
        "approve": "shortlisted",
        "approved": "shortlisted",
        "reject": "rejected",
        "rejected": "rejected",
        "needs_review": "needs_review",
        "needs review": "needs_review",
        "review": "needs_review",
        "maybe": "needs_review",
        "pending": "pending",
    }
    return mapping.get(normalized, normalized)


def _candidate_stage_for_recommendation(recommendation: str | None) -> CandidateStage:
    normalized = _normalize_screening_recommendation(recommendation)
    if normalized == "shortlisted":
        return CandidateStage.RESUME_SHORTLISTED
    if normalized == "rejected":
        return CandidateStage.RESUME_REJECTED
    return CandidateStage.RESUME_SCREENING_PENDING


def _add_stage_log(
    db: Session,
    *,
    candidate: Candidate,
    from_stage: CandidateStage,
    to_stage: CandidateStage,
    actor: User | None,
    notes: str | None = None,
) -> None:
    if from_stage == to_stage:
        return
    db.add(
        StageLog(
            candidate_id=candidate.id,
            from_stage=from_stage,
            to_stage=to_stage,
            changed_by=actor.id if actor else "system",
            changed_by_name=actor.name if actor else "System",
            notes=notes,
        )
    )


def _latest_resume_document(candidate: Candidate) -> Document | None:
    resumes = [document for document in candidate.documents if document.type == "resume"]
    if not resumes:
        return None
    resumes.sort(key=lambda document: document.created_at or datetime.min.replace(tzinfo=UTC), reverse=True)
    return resumes[0]


def _resume_document_payload(document: Document | None) -> dict[str, Any] | None:
    if document is None:
        return None
    return {
        "id": document.id,
        "type": document.type,
        "fileName": document.file_name,
        "mimeType": document.mime_type,
        "status": document.status,
        "uploadedAt": document.created_at,
    }


def _screening_status_for_candidate(candidate: Candidate) -> str:
    if candidate.llm_status == "processing":
        return "running"
    if candidate.llm_status == "completed":
        return "completed"
    if candidate.current_stage in {
        CandidateStage.RESUME_SHORTLISTED,
        CandidateStage.RESUME_REJECTED,
    }:
        return "completed"
    payload = candidate.screening_payload or {}
    recommendation = _normalize_screening_recommendation(payload.get("recommendation"))
    if payload.get("lastScreenedAt") and recommendation not in {None, "pending"}:
        return "completed"
    if candidate.resume_url:
        return "pending"
    return "missing_resume"


def _resume_parsed_details(candidate: Candidate) -> dict[str, Any] | None:
    if not any([candidate.resume_summary, candidate.resume_key_points, candidate.resume_text]):
        return None
    payload = {
        "summary": candidate.resume_summary,
        "keyPoints": candidate.resume_key_points or [],
        "resumeText": candidate.resume_text,
    }
    screening_payload = candidate.screening_payload or {}
    for key in ("skills", "totalExperienceYears", "currentRole", "education"):
        if screening_payload.get(key) is not None:
            payload[key] = screening_payload.get(key)
    return payload


def build_screening_record(candidate: Candidate) -> dict[str, Any]:
    screening_payload = candidate.screening_payload or {}
    manual_override = screening_payload.get("manualOverride")
    last_screened_at = screening_payload.get("lastScreenedAt")
    if isinstance(last_screened_at, str):
        try:
            last_screened_at = datetime.fromisoformat(last_screened_at.replace("Z", "+00:00"))
        except ValueError:
            last_screened_at = None

    resume_document = _latest_resume_document(candidate)
    recommendation = _normalize_screening_recommendation(
        screening_payload.get("recommendation")
        or ("shortlisted" if candidate.current_stage == CandidateStage.RESUME_SHORTLISTED else None)
        or ("rejected" if candidate.current_stage == CandidateStage.RESUME_REJECTED else None)
    )

    return {
        "candidateId": candidate.id,
        "candidateCode": candidate.candidate_code,
        "candidateName": candidate.full_name,
        "personalEmail": candidate.personal_email,
        "phone": candidate.phone,
        "positionId": candidate.position_id,
        "positionTitle": candidate.position.title if candidate.position else None,
        "currentStage": candidate.current_stage.value,
        "currentStatus": candidate.current_status,
        "screeningStatus": _screening_status_for_candidate(candidate),
        "llmStatus": candidate.llm_status,
        "screeningScore": candidate.resume_score,
        "matchScore": screening_payload.get("matchScore", screening_payload.get("score", candidate.resume_score)),
        "recommendation": recommendation,
        "screeningSummary": candidate.resume_summary or screening_payload.get("summary"),
        "parsedResumeDetails": _resume_parsed_details(candidate),
        "screeningPayload": screening_payload or None,
        "manualOverride": manual_override if isinstance(manual_override, dict) else None,
        "resumeUploadedAt": resume_document.created_at if resume_document else None,
        "lastScreenedAt": last_screened_at,
        "updatedAt": candidate.updated_at,
        "createdAt": candidate.created_at,
        "resumeDocument": _resume_document_payload(resume_document),
    }


def list_screening_candidates(
    db: Session,
    *,
    search: str | None = None,
    recommendation: str | None = None,
    page: int = 1,
    limit: int = 25,
) -> tuple[list[dict[str, Any]], int]:
    query: Select[tuple[Candidate]] = (
        select(Candidate)
        .options(joinedload(Candidate.position), selectinload(Candidate.documents))
        .where(
            or_(
                Candidate.resume_url.is_not(None),
                Candidate.documents.any(Document.type == "resume"),
            )
        )
        .where(Candidate.is_removed.is_(False))
        .order_by(Candidate.updated_at.desc(), Candidate.created_at.desc())
    )
    count_query = (
        select(func.count())
        .select_from(Candidate)
        .where(
            or_(
                Candidate.resume_url.is_not(None),
                Candidate.documents.any(Document.type == "resume"),
            )
        )
        .where(Candidate.is_removed.is_(False))
    )
    if search:
        term = f"%{search.lower()}%"
        search_condition = (
            func.lower(Candidate.full_name).like(term)
            | func.lower(Candidate.personal_email).like(term)
            | func.lower(func.coalesce(Candidate.ethara_email, "")).like(term)
            | func.lower(Candidate.candidate_code).like(term)
        )
        query = query.where(search_condition)
        count_query = count_query.where(search_condition)

    normalized_recommendation = _normalize_screening_recommendation(recommendation)
    if normalized_recommendation:
        candidates = list(db.scalars(query).unique())
        records = [build_screening_record(candidate) for candidate in candidates]
        records = [
            record
            for record in records
            if _normalize_screening_recommendation(record.get("recommendation")) == normalized_recommendation
        ]
        total = len(records)
        start = max(page - 1, 0) * limit
        end = start + limit
        return records[start:end], total

    total = db.scalar(count_query) or 0
    candidates = list(db.scalars(query.offset((page - 1) * limit).limit(limit)).unique())
    return [build_screening_record(candidate) for candidate in candidates], total


def get_screening_candidate_or_404(db: Session, *, candidate_id: str) -> Candidate:
    candidate = db.scalar(
        select(Candidate)
        .options(joinedload(Candidate.position), selectinload(Candidate.documents))
        .where(Candidate.id == candidate_id)
    )
    if candidate is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Candidate not found")
    if not candidate.resume_url and _latest_resume_document(candidate) is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate does not have a resume submitted for screening.",
        )
    return candidate


def upload_document(
    db: Session,
    *,
    candidate: Candidate,
    file: UploadFile,
    type_: str,
    actor: User,
) -> Document:
    storage = StorageService()
    file_url, _ = storage.save_upload(file, folder=f"candidates/{candidate.id}")
    document = Document(
        candidate_id=candidate.id,
        type=type_,
        file_name=file.filename or "upload",
        file_url=file_url,
        file_size=getattr(file, "size", None),
        mime_type=file.content_type,
        status="pending",
    )
    db.add(document)
    db.flush()

    if type_ == "resume":
        parsed_resume = _extract_resume_details(file=file, candidate=candidate)
        previous_stage = candidate.current_stage
        candidate.resume_url = file_url
        candidate.resume_text = parsed_resume.get("resumeText") or candidate.resume_text
        candidate.resume_summary = parsed_resume.get("summary") or candidate.resume_summary
        candidate.resume_key_points = parsed_resume.get("keyPoints") or candidate.resume_key_points
        existing_payload = dict(candidate.screening_payload or {})
        existing_payload.update(
            {
                "status": "queued",
                "resumeDocumentId": document.id,
                "resumeUploadedAt": document.created_at.isoformat() if document.created_at else None,
                "summary": candidate.resume_summary,
                "keyPoints": candidate.resume_key_points or [],
            }
        )
        for key in ("skills", "totalExperienceYears", "currentRole", "education"):
            if parsed_resume.get(key) is not None:
                existing_payload[key] = parsed_resume.get(key)
        candidate.screening_payload = existing_payload
        candidate.llm_status = "processing"
        candidate.current_stage = CandidateStage.RESUME_SCREENING_PENDING
        candidate.current_status = stage_to_status(CandidateStage.RESUME_SCREENING_PENDING)
        _add_stage_log(
            db,
            candidate=candidate,
            from_stage=previous_stage,
            to_stage=CandidateStage.RESUME_SCREENING_PENDING,
            actor=actor,
            notes="Resume uploaded and queued for screening.",
        )
    elif type_ in {"aadhaar", "aadhaar_card", "aadhaar-card"}:
        # Aadhaar OCR is dispatched to a Celery task elsewhere, but prod has no
        # worker and that task only writes to the Document row — the candidate
        # dashboard reads candidate.aadhaar_extracted. Extract inline here (same
        # as the public registration flow) so the details actually show up.
        try:
            from app.api.routes.candidates import extract_aadhaar_fields
            file.file.seek(0)
            aadhaar_ocr = extract_aadhaar_fields(file)
        except Exception:
            aadhaar_ocr = None
        if aadhaar_ocr and (
            aadhaar_ocr.get("aadhaarNumber")
            or aadhaar_ocr.get("dateOfBirth")
            or aadhaar_ocr.get("cardHolderName")
        ):
            candidate.aadhaar_extracted = aadhaar_ocr
            ocr_name = aadhaar_ocr.get("cardHolderName") or aadhaar_ocr.get("name")
            if ocr_name and not candidate.aadhaar_ocr_name:
                candidate.aadhaar_ocr_name = ocr_name
            ocr_number = re.sub(r"\D", "", aadhaar_ocr.get("aadhaarNumber") or "")
            if len(ocr_number) >= 4 and not candidate.aadhaar_last4:
                candidate.aadhaar_last4 = ocr_number[-4:]
    db.add(candidate)
    log_audit(
        db,
        entity_type="document",
        entity_id=document.id,
        action=f"document_uploaded:{type_}",
        actor=actor,
        candidate_id=candidate.id,
        new_value={"type": type_, "fileUrl": file_url},
    )
    return document


def verify_document(db: Session, *, document: Document, status_value: str, actor: User) -> Document:
    document.status = status_value
    document.verified_by = actor.id
    document.verified_at = datetime.now(UTC)
    db.add(document)
    log_audit(
        db,
        entity_type="document",
        entity_id=document.id,
        action=f"document_verified:{status_value}",
        actor=actor,
        candidate_id=document.candidate_id,
        new_value={"status": status_value},
    )
    return document


def create_evaluation_assignment(db: Session, *, candidate_id: str, evaluator_id: str, actor: User) -> Evaluation:
    evaluation = Evaluation(candidate_id=candidate_id, evaluator_id=evaluator_id)
    db.add(evaluation)
    db.flush()
    create_notification(
        db,
        user_id=evaluator_id,
        candidate_id=candidate_id,
        title="Evaluation Assigned",
        message="A candidate has been assigned to you for evaluation.",
        type_=NotificationType.ACTION,
    )
    log_audit(
        db,
        entity_type="evaluation",
        entity_id=evaluation.id,
        action="evaluation_assigned",
        actor=actor,
        candidate_id=candidate_id,
        new_value={"evaluatorId": evaluator_id},
    )
    return evaluation


def _get_active_evaluator_by_email(db: Session, *, email: str) -> User | None:
    user = db.scalar(
        select(User)
        .where(
            func.lower(User.email) == email.lower(),
            User.is_active.is_(True),
        )
        .order_by(User.created_at.asc())
    )
    return user if user and _has_any_role(user, {Role.EVALUATOR}) else None


def _get_default_evaluator(db: Session) -> User:
    preferred = _get_active_evaluator_by_email(db, email=DEFAULT_EVALUATOR_EMAIL)
    if preferred is not None:
        return preferred
    fallback = next(
        (
            user
            for user in db.scalars(select(User).where(User.is_active.is_(True)).order_by(User.created_at.asc())).all()
            if _has_any_role(user, {Role.EVALUATOR})
        ),
        None,
    )
    if fallback is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No active evaluator account is available.",
        )
    return fallback


def ensure_evaluation_assignment(
    db: Session,
    *,
    candidate_id: str,
    actor: User,
    evaluator_id: str | None = None,
) -> Evaluation:
    candidate = db.get(Candidate, candidate_id)
    if candidate is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Candidate not found")

    if evaluator_id:
        # A supplied evaluator must resolve to an active user holding the EVALUATOR
        # capability — never a candidate/staff id smuggled in by the caller.
        evaluator = db.get(User, evaluator_id)
        if evaluator is None or not _has_any_role(evaluator, {Role.EVALUATOR}) or not evaluator.is_active:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Assigned evaluator not found or is not an active evaluator.",
            )
        existing = db.scalar(
            select(Evaluation)
            .where(
                Evaluation.candidate_id == candidate_id,
                Evaluation.evaluator_id == evaluator_id,
            )
            .order_by(Evaluation.created_at.desc())
        )
        if existing is not None:
            return existing
        return create_evaluation_assignment(
            db,
            candidate_id=candidate_id,
            evaluator_id=evaluator_id,
            actor=actor,
        )

    latest = db.scalar(
        select(Evaluation)
        .where(Evaluation.candidate_id == candidate_id)
        .order_by(Evaluation.created_at.desc())
    )
    if latest is not None:
        return latest

    evaluator = _get_default_evaluator(db)
    return create_evaluation_assignment(
        db,
        candidate_id=candidate_id,
        evaluator_id=evaluator.id,
        actor=actor,
    )


def submit_evaluation(db: Session, *, evaluation: Evaluation, payload: dict[str, Any], actor: User) -> Evaluation:
    # First submission stamps completed_at. A second submit would silently overwrite
    # a finalized evaluation (scores + recommendation + candidate stage). Block that
    # unless the actor is staff explicitly re-opening the record.
    if evaluation.completed_at is not None and not _has_any_role(actor, _EVALUATION_STAFF_ROLES):
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="This evaluation has already been submitted. Ask an Admin/HR to reopen it before resubmitting.",
        )
    for field, value in payload.items():
        setattr(evaluation, field, value)
    evaluation.completed_at = datetime.now(UTC)
    if evaluation.total_score is None:
        score_values = [value for value in [
            evaluation.technical_skills,
            evaluation.communication,
            evaluation.problem_solving,
            evaluation.cultural_fit,
            evaluation.attitude,
        ] if value is not None]
        evaluation.total_score = round(sum(score_values) / len(score_values) * 10, 2) if score_values else None
    db.add(evaluation)

    candidate = db.get(Candidate, evaluation.candidate_id)
    if candidate:
        candidate.current_stage = (
            CandidateStage.EVALUATION_PASSED
            if (evaluation.recommendation or "").lower() in {"passed", "shortlisted", "strongly_recommended"}
            else CandidateStage.EVALUATION_FAILED
        )
        candidate.current_status = stage_to_status(candidate.current_stage)
        db.add(candidate)
    log_audit(
        db,
        entity_type="evaluation",
        entity_id=evaluation.id,
        action="evaluation_submitted",
        actor=actor,
        candidate_id=evaluation.candidate_id,
        new_value={"totalScore": evaluation.total_score, "recommendation": evaluation.recommendation},
    )
    return evaluation


def _ical_fold(line: str) -> str:
    """
    Fold iCal content lines to max 75 octets per RFC 5545 §3.1.
    Continuation lines begin with a single SPACE.
    """
    encoded = line.encode("utf-8")
    if len(encoded) <= 75:
        return line
    folded_parts: list[str] = []
    while encoded:
        chunk = encoded[:75]
        try:
            part = chunk.decode("utf-8")
        except UnicodeDecodeError:
            chunk = encoded[:74]
            part = chunk.decode("utf-8", errors="replace")
        folded_parts.append(part)
        encoded = encoded[len(chunk):]
    return "\r\n ".join(folded_parts)


def _build_ical_event(
    *,
    uid: str,
    subject: str,
    scheduled_at: datetime,
    duration_minutes: int = 60,
    organizer_email: str,
    attendee_emails: list[str],
    description: str = "",
    location: str = "",
    meeting_url: str | None = None,
) -> bytes:
    dt_stamp = datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")
    dt_start = scheduled_at.strftime("%Y%m%dT%H%M%SZ")
    end_dt = scheduled_at + timedelta(minutes=duration_minutes)
    dt_end = end_dt.strftime("%Y%m%dT%H%M%SZ")
    # Put the join link at the top of the DESCRIPTION too — calendar clients show
    # the description prominently (and make URLs clickable), whereas LOCATION alone
    # is often not surfaced as a join link.
    full_description = description
    if meeting_url:
        join_line = f"Join the meeting: {meeting_url}"
        full_description = f"{join_line}\n\n{description}" if description else join_line
    escaped_description = full_description.replace("\n", r"\n")

    raw_lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Ethara HRMS//PI Interview//EN",
        "METHOD:REQUEST",
        "BEGIN:VEVENT",
        f"UID:{uid}",
        f"DTSTAMP:{dt_stamp}",
        f"DTSTART:{dt_start}",
        f"DTEND:{dt_end}",
        f"SUMMARY:{subject}",
        f"DESCRIPTION:{escaped_description}",
        f"LOCATION:{location}",
        f"ORGANIZER:mailto:{organizer_email}",
    ]
    if meeting_url:
        # URL is the standard iCal property; X-GOOGLE-CONFERENCE makes Google
        # Calendar render a "Join" button on import.
        raw_lines.append(f"URL:{meeting_url}")
        raw_lines.append(f"X-GOOGLE-CONFERENCE:{meeting_url}")
    for email in attendee_emails:
        raw_lines.append(f"ATTENDEE;RSVP=TRUE;ROLE=REQ-PARTICIPANT:mailto:{email}")
    raw_lines += ["END:VEVENT", "END:VCALENDAR"]

    folded = [_ical_fold(line) for line in raw_lines]
    return ("\r\n".join(folded) + "\r\n").encode("utf-8")


def _send_interview_email(
    *,
    to_email: str,
    to_name: str,
    subject: str,
    scheduled_at: datetime,
    mode: str | None,
    notes: str | None,
    ical_bytes: bytes,
    duration_minutes: int | None = None,
) -> None:
    from app.services.integrations import EmailService
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from app.core.config import get_settings

    settings = get_settings()
    mode_label = {
        "google_meet": "Google Meet",
        "offline": "Offline / In-person",
        "phone": "Phone Call",
        "teams": "Microsoft Teams",
        "zoom": "Zoom",
    }.get(mode or "", mode or "Not specified")

    online_modes = {"google_meet", "teams", "zoom"}
    mode_email_label = f"{mode_label} (online)" if mode in online_modes else mode_label
    date_str = format_app_datetime(scheduled_at, "%d %B %Y, %H:%M")
    duration_line = f"Duration: {duration_minutes} minutes\n" if duration_minutes else ""
    join_instructions = (
        "Please ensure you join the meeting on time using a stable internet connection and a "
        "working camera/microphone setup. The meeting link will be shared with you shortly before "
        "the interview (or via the same communication channel).\n\n"
        if mode in online_modes
        else "Please ensure you arrive on time for your interview.\n\n"
    )
    body_text = (
        f"Dear {to_name},\n\n"
        "We are pleased to inform you that your Personal Interview (PI) has been scheduled "
        "as per the details below:\n\n"
        f"Meeting Title: {subject}\n"
        f"Date & Time: {date_str}\n"
        f"Mode: {mode_email_label}\n"
        f"{duration_line}"
    )
    if notes:
        body_text += f"\nAdditional Notes:\n{notes}\n"
    body_text += (
        "\n" + join_instructions
        + "Kindly be available at least 5–10 minutes prior to the scheduled time.\n\n"
        "We look forward to your participation.\n\n"
        "Best regards,\nEthara Hiring Team"
    )

    duration_row = (
        f"<tr><td style='padding:4px 8px;font-weight:bold;'>Duration</td><td style='padding:4px 8px;'>{duration_minutes} minutes</td></tr>"
        if duration_minutes else ""
    )
    body_html = (
        f"<p>Dear <strong>{to_name}</strong>,</p>"
        "<p>We are pleased to inform you that your Personal Interview (PI) has been scheduled as per the details below:</p>"
        f"<table style='border-collapse:collapse;'>"
        f"<tr><td style='padding:4px 8px;font-weight:bold;'>Meeting Title</td><td style='padding:4px 8px;'>{subject}</td></tr>"
        f"<tr><td style='padding:4px 8px;font-weight:bold;'>Date &amp; Time</td><td style='padding:4px 8px;'>{date_str}</td></tr>"
        f"<tr><td style='padding:4px 8px;font-weight:bold;'>Mode</td><td style='padding:4px 8px;'>{mode_email_label}</td></tr>"
        + duration_row
        + (f"<tr><td style='padding:4px 8px;font-weight:bold;'>Notes</td><td style='padding:4px 8px;'>{notes}</td></tr>" if notes else "")
        + "</table>"
        f"<p>{join_instructions.strip()}</p>"
        "<p>Kindly be available at least 5–10 minutes prior to the scheduled time. We look forward to your participation.</p>"
        f"<p>Best regards,<br/>Ethara Hiring Team</p>"
    )

    if settings.email_backend == "console":
        print(f"[calendar-invite] to={to_email} subject='{subject}' at={date_str} mode={mode_label}")
        if notes:
            print(f"[calendar-invite] notes={notes}")
        return

    if settings.email_backend == "smtp":
        if not settings.smtp_host or not settings.smtp_username or not settings.smtp_password:
            return
        msg = MIMEMultipart("mixed")
        msg["From"] = settings.email_from
        msg["To"] = to_email
        msg["Subject"] = f"Invitation: {subject} ({mode_email_label})"
        alt = MIMEMultipart("alternative")
        alt.attach(MIMEText(body_text, "plain"))
        alt.attach(MIMEText(body_html, "html"))
        msg.attach(alt)
        ical_part = MIMEText(ical_bytes.decode("utf-8"), "calendar", "utf-8")
        ical_part.replace_header("Content-Type", "text/calendar; charset=UTF-8; method=REQUEST")
        ical_part["Content-Disposition"] = 'attachment; filename="interview.ics"'
        msg.attach(ical_part)
        try:
            with smtplib.SMTP(settings.smtp_host, settings.smtp_port, timeout=10) as smtp:
                if settings.smtp_use_tls:
                    smtp.ehlo()
                    smtp.starttls()
                    smtp.ehlo()
                smtp.login(settings.smtp_username, settings.smtp_password)
                smtp.send_message(msg)
        except Exception:
            pass
        return

    svc = EmailService()
    try:
        svc.send_email(to_email=to_email, subject=f"Invitation: {subject} ({mode_email_label})", body_text=body_text, body_html=body_html)
    except Exception:
        pass


def _send_pms_meeting_email(
    *,
    to_email: str,
    to_name: str,
    title: str,
    scheduled_at: datetime,
    duration_minutes: int,
    location: str | None,
    notes: str | None,
    ical_bytes: bytes,
    meeting_url: str | None = None,
) -> None:
    """Email a single attendee a PMS review-meeting calendar invite (online only)."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText

    from app.core.config import get_settings
    from app.services.integrations import EmailService

    settings = get_settings()
    date_str = format_app_datetime(scheduled_at, "%d %B %Y, %H:%M")
    join_text = f"Join the meeting: {meeting_url}\n" if meeting_url else ""
    location_line = "" if meeting_url else (f"Location: {location}\n" if location else "")
    body_text = (
        f"Dear {to_name},\n\n"
        "A Performance Management System (PMS) review meeting has been scheduled "
        "with the details below:\n\n"
        f"Meeting Title: {title}\n"
        f"Date & Time: {date_str}\n"
        f"Mode: Online\n"
        f"Duration: {duration_minutes} minutes\n"
        f"{join_text}{location_line}"
    )
    if notes:
        body_text += f"\nAdditional Notes:\n{notes}\n"
    body_text += (
        "\nPlease join on time using a stable internet connection. Kindly be available "
        "at least 5 minutes prior to the scheduled time.\n\n"
        "Best regards,\nEthara HR Team"
    )

    link_row = (
        f"<tr><td style='padding:4px 8px;font-weight:bold;'>Meeting Link</td>"
        f"<td style='padding:4px 8px;'><a href='{meeting_url}'>{meeting_url}</a></td></tr>"
        if meeting_url else (
            f"<tr><td style='padding:4px 8px;font-weight:bold;'>Location</td>"
            f"<td style='padding:4px 8px;'>{location}</td></tr>" if location else ""
        )
    )
    join_button = (
        f"<p style='margin:16px 0;'><a href='{meeting_url}' "
        "style='display:inline-block;padding:10px 20px;background:#ED00ED;color:#ffffff;"
        "border-radius:8px;text-decoration:none;font-weight:bold;'>Join the meeting</a></p>"
        if meeting_url else ""
    )
    body_html = (
        f"<p>Dear <strong>{to_name}</strong>,</p>"
        "<p>A Performance Management System (PMS) review meeting has been scheduled with the details below:</p>"
        "<table style='border-collapse:collapse;'>"
        f"<tr><td style='padding:4px 8px;font-weight:bold;'>Meeting Title</td><td style='padding:4px 8px;'>{title}</td></tr>"
        f"<tr><td style='padding:4px 8px;font-weight:bold;'>Date &amp; Time</td><td style='padding:4px 8px;'>{date_str}</td></tr>"
        "<tr><td style='padding:4px 8px;font-weight:bold;'>Mode</td><td style='padding:4px 8px;'>Online</td></tr>"
        f"<tr><td style='padding:4px 8px;font-weight:bold;'>Duration</td><td style='padding:4px 8px;'>{duration_minutes} minutes</td></tr>"
        + link_row
        + (f"<tr><td style='padding:4px 8px;font-weight:bold;'>Notes</td><td style='padding:4px 8px;'>{notes}</td></tr>" if notes else "")
        + "</table>"
        + join_button
        + "<p>Please join on time using a stable internet connection. Kindly be available at least 5 minutes prior to the scheduled time.</p>"
        "<p>Best regards,<br/>Ethara HR Team</p>"
    )

    subject_line = f"Invitation: {title} (PMS Review)"

    if settings.email_backend == "console":
        print(f"[pms-meeting-invite] to={to_email} title='{title}' at={date_str}")
        return

    if settings.email_backend == "smtp":
        if not settings.smtp_host or not settings.smtp_username or not settings.smtp_password:
            return
        msg = MIMEMultipart("mixed")
        msg["From"] = settings.email_from
        msg["To"] = to_email
        msg["Subject"] = subject_line
        alt = MIMEMultipart("alternative")
        alt.attach(MIMEText(body_text, "plain"))
        alt.attach(MIMEText(body_html, "html"))
        msg.attach(alt)
        ical_part = MIMEText(ical_bytes.decode("utf-8"), "calendar", "utf-8")
        ical_part.replace_header("Content-Type", "text/calendar; charset=UTF-8; method=REQUEST")
        ical_part["Content-Disposition"] = 'attachment; filename="pms-meeting.ics"'
        msg.attach(ical_part)
        try:
            with smtplib.SMTP(settings.smtp_host, settings.smtp_port, timeout=10) as smtp:
                if settings.smtp_use_tls:
                    smtp.ehlo()
                    smtp.starttls()
                    smtp.ehlo()
                smtp.login(settings.smtp_username, settings.smtp_password)
                smtp.send_message(msg)
        except Exception:
            pass
        return

    try:
        EmailService().send_email(
            to_email=to_email, subject=subject_line, body_text=body_text, body_html=body_html
        )
    except Exception:
        pass


def send_pms_meeting_invites(
    *,
    meeting_id: str,
    title: str,
    scheduled_at: datetime,
    duration_minutes: int,
    location: str | None,
    notes: str | None,
    organizer: User,
    employee_name: str,
    employee_email: str | None,
    invite_employee: bool,
    extra_attendees: list[str],
) -> list[str]:
    """Build a calendar invite for an online PMS meeting and email every participant.

    The organizer (the account that set up the meeting) is always included. The
    employee being reviewed is included when ``invite_employee`` is set and an
    email is on file. Returns the de-duplicated list of emails actually notified.
    """
    import logging as _logging

    from app.core.config import get_settings

    _logger = _logging.getLogger(__name__)
    settings = get_settings()
    owner_email = settings.email_from  # noreply@ethara.ai owns the shared calendar

    # Build the attendee list: organizer first, then employee, then extras.
    name_by_email: dict[str, str] = {}
    attendee_emails: list[str] = []

    def _add(email: str | None, name: str) -> None:
        if not email:
            return
        cleaned = email.strip()
        if not cleaned or cleaned == owner_email or cleaned in name_by_email:
            return
        name_by_email[cleaned] = name
        attendee_emails.append(cleaned)

    _add(organizer.email, organizer.name or "Organizer")
    if invite_employee:
        _add(employee_email, employee_name or "Employee")
    for member in extra_attendees:
        _add(member, "Attendee")

    # The "location" field for an online PMS meeting is the join link. Normalise it
    # to a URL so it lands in the calendar event as a clickable join link.
    meeting_url: str | None = None
    if location and location.strip():
        loc = location.strip()
        if loc.lower().startswith(("http://", "https://")):
            meeting_url = loc
        elif "." in loc and " " not in loc:  # bare link e.g. meet.google.com/abc
            meeting_url = f"https://{loc}"

    ical_uid = f"pms-{meeting_id}@ethara.ai"
    ical_bytes = _build_ical_event(
        uid=ical_uid,
        subject=title,
        scheduled_at=scheduled_at,
        duration_minutes=duration_minutes,
        organizer_email=owner_email,
        attendee_emails=attendee_emails,
        description=notes or "",
        location=meeting_url or location or "Online",
        meeting_url=meeting_url,
    )

    for email in attendee_emails:
        try:
            _send_pms_meeting_email(
                to_email=email,
                to_name=name_by_email.get(email, "Attendee"),
                title=title,
                scheduled_at=scheduled_at,
                duration_minutes=duration_minutes,
                location=location,
                notes=notes,
                ical_bytes=ical_bytes,
                meeting_url=meeting_url,
            )
        except Exception as exc:
            _logger.warning("Failed to send PMS meeting invite to %s: %s", email, exc)

    # Send to the noreply inbox too so the event lands in the shared Google Calendar.
    try:
        _send_pms_meeting_email(
            to_email=owner_email,
            to_name="Ethara PMS",
            title=title,
            scheduled_at=scheduled_at,
            duration_minutes=duration_minutes,
            location=location,
            notes=notes,
            ical_bytes=ical_bytes,
            meeting_url=meeting_url,
        )
    except Exception as exc:
        _logger.warning("Failed to send PMS meeting invite to organizer inbox %s: %s", owner_email, exc)

    return attendee_emails


def schedule_interview(
    db: Session,
    *,
    evaluation: Evaluation,
    subject: str,
    scheduled_at: datetime,
    notes: str | None,
    actor: User,
    mode: str | None = None,
    duration_minutes: int = 60,
    round_number: int | None = None,
    evaluator_id: str | None = None,
    panel_label: str | None = None,
    panel_members: list[str] | None = None,
) -> Evaluation:
    rounds = ensure_legacy_pi_rounds(db, evaluation=evaluation)
    rounds = sorted(rounds, key=lambda item: item.round_number)

    normalized_panel_members = [member.strip() for member in (panel_members or []) if member and member.strip()] or None
    requested_round = round_number if round_number is not None else None
    if requested_round is not None and (requested_round < 1 or requested_round > MAX_PI_ROUNDS):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"PI round number must be between 1 and {MAX_PI_ROUNDS}.",
        )

    target_round = None
    if requested_round is not None:
        target_round = next((item for item in rounds if item.round_number == requested_round), None)
    elif rounds:
        latest_round = rounds[-1]
        if latest_round.status in PI_ACTIVE_STATUSES and latest_round.completed_at is None:
            target_round = latest_round

    if target_round is not None and target_round.completed_at is not None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Completed PI rounds cannot be rescheduled. Create the next round instead.",
        )

    assigned_round_number = requested_round
    if target_round is None:
        if assigned_round_number is None:
            assigned_round_number = (rounds[-1].round_number + 1) if rounds else 1
        if assigned_round_number > MAX_PI_ROUNDS:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Only {MAX_PI_ROUNDS} PI rounds are supported per candidate.",
            )
        if any(item.round_number == assigned_round_number for item in rounds):
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=f"PI round {assigned_round_number} already exists for this candidate.",
            )
        target_round = PiInterviewRound(
            evaluation_id=evaluation.id,
            candidate_id=evaluation.candidate_id,
            round_number=assigned_round_number,
            status="scheduled",
        )
        db.add(target_round)
        db.flush()
        rounds.append(target_round)
        is_reschedule = False
    else:
        is_reschedule = target_round.scheduled_at is not None
        assigned_round_number = target_round.round_number

    assigned_evaluator = None
    assigned_evaluator_id = evaluator_id or target_round.evaluator_id or evaluation.evaluator_id
    if assigned_evaluator_id:
        assigned_evaluator = db.get(User, assigned_evaluator_id)
        if assigned_evaluator is None or not _has_any_role(assigned_evaluator, {Role.EVALUATOR}) or not assigned_evaluator.is_active:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Assigned PI evaluator not found.")
        evaluation.evaluator_id = assigned_evaluator.id
        target_round.evaluator_id = assigned_evaluator.id

    target_round.subject = subject
    target_round.scheduled_at = scheduled_at
    target_round.completed_at = None
    target_round.status = "rescheduled" if is_reschedule else "scheduled"
    target_round.mode = mode
    target_round.duration_minutes = duration_minutes or target_round.duration_minutes or 60
    target_round.panel_label = panel_label.strip() if panel_label and panel_label.strip() else None
    target_round.panel_members = normalized_panel_members
    target_round.round_decision = None
    target_round.no_further_pi_required = False
    target_round.final_verdict = None
    target_round.score = None
    if notes is not None:
        target_round.remarks = notes.strip() or None
    db.add(target_round)

    sync_evaluation_interview_summary(evaluation, latest_round=target_round)
    db.add(evaluation)

    candidate = db.get(Candidate, evaluation.candidate_id)
    evaluator = assigned_evaluator or db.get(User, target_round.evaluator_id or evaluation.evaluator_id)

    from app.core.config import get_settings
    settings = get_settings()
    organizer_email = settings.email_from  # noreply@ethara.ai owns the calendar

    attendee_emails: list[str] = []
    if candidate:
        attendee_emails.append(candidate.personal_email)
    if evaluator and evaluator.email != organizer_email:
        attendee_emails.append(evaluator.email)
    # actor (HR/manager who scheduled) gets the invite too if different from evaluator
    if actor.email != organizer_email and actor.email not in attendee_emails:
        attendee_emails.append(actor.email)
    # Interviewer / panel emails entered while scheduling also receive the invite.
    for _member in (normalized_panel_members or []):
        if _member and _member != organizer_email and _member not in attendee_emails:
            attendee_emails.append(_member)

    ical_uid = f"pi-{evaluation.id}-round-{assigned_round_number}@ethara.ai"
    ical_bytes = _build_ical_event(
        uid=ical_uid,
        subject=subject,
        scheduled_at=scheduled_at,
        duration_minutes=target_round.duration_minutes,
        organizer_email=organizer_email,
        attendee_emails=attendee_emails,
        description=notes or "",
        location=mode or "",
    )

    import logging as _logging
    _logger = _logging.getLogger(__name__)

    if candidate:
        try:
            _send_interview_email(
                to_email=candidate.personal_email,
                to_name=candidate.full_name,
                subject=subject,
                scheduled_at=scheduled_at,
                mode=mode,
                notes=notes,
                ical_bytes=ical_bytes,
            )
        except Exception as exc:
            _logger.warning("Failed to send PI invite to candidate %s: %s", candidate.personal_email, exc)

    if evaluator and evaluator.email != (candidate.personal_email if candidate else None):
        try:
            _send_interview_email(
                to_email=evaluator.email,
                to_name=evaluator.name,
                subject=subject,
                scheduled_at=scheduled_at,
                mode=mode,
                notes=notes,
                ical_bytes=ical_bytes,
            )
        except Exception as exc:
            _logger.warning("Failed to send PI invite to evaluator %s: %s", evaluator.email, exc)

    # Send the invite to each interviewer / panel email entered while scheduling.
    _already_emailed = {organizer_email, (candidate.personal_email if candidate else None), (evaluator.email if evaluator else None)}
    for _member in (normalized_panel_members or []):
        if not _member or _member in _already_emailed:
            continue
        _already_emailed.add(_member)
        try:
            _send_interview_email(
                to_email=_member,
                to_name="Interviewer",
                subject=subject,
                scheduled_at=scheduled_at,
                mode=mode,
                notes=notes,
                ical_bytes=ical_bytes,
                duration_minutes=target_round.duration_minutes,
            )
        except Exception as exc:
            _logger.warning("Failed to send PI invite to interviewer %s: %s", _member, exc)

    # Send invite to noreply so the event appears in the noreply Google Calendar
    try:
        _send_interview_email(
            to_email=organizer_email,
            to_name="Ethara Interviews",
            subject=subject,
            scheduled_at=scheduled_at,
            mode=mode,
            notes=notes,
            ical_bytes=ical_bytes,
            duration_minutes=duration_minutes,
        )
    except Exception as exc:
        _logger.warning("Failed to send PI invite to organizer inbox %s: %s", organizer_email, exc)

    action_label = "rescheduled" if is_reschedule else "scheduled"
    if target_round.evaluator_id:
        create_notification(
            db,
            user_id=target_round.evaluator_id,
            candidate_id=evaluation.candidate_id,
            title=f"PI Round {assigned_round_number} {action_label.capitalize()}",
            message=f"PI round {assigned_round_number} {action_label}: {subject} at {format_app_datetime(scheduled_at, '%d %b %Y %H:%M')}",
            type_=NotificationType.INFO,
        )
    if candidate and candidate.portal_user_id:
        create_notification(
            db,
            user_id=candidate.portal_user_id,
            candidate_id=evaluation.candidate_id,
            title=f"PI Round {assigned_round_number} {action_label.capitalize()}",
            message=f"Your PI round {assigned_round_number} has been {action_label}: {subject} on {format_app_datetime(scheduled_at, '%d %b %Y at %H:%M')}",
            type_=NotificationType.INFO,
        )
    log_audit(
        db,
        entity_type="evaluation",
        entity_id=evaluation.id,
        action=f"interview_{action_label}",
        actor=actor,
        candidate_id=evaluation.candidate_id,
        new_value={
            "roundNumber": assigned_round_number,
            "subject": subject,
            "scheduledAt": scheduled_at.isoformat(),
            "interviewStatus": target_round.status,
            "mode": mode,
            "evaluatorId": target_round.evaluator_id,
            "panelLabel": target_round.panel_label,
            "panelMembers": target_round.panel_members or [],
        },
    )
    return evaluation


def complete_interview(
    db: Session,
    *,
    evaluation: Evaluation,
    decision: str,
    notes: str | None,
    actor: User,
    pi_score: float | None = None,
    round_id: str | None = None,
    round_number: int | None = None,
    no_further_pi_required: bool = False,
    final_verdict: str | None = None,
) -> Evaluation:
    rounds = ensure_legacy_pi_rounds(db, evaluation=evaluation)
    rounds = sorted(rounds, key=lambda item: item.round_number)
    if not rounds:
        raise HTTPException(status_code=400, detail="No PI round exists for this evaluation yet.")

    target_round = None
    if round_id:
        target_round = next((item for item in rounds if item.id == round_id), None)
    elif round_number is not None:
        target_round = next((item for item in rounds if item.round_number == round_number), None)
    else:
        target_round = next(
            (item for item in reversed(rounds) if item.status in PI_ACTIVE_STATUSES or item.completed_at is None),
            rounds[-1],
        )

    if target_round is None:
        raise HTTPException(status_code=404, detail="PI round not found for this evaluation.")

    # Round-level authorization: when a PI round is assigned to a specific evaluator,
    # only that evaluator (or staff) may complete/score it — another evaluator who
    # happens to own the parent Evaluation must not close someone else's round.
    if (
        target_round.evaluator_id
        and target_round.evaluator_id != actor.id
        and not _has_any_role(actor, _EVALUATION_STAFF_ROLES)
    ):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="This interview round is assigned to another evaluator.",
        )

    normalized_final_verdict = _normalize_pi_final_verdict(final_verdict)
    normalized_decision = _normalize_pi_round_decision(decision)
    final_round = no_further_pi_required

    # Preserve the current one-click pass/fail flow by treating legacy decisions as final verdicts.
    if normalized_decision in PI_FINAL_VERDICTS and normalized_final_verdict is None:
        normalized_final_verdict = normalized_decision
        final_round = True

    if final_round and normalized_final_verdict is None:
        raise HTTPException(
            status_code=400,
            detail="A final PI verdict of Selected or Rejected is required when no further PI is needed.",
        )
    if not final_round and normalized_decision not in {"proceed_to_next_round", None}:
        normalized_decision = "proceed_to_next_round"
    if not final_round and target_round.round_number >= MAX_PI_ROUNDS:
        raise HTTPException(
            status_code=400,
            detail=f"PI round {MAX_PI_ROUNDS} must be closed with a final Selected or Rejected verdict.",
        )

    target_round.evaluator_id = target_round.evaluator_id or actor.id
    target_round.completed_at = datetime.now(UTC)
    target_round.status = "no_further_pi_required" if final_round else "completed"
    target_round.score = pi_score if pi_score is not None else target_round.score
    target_round.remarks = notes.strip() if notes and notes.strip() else target_round.remarks
    target_round.round_decision = normalized_decision or ("selected" if normalized_final_verdict == "selected" else "rejected" if normalized_final_verdict == "rejected" else None)
    target_round.no_further_pi_required = final_round
    target_round.final_verdict = normalized_final_verdict
    db.add(target_round)

    sync_evaluation_interview_summary(evaluation)
    if normalized_final_verdict:
        evaluation.recommendation = normalized_final_verdict
    if pi_score is not None:
        evaluation.pi_score = pi_score
    db.add(evaluation)

    candidate = db.get(Candidate, evaluation.candidate_id)
    if candidate and normalized_final_verdict:
        previous_stage = candidate.current_stage
        if normalized_final_verdict == "rejected":
            candidate.current_stage = CandidateStage.EVALUATION_FAILED
            candidate.current_status = f"Rejected after PI Round {target_round.round_number}"
        else:
            if candidate.current_stage in {
                CandidateStage.SELECTION_FORM_SENT,
                CandidateStage.SELECTION_FORM_SUBMITTED,
                CandidateStage.SELECTION_FORM_VALIDATED,
            }:
                if candidate.current_stage != CandidateStage.SELECTION_FORM_VALIDATED:
                    candidate.current_stage = CandidateStage.SELECTION_FORM_VALIDATED
                candidate.current_status = f"Selected after PI Round {target_round.round_number}"
            else:
                candidate.current_stage = CandidateStage.EVALUATION_PASSED
                candidate.current_status = f"Selected after PI Round {target_round.round_number}"
        db.add(candidate)
        _add_stage_log(
            db,
            candidate=candidate,
            from_stage=previous_stage,
            to_stage=candidate.current_stage,
            actor=actor,
            notes=(
                f"PI round {target_round.round_number} closed as {normalized_final_verdict}."
                if not notes else f"PI round {target_round.round_number} closed as {normalized_final_verdict}. {notes.strip()}"
            ),
        )

    log_audit(
        db,
        entity_type="evaluation",
        entity_id=evaluation.id,
        action="interview_completed",
        actor=actor,
        candidate_id=evaluation.candidate_id,
        new_value={
            "roundId": target_round.id,
            "roundNumber": target_round.round_number,
            "decision": target_round.round_decision,
            "finalVerdict": target_round.final_verdict,
            "interviewStatus": target_round.status,
            "score": target_round.score,
            "noFurtherPiRequired": target_round.no_further_pi_required,
        },
    )
    return evaluation


def bypass_pi_interview(
    db: Session,
    *,
    evaluation: Evaluation,
    actor: User,
    final_verdict: str,
    notes: str | None = None,
    pi_score: float | None = None,
) -> Evaluation:
    normalized_final_verdict = _normalize_pi_final_verdict(final_verdict)
    if normalized_final_verdict is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Final PI verdict must be Selected or Rejected.",
        )
    if pi_score is not None and (pi_score < 0 or pi_score > 100):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="PI score must be between 0 and 100.",
        )

    rounds = sorted(evaluation.pi_rounds or [], key=lambda item: item.round_number)
    target_round = next(
        (
            item
            for item in reversed(rounds)
            if item.status in PI_ACTIVE_STATUSES or item.completed_at is None
        ),
        None,
    )
    now = datetime.now(UTC)

    if target_round is None:
        next_round = (rounds[-1].round_number + 1) if rounds else 1
        if next_round > MAX_PI_ROUNDS:
            target_round = rounds[-1]
        else:
            target_round = PiInterviewRound(
                evaluation_id=evaluation.id,
                candidate_id=evaluation.candidate_id,
                round_number=next_round,
            )
            db.add(target_round)
            evaluation.pi_rounds.append(target_round)

    if (
        target_round.evaluator_id
        and target_round.evaluator_id != actor.id
        and not _has_any_role(actor, _EVALUATION_STAFF_ROLES)
    ):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="This interview round is assigned to another evaluator.",
        )

    target_round.evaluator_id = target_round.evaluator_id or actor.id
    target_round.subject = target_round.subject or "Interview Bypassed"
    target_round.completed_at = now
    target_round.status = "pi_bypassed"
    target_round.score = pi_score if pi_score is not None else target_round.score
    target_round.remarks = notes.strip() if notes and notes.strip() else "PI bypassed."
    target_round.round_decision = normalized_final_verdict
    target_round.no_further_pi_required = True
    target_round.final_verdict = normalized_final_verdict
    db.add(target_round)

    sync_evaluation_interview_summary(evaluation, latest_round=target_round)
    evaluation.recommendation = normalized_final_verdict
    if pi_score is not None:
        evaluation.pi_score = pi_score
    db.add(evaluation)

    candidate = db.get(Candidate, evaluation.candidate_id)
    if candidate:
        previous_stage = candidate.current_stage
        if normalized_final_verdict == "rejected":
            candidate.current_stage = CandidateStage.EVALUATION_FAILED
            candidate.current_status = "Rejected after PI bypass"
        elif candidate.current_stage in {
            CandidateStage.SELECTION_FORM_SENT,
            CandidateStage.SELECTION_FORM_SUBMITTED,
            CandidateStage.SELECTION_FORM_VALIDATED,
        }:
            candidate.current_status = "Selected after PI bypass"
        else:
            candidate.current_stage = CandidateStage.SELECTION_FORM_SENT
            candidate.current_status = "Selection Form Sent (Interview Bypassed)"
            apply_stage_side_effects(db, candidate, actor=actor)
        db.add(candidate)
        _add_stage_log(
            db,
            candidate=candidate,
            from_stage=previous_stage,
            to_stage=candidate.current_stage,
            actor=actor,
            notes=(
                f"PI bypassed as {normalized_final_verdict}."
                if not notes else f"PI bypassed as {normalized_final_verdict}. {notes.strip()}"
            ),
        )

    log_audit(
        db,
        entity_type="evaluation",
        entity_id=evaluation.id,
        action="pi_bypassed",
        actor=actor,
        candidate_id=evaluation.candidate_id,
        new_value={
            "roundId": target_round.id,
            "roundNumber": target_round.round_number,
            "finalVerdict": target_round.final_verdict,
            "score": target_round.score,
            "notes": notes,
        },
    )
    return evaluation


def get_or_create_selection_form(db: Session, *, candidate_id: str) -> SelectionForm:
    return ensure_selection_form(db, candidate_id)


def serialize_candidate_id_card_form(
    record: CandidateIdCardForm | None,
    *,
    candidate: Candidate,
) -> dict[str, Any]:
    return {
        "id": record.id if record else None,
        "candidateId": candidate.id,
        "name": record.name if record and record.name else candidate.full_name,
        # Auto-filled from the candidate's system-generated GRP code even before first submit.
        "employeeId": (record.employee_id if record and record.employee_id else candidate.employee_code),
        "bloodGroup": record.blood_group if record else None,
        "emergencyNo": record.emergency_no if record else None,
        "submittedAt": record.submitted_at if record else None,
        "submittedBy": record.submitted_by if record else None,
        "itCompletedAt": record.it_completed_at if record else None,
        "itCompletedBy": record.it_completed_by if record else None,
        "createdAt": record.created_at if record else None,
        "updatedAt": record.updated_at if record else None,
    }


def _candidate_id_card_queue_status(
    record: CandidateIdCardForm | None,
) -> str:
    if record and record.it_completed_at:
        return "done"
    if record and record.submitted_at:
        return "ready"
    return "awaiting_details"


def serialize_candidate_id_card_queue_item(
    *,
    candidate: Candidate,
    record: CandidateIdCardForm | None,
    designation: str | None = None,
    photo_url: str | None = None,
) -> dict[str, Any]:
    status = _candidate_id_card_queue_status(record)
    payload = serialize_candidate_id_card_form(record, candidate=candidate)
    payload.update({
        "candidateName": candidate.full_name,
        "personalEmail": candidate.personal_email,
        "etharaEmail": candidate.ethara_email,
        "currentStage": candidate.current_stage.value if candidate.current_stage else None,
        "currentStatus": candidate.current_status,
        "designation": designation,
        "photoUrl": photo_url,
        "status": status,
        "canMarkDone": status == "ready",
    })
    return payload


def get_candidate_id_card_form(db: Session, *, candidate: Candidate) -> dict[str, Any]:
    record = db.scalar(
        select(CandidateIdCardForm).where(CandidateIdCardForm.candidate_id == candidate.id)
    )
    return serialize_candidate_id_card_form(record, candidate=candidate)


def list_candidate_id_card_queue(db: Session) -> list[dict[str, Any]]:
    rows = db.execute(
        select(Candidate, CandidateIdCardForm, Position.title)
        .outerjoin(CandidateIdCardForm, CandidateIdCardForm.candidate_id == Candidate.id)
        .outerjoin(Position, Candidate.position_id == Position.id)
        .where(
            func.length(func.trim(func.coalesce(Candidate.ethara_email, ""))) > 0,
            Candidate.is_removed.is_(False),
        )
        .order_by(Candidate.full_name.asc(), Candidate.created_at.asc())
    ).all()

    photo_types = {"photo", "passport_photo", "passport_size_photo", "selection_form_passport_size_photo", "profile_photo", "photograph"}
    photo_by_candidate: dict[str, str] = {}
    candidate_ids = [candidate.id for candidate, _, _ in rows]
    if candidate_ids:
        # Later rows overwrite earlier ones, so the newest photo wins.
        for candidate_id, file_url in db.execute(
            select(Document.candidate_id, Document.file_url)
            .where(Document.candidate_id.in_(candidate_ids), Document.type.in_(photo_types))
            .order_by(Document.created_at.asc())
        ).all():
            if file_url:
                photo_by_candidate[candidate_id] = file_url

    # Photos uploaded via the employee detail form live in EmployeeDocument, not the
    # candidate document table — match profiles by employee code / email and use those
    # photos for candidates that have none of their own.
    if rows:
        from app.db.models import EmployeeDocument, EmployeeProfile

        profiles = db.scalars(select(EmployeeProfile)).all()
        profile_by_code: dict[str, EmployeeProfile] = {}
        profile_by_email: dict[str, EmployeeProfile] = {}
        for profile in profiles:
            if profile.employee_code:
                profile_by_code[profile.employee_code.strip().lower()] = profile
            for email in (profile.ethara_email, profile.personal_email):
                if email:
                    profile_by_email.setdefault(email.strip().lower(), profile)

        employee_photo_by_profile: dict[str, str] = {}
        for profile_id, doc_type, file_url in db.execute(
            select(EmployeeDocument.employee_profile_id, EmployeeDocument.type, EmployeeDocument.file_url)
            .order_by(EmployeeDocument.created_at.asc())
        ).all():
            if file_url and str(doc_type or "").strip().lower() in photo_types:
                employee_photo_by_profile[profile_id] = file_url

        for candidate, record, _ in rows:
            if candidate.id in photo_by_candidate:
                continue
            profile = None
            code = (record.employee_id if record and record.employee_id else candidate.employee_code) or ""
            if code:
                profile = profile_by_code.get(code.strip().lower())
            if profile is None:
                for email in (candidate.ethara_email, candidate.personal_email):
                    if email:
                        profile = profile_by_email.get(email.strip().lower())
                        if profile is not None:
                            break
            if profile is not None:
                file_url = employee_photo_by_profile.get(profile.id)
                if file_url:
                    photo_by_candidate[candidate.id] = file_url

    def _signed_photo(candidate_id: str) -> str | None:
        file_url = photo_by_candidate.get(candidate_id)
        if not file_url:
            return None
        try:
            return make_signed_upload_url(file_url)
        except Exception:
            return None

    items = [
        serialize_candidate_id_card_queue_item(
            candidate=candidate,
            record=record,
            designation=position_title,
            photo_url=_signed_photo(candidate.id),
        )
        for candidate, record, position_title in rows
    ]
    return sorted(
        items,
        key=lambda item: (
            {"ready": 0, "awaiting_details": 1, "done": 2}.get(str(item.get("status")), 3),
            str(item.get("candidateName") or "").lower(),
        ),
    )


def submit_candidate_id_card_form(
    db: Session,
    *,
    candidate: Candidate,
    actor: User,
    payload: dict[str, Any],
) -> dict[str, Any]:
    if not candidate.ethara_email:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Create the Ethara email ID before submitting the ID card form.",
        )

    # The employee code is system-generated (sequential GRP, allocated at contract signing)
    # and is NOT taken from client input — it is auto-filled here so it can never be tampered
    # with or left blank. Allocate one now for any legacy candidate that doesn't have it yet.
    from app.services.employees import assign_candidate_employee_code, normalize_blood_group

    employee_code = candidate.employee_code or assign_candidate_employee_code(db, candidate)

    cleaned = {
        "name": str(payload.get("name") or "").strip(),
        "employee_id": employee_code,
        "blood_group": normalize_blood_group(payload.get("blood_group")),
        "emergency_no": str(payload.get("emergency_no") or "").strip(),
    }
    if not (cleaned["name"] and cleaned["blood_group"] and cleaned["emergency_no"]):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="Name, blood group, and emergency number are required.",
        )

    record = db.scalar(
        select(CandidateIdCardForm).where(CandidateIdCardForm.candidate_id == candidate.id)
    )
    is_first_submit = record is None or record.submitted_at is None
    if record is None:
        record = CandidateIdCardForm(candidate_id=candidate.id)

    record.name = cleaned["name"]
    record.employee_id = cleaned["employee_id"]
    record.blood_group = cleaned["blood_group"]
    record.emergency_no = cleaned["emergency_no"]
    record.submitted_at = datetime.now(UTC)
    record.submitted_by = actor.id
    record.it_completed_at = None
    record.it_completed_by = None
    db.add(record)
    db.flush()

    notify_roles(
        db,
        roles={Role.ADMIN, Role.LEADERSHIP, Role.HR, Role.TA, Role.IT_TEAM},
        candidate_id=candidate.id,
        title="Candidate ID card form submitted" if is_first_submit else "Candidate ID card form updated",
        message=(
            f"{candidate.full_name} submitted ID card details for {cleaned['employee_id']}."
            if is_first_submit
            else f"{candidate.full_name}'s ID card details were updated."
        ),
        type_=NotificationType.ACTION,
    )
    log_audit(
        db,
        entity_type="candidate_id_card_form",
        entity_id=record.id,
        action="candidate_id_card_form_submitted" if is_first_submit else "candidate_id_card_form_updated",
        actor=actor,
        candidate_id=candidate.id,
        new_value={
            "name": record.name,
            "employeeId": record.employee_id,
            "bloodGroup": record.blood_group,
            "emergencyNo": record.emergency_no,
        },
    )
    return serialize_candidate_id_card_form(record, candidate=candidate)


def mark_candidate_id_card_forms_done(
    db: Session,
    *,
    candidate_ids: list[str],
    actor: User,
) -> dict[str, Any]:
    normalized_ids = [candidate_id.strip() for candidate_id in candidate_ids if candidate_id.strip()]
    if not normalized_ids:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="Select at least one candidate ID card to mark as done.",
        )

    rows = db.execute(
        select(Candidate, CandidateIdCardForm)
        .outerjoin(CandidateIdCardForm, CandidateIdCardForm.candidate_id == Candidate.id)
        .where(Candidate.id.in_(normalized_ids))
    ).all()

    records_by_candidate = {
        candidate.id: (candidate, record)
        for candidate, record in rows
    }
    missing_ids = [candidate_id for candidate_id in normalized_ids if candidate_id not in records_by_candidate]
    if missing_ids:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="One or more candidates were not found.")

    updated_candidate_ids: list[str] = []
    now = datetime.now(UTC)
    for candidate_id in normalized_ids:
        candidate, record = records_by_candidate[candidate_id]
        if not candidate.ethara_email or record is None or record.submitted_at is None:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="Only submitted ID card forms can be marked as done.",
            )
        if record.it_completed_at is not None:
            continue
        record.it_completed_at = now
        record.it_completed_by = actor.id
        db.add(record)
        log_audit(
            db,
            entity_type="candidate_id_card_form",
            entity_id=record.id,
            action="candidate_id_card_marked_done",
            actor=actor,
            candidate_id=candidate.id,
            new_value={
                "status": "done",
                "itCompletedBy": record.it_completed_by,
            },
        )
        updated_candidate_ids.append(candidate.id)

    return {
        "updatedCount": len(updated_candidate_ids),
        "updatedCandidateIds": updated_candidate_ids,
    }


# Tokens recognised in an uploaded ID-card status sheet. "Done" means the physical card
# has been created/issued; "Pending" means it is still outstanding. Anything else is
# reported back to the uploader rather than silently ignored.
_ID_CARD_DONE_TOKENS = {"done", "created", "issued", "complete", "completed", "yes", "y", "true", "1"}
_ID_CARD_PENDING_TOKENS = {"pending", "not done", "notdone", "not created", "incomplete", "no", "n", "false", "0"}


def id_card_status_template_csv() -> str:
    """A tiny example status sheet for office admin / HR to fill in and re-upload.

    One row per member: their Ethara email and whether the physical ID card has been
    created ("Done") or is still outstanding ("Pending").
    """
    buffer = StringIO()
    writer = csv.writer(buffer)
    writer.writerow(["Email", "Status"])
    writer.writerow(["jane.doe@ethara.ai", "Done"])
    writer.writerow(["john.smith@ethara.ai", "Pending"])
    return buffer.getvalue()


def parse_id_card_status_csv(raw: bytes) -> list[dict[str, str]]:
    """Parse an uploaded status sheet into [{email, status}]. Raises ValueError."""
    if len(raw) > 5 * 1024 * 1024:
        raise ValueError("File too large (max 5MB).")
    try:
        text = raw.decode("utf-8-sig")
    except UnicodeDecodeError:
        text = raw.decode("latin-1", errors="ignore")
    reader = csv.DictReader(StringIO(text))
    if not reader.fieldnames:
        raise ValueError("The file appears to be empty or has no header row.")
    norm = {(header or "").strip().lower(): header for header in reader.fieldnames}

    def pick(*aliases: str) -> str | None:
        for alias in aliases:
            if alias in norm:
                return norm[alias]
        return None

    email_col = pick("email", "email id", "emailid", "email address", "ethara email", "ethara_email")
    status_col = pick("status", "id card", "id card status", "result", "state", "done/pending", "done / pending")
    if not email_col or not status_col:
        raise ValueError("The sheet must include an 'Email' column and a 'Status' (Done/Pending) column.")
    rows: list[dict[str, str]] = []
    for raw_row in reader:
        rows.append(
            {
                "email": (raw_row.get(email_col) or "").strip(),
                "status": (raw_row.get(status_col) or "").strip(),
            }
        )
    return rows


def apply_id_card_status_results(
    db: Session,
    *,
    rows: list[dict[str, str]],
    actor: User,
) -> dict[str, Any]:
    """Bulk-update candidate ID-card status by Ethara email from an uploaded sheet.

    "Done" marks the card as created/issued (only once the member has submitted their
    ID-card details); "Pending" reverts a previously-issued card back to outstanding.
    Rows that cannot be applied are reported back rather than aborting the whole upload.
    """
    candidates = db.scalars(
        select(Candidate).where(
            func.length(func.trim(func.coalesce(Candidate.ethara_email, ""))) > 0,
            Candidate.is_removed.is_(False),
        )
    ).all()
    by_email: dict[str, Candidate] = {}
    for candidate in candidates:
        for email in (candidate.ethara_email, candidate.personal_email):
            if email and email.strip():
                by_email.setdefault(email.strip().lower(), candidate)

    records_by_candidate = {
        record.candidate_id: record
        for record in db.scalars(select(CandidateIdCardForm)).all()
    }

    now = datetime.now(UTC)
    marked_done = 0
    marked_pending = 0
    not_found: list[str] = []
    skipped: list[dict[str, str]] = []

    for row in rows:
        email = (row.get("email") or "").strip()
        if not email:
            continue
        candidate = by_email.get(email.lower())
        if candidate is None:
            not_found.append(email)
            continue
        token = (row.get("status") or "").strip().lower()
        record = records_by_candidate.get(candidate.id)

        if token in _ID_CARD_DONE_TOKENS:
            if record is None or record.submitted_at is None:
                skipped.append({"email": email, "reason": "ID card details not submitted yet."})
                continue
            if record.it_completed_at is not None:
                continue  # already issued — idempotent no-op
            record.it_completed_at = now
            record.it_completed_by = actor.id
            db.add(record)
            log_audit(
                db,
                entity_type="candidate_id_card_form",
                entity_id=record.id,
                action="candidate_id_card_marked_done",
                actor=actor,
                candidate_id=candidate.id,
                new_value={"status": "done", "itCompletedBy": actor.id, "source": "status_sheet_upload"},
            )
            marked_done += 1
        elif token in _ID_CARD_PENDING_TOKENS:
            if record is None or record.it_completed_at is None:
                continue  # already not issued — nothing to revert
            record.it_completed_at = None
            record.it_completed_by = None
            db.add(record)
            log_audit(
                db,
                entity_type="candidate_id_card_form",
                entity_id=record.id,
                action="candidate_id_card_marked_pending",
                actor=actor,
                candidate_id=candidate.id,
                new_value={"status": "ready", "source": "status_sheet_upload"},
            )
            marked_pending += 1
        else:
            skipped.append({"email": email, "reason": f"Unrecognised status '{row.get('status')}'."})

    return {
        "markedDone": marked_done,
        "markedPending": marked_pending,
        "notFound": not_found,
        "skipped": skipped,
    }


def _advance_stage_no_regress(
    db: Session,
    *,
    candidate: Candidate,
    target: CandidateStage,
    actor: User | None,
    notes: str | None = None,
) -> bool:
    """Set a candidate's stage to ``target`` ONLY if it does not move them backward.

    The selection-form submit/validate endpoints can legitimately be hit again after a
    candidate has already advanced — e.g. a still-live form link is re-submitted, even
    post-onboarding. Re-applying the stage unconditionally used to regress a fully onboarded
    employee all the way back to "Selection Form Submitted". Here we no-op the stage change
    when the candidate is already in a LATER bucket, and record a StageLog when it does move.

    Returns True if the stage actually changed.
    """
    current = candidate.current_stage
    cur_bucket = stage_bucket(current)
    to_bucket = stage_bucket(target)
    # Refuse to move to an earlier bucket. If either stage is off-pipeline (unknown bucket),
    # fall back to applying the target to preserve prior behavior.
    if cur_bucket is not None and to_bucket is not None and to_bucket < cur_bucket:
        return False
    if current == target:
        return False
    _add_stage_log(db, candidate=candidate, from_stage=current, to_stage=target, actor=actor, notes=notes)
    candidate.current_stage = target
    candidate.current_status = stage_to_status(target)
    return True


def submit_selection_form(
    db: Session, *, selection_form: SelectionForm, form_data: dict[str, Any], actor: User
) -> SelectionForm:
    normalized_form_data = _normalize_selection_form_payload(form_data)
    selection_form.form_data = normalized_form_data
    selection_form.submitted_at = datetime.now(UTC)
    db.add(selection_form)
    candidate = db.get(Candidate, selection_form.candidate_id)
    if candidate:
        _sync_candidate_from_selection_form(candidate, normalized_form_data)
        _advance_stage_no_regress(
            db,
            candidate=candidate,
            target=CandidateStage.SELECTION_FORM_SUBMITTED,
            actor=actor,
            notes="Selection form submitted.",
        )
        db.add(candidate)
    log_audit(
        db,
        entity_type="selection_form",
        entity_id=selection_form.id,
        action="selection_form_submitted",
        actor=actor,
        candidate_id=selection_form.candidate_id,
        new_value={"submitted": True},
    )
    return selection_form


def validate_selection_form(db: Session, *, selection_form: SelectionForm, actor: User) -> SelectionForm:
    selection_form.validated_at = datetime.now(UTC)
    db.add(selection_form)
    candidate = db.get(Candidate, selection_form.candidate_id)
    if candidate:
        _advance_stage_no_regress(
            db,
            candidate=candidate,
            target=CandidateStage.SELECTION_FORM_VALIDATED,
            actor=actor,
            notes="Selection form validated.",
        )
        db.add(candidate)
    log_audit(
        db,
        entity_type="selection_form",
        entity_id=selection_form.id,
        action="selection_form_validated",
        actor=actor,
        candidate_id=selection_form.candidate_id,
        new_value={"validated": True},
    )
    return selection_form


def update_contract(db: Session, *, contract: Contract, payload: dict[str, Any], actor: User) -> Contract:
    for field, value in payload.items():
        setattr(contract, field, value)
    db.add(contract)
    if contract.status == ContractStatus.SIGNED:
        contract.signed_at = contract.signed_at or datetime.now(UTC)
        candidate = db.get(Candidate, contract.candidate_id)
        if candidate:
            candidate.current_stage = CandidateStage.CONTRACT_SIGNED
            candidate.current_status = stage_to_status(candidate.current_stage)
            # Auto code allocation gated off by default — codes come from the IT
            # bulk-register upload. See AUTO_EMPLOYEE_PROVISIONING.
            if get_settings().auto_employee_provisioning:
                from app.services.employees import assign_candidate_employee_code

                assign_candidate_employee_code(db, candidate)
            db.add(candidate)
    log_audit(
        db,
        entity_type="contract",
        entity_id=contract.id,
        action="contract_updated",
        actor=actor,
        candidate_id=contract.candidate_id,
        new_value={"status": contract.status.value},
    )
    return contract


def submit_compliance_form(
    db: Session, *, record: ComplianceForm, form_data: dict[str, Any], actor: User
) -> ComplianceForm:
    # Once a compliance form is verified it is locked. Re-submission (by the
    # candidate or anyone else) would silently revert a staff verification, so it is
    # rejected — staff must explicitly reopen the form before it can be resubmitted.
    if record.status == "verified" and not _has_any_role(actor, _EVALUATION_STAFF_ROLES):
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="This compliance form has already been verified and is locked. Ask an Admin/HR to reopen it.",
        )
    record.form_data = form_data
    record.status = "submitted"
    record.submitted_at = datetime.now(UTC)
    db.add(record)
    log_audit(
        db,
        entity_type="compliance_form",
        entity_id=record.id,
        action="compliance_form_submitted",
        actor=actor,
        candidate_id=record.candidate_id,
        new_value={"status": "submitted"},
    )
    return record


def verify_compliance_form(db: Session, *, record: ComplianceForm, actor: User) -> ComplianceForm:
    record.status = "verified"
    record.verified_at = datetime.now(UTC)
    db.add(record)
    # ComplianceForm has no verified_by column, so the verifier identity is persisted
    # via the audit log (performed_by / verifiedBy in new_value). See cross_cutting_notes
    # for adding a dedicated column + migration if a first-class field is wanted.
    log_audit(
        db,
        entity_type="compliance_form",
        entity_id=record.id,
        action="compliance_form_verified",
        actor=actor,
        candidate_id=record.candidate_id,
        new_value={
            "status": "verified",
            "verifiedBy": actor.id,
            "verifiedByName": actor.name,
            "verifiedAt": record.verified_at.isoformat() if record.verified_at else None,
        },
    )
    return record


def run_resume_screening(
    db: Session,
    *,
    candidate_id: str,
    job_description: str | None = None,
    actor: User | None = None,
    request: Request | None = None,
) -> Candidate:
    candidate = db.scalar(
        select(Candidate)
        .options(joinedload(Candidate.position), selectinload(Candidate.documents))
        .where(Candidate.id == candidate_id)
    )
    if candidate is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Candidate not found")

    position = candidate.position

    # If the resume text wasn't captured at upload time, extract it now from the stored
    # file (PyMuPDF for PDF, python-docx for DOC) so screening has real content.
    if not (candidate.resume_text and candidate.resume_text.strip()) and candidate.resume_url:
        try:
            from pathlib import Path as _Path
            from app.api.routes.candidates import _extract_resume_text
            from app.core.config import get_settings
            settings = get_settings()
            _url = candidate.resume_url
            _rel = _url.removeprefix("/uploads/") if _url.startswith("/uploads/") else _url
            _fpath = settings.local_storage_path / _rel
            if _fpath.exists():
                _text = _extract_resume_text(_fpath.read_bytes(), _Path(_fpath).suffix, None)
                if _text and _text.strip():
                    candidate.resume_text = _text[:20000]
                    db.add(candidate)
        except Exception:
            pass

    llm = LLMService()
    result = llm.screen_resume(
        candidate_name=candidate.full_name,
        resume_text=candidate.resume_text,
        resume_url=candidate.resume_url,
        job_title=position.title if position else None,
        job_description=job_description or (position.description if position else None),
        screening_prompt=position.screening_prompt if position else None,
    )

    previous_stage = candidate.current_stage
    previous_status = candidate.current_status
    candidate.resume_score = float(result.get("score", result.get("matchScore", 0)))
    candidate.resume_summary = result.get("summary")
    payload = dict(candidate.screening_payload or {})
    payload.update(result)
    payload["score"] = candidate.resume_score
    payload["matchScore"] = float(result.get("matchScore", result.get("score", candidate.resume_score or 0)))
    payload["lastScreenedAt"] = datetime.now(UTC).isoformat()
    latest_resume = _latest_resume_document(candidate)
    if latest_resume is not None:
        payload["resumeDocumentId"] = latest_resume.id
    recommendation = result.get("recommendation", "pending")
    normalized_recommendation = _normalize_screening_recommendation(recommendation)
    payload["recommendation"] = normalized_recommendation or recommendation
    candidate.screening_payload = payload

    if normalized_recommendation in {None, "pending"}:
        candidate.llm_status = "needs_configuration"
        target_stage = CandidateStage.RESUME_SCREENING_PENDING
    else:
        candidate.llm_status = "completed"
        target_stage = _candidate_stage_for_recommendation(normalized_recommendation)
    db.add(candidate)
    # No-regress guard: a (re-)run of screening must never pull a candidate who has already
    # advanced past the screening bucket back to a screening stage. This previously reset
    # candidates who had already submitted their selection form. _advance_stage_no_regress
    # only moves (and logs) the stage when the target is not an earlier pipeline bucket.
    _advance_stage_no_regress(
        db,
        candidate=candidate,
        target=target_stage,
        actor=actor,
        notes=f"Resume screening completed with recommendation: {payload.get('recommendation') or 'pending'}.",
    )
    log_audit(
        db,
        entity_type="resume_screening",
        entity_id=candidate.id,
        action="resume_screening_completed",
        actor=actor,
        request=request,
        candidate_id=candidate.id,
        old_value={"stage": previous_stage.value, "status": previous_status},
        new_value={
            "stage": candidate.current_stage.value,
            "status": candidate.current_status,
            "score": candidate.resume_score,
            "matchScore": payload.get("matchScore"),
            "recommendation": payload.get("recommendation"),
            "summary": candidate.resume_summary,
        },
    )
    return candidate


def override_resume_screening(
    db: Session,
    *,
    candidate: Candidate,
    recommendation: str,
    reason: str,
    actor: User,
    request: Request | None = None,
) -> Candidate:
    normalized_recommendation = _normalize_screening_recommendation(recommendation)
    if normalized_recommendation not in {"shortlisted", "rejected", "needs_review"}:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Recommendation must be shortlisted, rejected, or needs_review.",
        )

    previous_stage = candidate.current_stage
    previous_status = candidate.current_status
    payload = dict(candidate.screening_payload or {})
    payload["recommendation"] = normalized_recommendation
    payload["manualOverride"] = {
        "recommendation": normalized_recommendation,
        "reason": reason.strip(),
        "performedBy": actor.id,
        "performedByName": actor.name,
        "performedAt": datetime.now(UTC).isoformat(),
    }
    payload["lastScreenedAt"] = payload.get("lastScreenedAt") or datetime.now(UTC).isoformat()
    candidate.screening_payload = payload
    candidate.llm_status = "completed"
    target_stage = _candidate_stage_for_recommendation(normalized_recommendation)
    # No-regress guard: a manual screening override must NOT yank a candidate who
    # has already advanced past resume screening (e.g. already submitted/validated
    # their selection form, or further) back into the screening/assessment funnel.
    # That regression is exactly what stranded already-submitted candidates back at
    # "Selection Form Sent" and kept re-prompting them to submit. The screening
    # verdict is still recorded in screening_payload above and the audit log below;
    # we only refuse to move the pipeline stage backward. An explicit "rejected"
    # verdict is still allowed to regress, since that is a deliberate removal.
    cur_bucket = stage_bucket(candidate.current_stage)
    to_bucket = stage_bucket(target_stage)
    is_regression = cur_bucket is not None and to_bucket is not None and to_bucket < cur_bucket
    if normalized_recommendation == "rejected" or not is_regression:
        candidate.current_stage = target_stage
        candidate.current_status = stage_to_status(target_stage)
        _add_stage_log(
            db,
            candidate=candidate,
            from_stage=previous_stage,
            to_stage=candidate.current_stage,
            actor=actor,
            notes=f"Manual screening override: {reason.strip()}",
        )
    db.add(candidate)
    log_audit(
        db,
        entity_type="resume_screening",
        entity_id=candidate.id,
        action="resume_screening_overridden",
        actor=actor,
        request=request,
        candidate_id=candidate.id,
        old_value={"stage": previous_stage.value, "status": previous_status},
        new_value={
            "stage": candidate.current_stage.value,
            "status": candidate.current_status,
            "recommendation": normalized_recommendation,
            "reason": reason.strip(),
        },
    )
    return candidate


def run_document_ocr(db: Session, *, document_id: str) -> Document:
    document = db.get(Document, document_id)
    if document is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    ocr = OCRService()
    llm = LLMService()
    ocr_payload = ocr.extract(document_type=document.type, file_url=document.file_url)
    llm_payload = llm.extract_document(document_type=document.type, extracted_text=ocr_payload)
    document.ocr_status = "completed"
    document.ocr_provider = ocr_payload.get("provider")
    document.extracted_data = ocr_payload
    document.llm_extracted_data = llm_payload
    db.add(document)
    return document


def complete_it_request(db: Session, *, record: ITRequest, created_email: str, actor: User) -> ITRequest:
    record.created_email = created_email
    record.status = "completed"
    record.completed_at = datetime.now(UTC)
    db.add(record)
    candidate = db.get(Candidate, record.candidate_id)
    if candidate:
        candidate.ethara_email = created_email
        if candidate.current_stage in {
            CandidateStage.CONTRACT_SIGNED,
            CandidateStage.INDUCTION_COMPLETED,
        }:
            candidate.current_stage = CandidateStage.IT_EMAIL_CREATED
            candidate.current_status = stage_to_status(candidate.current_stage)
        db.add(candidate)
    log_audit(
        db,
        entity_type="it_request",
        entity_id=record.id,
        action="it_request_completed",
        actor=actor,
        candidate_id=record.candidate_id,
        new_value={"createdEmail": created_email},
    )
    return record


def pick_responsible_user_for_stage(db: Session, *, stage: CandidateStage) -> User | None:
    preferred_roles = [Role.HR, Role.TA, Role.LEADERSHIP, Role.ADMIN] if stage != CandidateStage.IT_EMAIL_CREATED else [Role.IT_TEAM, Role.LEADERSHIP, Role.ADMIN]
    active_users = list(db.scalars(select(User).where(User.is_active.is_(True)).order_by(User.created_at.asc())).all())
    for role in preferred_roles:
        user = next((item for item in active_users if _has_any_role(item, {role})), None)
        if user:
            return user
    return active_users[0] if active_users else None


def run_sla_checks(db: Session) -> list[Escalation]:
    now = datetime.now(UTC)
    created: list[Escalation] = []
    for stage, hours in SLA_HOURS.items():
        deadline = now - timedelta(hours=hours)
        candidates = list(
            db.scalars(
                select(Candidate)
                .options(selectinload(Candidate.escalations))
                .where(Candidate.current_stage == stage, Candidate.updated_at <= deadline, Candidate.is_removed.is_(False))
            )
        )
        for candidate in candidates:
            open_escalation = next(
                (item for item in candidate.escalations if item.stage == stage.value and item.status == EscalationStatus.OPEN),
                None,
            )
            delay_hours = int((now - candidate.updated_at).total_seconds() // 3600)
            level = 3 if delay_hours > hours * 2 else 2 if delay_hours > hours * 1.5 else 1
            delayed_by = f"{delay_hours // 24}d {delay_hours % 24}h" if delay_hours >= 24 else f"{delay_hours}h"

            if open_escalation:
                open_escalation.delayed_by = delayed_by
                open_escalation.escalation_level = max(open_escalation.escalation_level, level)
                db.add(open_escalation)
                continue

            responsible = pick_responsible_user_for_stage(db, stage=stage)
            if responsible is None:
                continue
            escalation = Escalation(
                candidate_id=candidate.id,
                stage=stage.value,
                responsible_user_id=responsible.id,
                sla_deadline=deadline,
                delayed_by=delayed_by,
                escalation_level=level,
                status=EscalationStatus.OPEN,
            )
            db.add(escalation)
            db.flush()
            create_notification(
                db,
                user_id=responsible.id,
                candidate_id=candidate.id,
                title=f"SLA breach: {candidate.full_name}",
                message=f"{candidate.full_name} is delayed in {stage_to_status(stage)} by {delayed_by}.",
                type_=NotificationType.WARNING if level < 3 else NotificationType.ERROR,
            )
            created.append(escalation)
    return created
